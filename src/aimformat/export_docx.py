"""Export: .aim → DOCX (``python-docx``, behind the ``aimformat[docx]`` extra).

The exporter walks the accepted body and maps chunks onto Word constructs.
Its distinguishing feature — a pattern adapted from a previous project's
editor, where the DOCX download replayed the change log with a caller-chosen
default — is how the **pending lane** is handled:

``pending="tracked"`` (default)
    Pending proposals are emitted as *real Word revision markup*
    (``w:ins``/``w:del``), attributed to the proposing actor and timestamped,
    so a counterparty opening the file in Word sees reviewable tracked
    changes. Granularity is the chunk (whole-block delete + insert), matching
    what the format stores; word-level diffs are a viewer concern. This
    covers body chunks, chunks inside slides, list items and whole list
    containers, table rows and whole table containers, and anchored adds.
``pending="accept-all"`` / ``pending="reject-all"``
    The pending lane is resolved on a throwaway copy of the document —
    through the ordinary accept/reject machinery, decided by the ``external``
    actor ``docx-export`` — and the resolved body is exported clean.

Slides linearize (Word has no fixed canvas): a page break, then the slide's
chunks as flowing blocks in reading order — geometry and classes drop; text,
structure, marks, and the per-chunk pending lane survive. The same
degradation contract as the Markdown exporter, on pages instead of ``---``.
A faithful canvas export is the PDF's job (and a future PPTX exporter's).

Not represented in v0.1: ``move`` proposals, ``aim:theme``/``aim:doc``
proposals, and proposals targeting a *whole slide* (all export as unchanged
current content), plus hyperlink relationships (links render as text with
the URL in parentheses). These are deliberate scope cuts, not oversights.
"""

from __future__ import annotations

import base64
import io
import re
from collections.abc import Mapping
from pathlib import Path
from typing import TYPE_CHECKING

from .document import AimDocument, Proposal
from .dom import Element, Text, parse_fragment
from .errors import InvalidOperation
from .events import external
from .paint import BorderSide, Paint, PaintResolver, brand_defaults
from .registry import REGISTRY

if TYPE_CHECKING:  # pragma: no cover
    pass

__all__ = ["to_docx"]

_PENDING_MODES = ("tracked", "accept-all", "reject-all")
_MONO = "Consolas"

_BOLD_TAGS = {"strong", "b"}
_ITALIC_TAGS = {"em", "i"}
_HEADINGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
_INLINE_TAGS = {
    "strong",
    "b",
    "em",
    "i",
    "u",
    "s",
    "mark",
    "code",
    "sub",
    "sup",
    "a",
    "img",
    "br",
    "span",
    "svg",
}


def _require_docx():
    try:
        import docx  # noqa: F401

        return docx
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise ImportError(
            "DOCX export needs python-docx. Install the extra:\n    pip install 'aimformat[docx]'"
        ) from exc


# --------------------------------------------------------------------------
# inline content -> (text, formatting) run specs
#
# Paint is not read off the element here: `aimformat.paint` has already
# resolved the whole tree against the generated stylesheet, so a leaf emitter
# looks its element up instead of re-deriving a cascade it would get subtly
# wrong (and each of ~40 leaves would get wrong differently).


def _palette_of(doc: AimDocument) -> dict[str, str]:
    """The document's brand slots over the registry defaults."""
    theme = dict(getattr(doc, "theme", None) or {})
    defaults = brand_defaults()
    return {**defaults, **{k: v for k, v in theme.items() if k in defaults}}


def _paint_fmt(paint: Paint, fmt: dict, *, inline: bool) -> dict:
    """*fmt* with this element's run-level paint folded in.

    ``inline`` distinguishes an inline element inside a block (whose own
    background and border become RUN properties) from the block itself
    (whose box becomes paragraph or cell properties) — otherwise a tinted
    paragraph would shade its own paragraph AND every run inside it.
    """
    out = dict(fmt)
    # colour inherits, so the record already carries the effective value —
    # including "nothing", which must CLEAR an inherited fmt colour (a link's
    # own base-layer colour beats the red it sits in)
    if paint.color:
        out["color"] = paint.color
    else:
        out.pop("color", None)
    if inline:
        if paint.own_background:
            out["shade"] = paint.own_background
        elif paint.background is None:
            # A base-layer opaque background (for example `code`) hides an
            # ancestor box in the browser. Drop a tracked block's inherited
            # run-shading approximation here too.
            out.pop("shade", None)
        border = paint.any_border
        if border is not None:
            out["border"] = border
    return out


def _first_family(stack: str | None) -> str | None:
    """The first concrete family of a font-stack slot ("Georgia, serif" →
    "Georgia"), or None. Word styles name one face, not a CSS stack."""
    if not stack:
        return None
    first = stack.split(",")[0].strip().strip("'\"")
    return first or None


def _run_typography(el: Element) -> dict:
    """A run element's literal typography for export: font size in points and
    font family. Class-driven size (the type scale) resolves through the
    normative pt table; an inline ``font-size``/``font-family`` overrides it,
    mirroring CSS specificity."""
    out: dict = {}
    for token in (el.get("class") or "").split():
        if token.startswith("text-"):
            pt = REGISTRY.type_scale_pt.get(token[len("text-") :])
            if pt:
                out["size_pt"] = float(pt)
        elif token == "uppercase":
            out["all_caps"] = True  # the importer's caps mapping, reversed
    for piece in (el.get("style") or "").split(";"):
        prop, sep, val = piece.partition(":")
        if not sep:
            continue
        prop, val = prop.strip(), val.strip()
        if prop == "font-size" and val.endswith("pt"):
            try:
                out["size_pt"] = float(val[:-2])
            except ValueError:
                pass
        elif prop == "font-family" and val:
            # the grammar allows a stack ("Segoe UI, sans-serif"); Word run
            # properties name exactly one face, so take the first family
            family = _first_family(val)
            if family:
                out["font_name"] = family
    return out


def _alignment_of(el: Element):
    """A block's alignment class → a Word paragraph alignment, or None."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    mapping = {
        "text-left": WD_ALIGN_PARAGRAPH.LEFT,
        "text-center": WD_ALIGN_PARAGRAPH.CENTER,
        "text-right": WD_ALIGN_PARAGRAPH.RIGHT,
        "text-justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    for token in (el.get("class") or "").split():
        if token in mapping:
            return mapping[token]
    return None


def _runs_of(
    el: Element,
    fmt: dict | None,
    paint: PaintResolver,
    *,
    inline: bool = False,
) -> list[dict]:
    """Run specs for *el*'s content.

    *paint* is required rather than defaulted: a resolver is the only thing
    that knows this element's computed colour, and an omitted one would
    silently export an unpainted document. ``inline`` is False for the block
    whose paragraph carries the box, True for everything nested inside it.
    """
    fmt = _paint_fmt(paint.of(el), dict(fmt or {}), inline=inline)
    # literal typography inherits into descendants like colour does: fold this
    # element's own size/family over what it inherited, children then override
    fmt.update(_run_typography(el))
    tag = el.tag
    if tag in _BOLD_TAGS:
        fmt["bold"] = True
    elif tag in _ITALIC_TAGS:
        fmt["italic"] = True
    elif tag == "u":
        fmt["underline"] = True
    elif tag == "s":
        fmt["strike"] = True
    elif tag == "mark":
        fmt["highlight"] = True
    elif tag == "code":
        fmt["mono"] = True
    elif tag == "sub":
        fmt["subscript"] = True
    elif tag == "sup":
        fmt["superscript"] = True

    runs: list[dict] = []
    for child in el.children:
        if isinstance(child, Text):
            if child.data:
                runs.append({"text": child.data, **fmt})
        elif isinstance(child, Element):
            if child.tag == "br":
                runs.append({"break": True, **fmt})
            elif child.tag == "a":
                link_fmt = _paint_fmt(paint.of(child), fmt, inline=True)
                runs += _runs_of(child, {**fmt, "underline": True}, paint, inline=True)
                href = child.get("href") or ""
                if href and not href.startswith("#"):
                    runs.append({"text": f" ({href})", **link_fmt})
            elif child.tag == "img":
                # inline image: honest placeholder (figure-fallback style),
                # URL kept unless it's an embedded data blob
                label = f"[image: {child.get('alt') or 'image'}]"
                src = child.get("src") or ""
                if src and not src.startswith("data:"):
                    label += f" ({src})"
                runs.append({"text": label, **fmt, "italic": True})
            else:
                runs += _runs_of(child, fmt, paint, inline=True)
    return runs


def _block_runs(el: Element, paint: PaintResolver) -> list[dict]:
    """Run specs for one block: a page-break chunk is a single page-break
    run (it has no text runs — without this, the tracked lane would emit a
    pending break as an empty revision paragraph)."""
    if el.tag == "aim-page-break":
        return [{"page_break": True}]
    if el.tag == "hr":
        spec = {"text": "—" * 12}
        border = paint.of(el).any_border
        if border is not None:
            spec["color"] = border.color
        return [spec]
    return _runs_of(el, None, paint)


def _clean_runs_and_box(el: Element, paint: PaintResolver) -> tuple[list[dict], Paint]:
    """Clean-export runs plus the box paint Word can apply safely.

    A Word paragraph/cell shade sits behind every run. When a descendant's
    opaque base background (for example ``code``) stops the AIM ancestor box,
    that one box shade would leak through it. Fall back to run shading for the
    parts where the authored box really shows and leave the base-backed run
    clear for the recipient template.
    """
    box = paint.of(el)
    if box.background is None or not any(
        descendant is not el and paint.of(descendant).background is None for descendant in el.iter()
    ):
        return _runs_of(el, None, paint), box
    runs = _runs_of(el, {"shade": box.background}, paint)
    return runs, Paint(
        color=box.color,
        background=None,
        own_background=box.own_background,
        borders=box.borders,
    )


def _tracked_box_fmt(el: Element, paint: PaintResolver, *, border: bool = True) -> dict:
    """Run-level approximation of a block/cell box inside a revision.

    Word has one paragraph/cell property outside `w:ins`/`w:del`. Putting a
    proposed box there makes rejection keep the proposed paint. Revision runs
    can carry their own shading and one whole-run border, so old and new paint
    remain independently reviewable.
    """
    resolved = paint.of(el)
    out: dict = {}
    if resolved.background:
        out["shade"] = resolved.background
    side = resolved.any_border if border else None
    if side is not None:
        out["border"] = side
    return out


def _tracked_runs(el: Element, paint: PaintResolver) -> list[dict]:
    fmt = _tracked_box_fmt(el, paint)
    runs = _runs_of(el, fmt or None, paint)
    return runs or ([{"text": "", **fmt}] if fmt else [])


def _tracked_block_runs(el: Element, paint: PaintResolver) -> list[dict]:
    if el.tag == "aim-page-break":
        return _block_runs(el, paint)
    if el.tag == "hr":
        # The clean degradation paints the em-dash rule's ink rather than
        # adding a second Word border. Keep that shape inside revisions.
        fmt = _tracked_box_fmt(el, paint, border=False)
        return [{**run, **fmt} for run in _block_runs(el, paint)]
    fmt = _tracked_box_fmt(el, paint)
    runs = _runs_of(el, fmt or None, paint)
    return runs or ([{"text": "", **fmt}] if fmt else [])


def _apply_runs(paragraph, runs: list[dict]) -> None:
    for spec in runs:
        run = paragraph.add_run()
        if spec.get("page_break"):
            from docx.enum.text import WD_BREAK

            run.add_break(WD_BREAK.PAGE)
            continue
        if spec.get("break"):
            run.add_break()
            continue
        run.text = spec.get("text", "")
        _format_run(run, spec)


def _format_run(run, spec: dict) -> None:
    run.bold = spec.get("bold") or None
    run.italic = spec.get("italic") or None
    run.underline = True if spec.get("underline") else None
    if spec.get("strike"):
        run.font.strike = True
    if spec.get("mono"):
        run.font.name = _MONO
    if spec.get("font_name"):  # explicit family wins over the mono default
        run.font.name = spec["font_name"]
    if spec.get("all_caps"):
        run.font.all_caps = True
    if spec.get("size_pt"):
        from docx.shared import Pt

        run.font.size = Pt(spec["size_pt"])
    if spec.get("subscript"):
        run.font.subscript = True
    if spec.get("superscript"):
        run.font.superscript = True
    if spec.get("highlight"):
        from docx.enum.text import WD_COLOR_INDEX

        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if spec.get("color"):
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor.from_string(spec["color"])
    if spec.get("shade"):
        # real shading, not Word's 16-value highlight enum — the format's
        # background is an arbitrary RGB and the enum would round it
        _set_shading(run._r.get_or_add_rPr(), spec["shade"], ordered_in="w:rPr")
    if spec.get("border"):
        _set_run_border(run._r.get_or_add_rPr(), spec["border"])


# --------------------------------------------------------------------------
# paint -> OOXML properties
#
# Every helper is idempotent (it replaces its own element rather than
# appending a second one), because tracked runs are built, formatted, and
# then detached — and a paragraph can be reached twice through a grouping
# element.

_EIGHTHS_PER_PX = 6  # 1 px = 0.75 pt = 6 eighths of a point

# WordprocessingML property children are a SEQUENCE, not a set: an element
# appended out of order makes the file invalid. python-docx only knows the
# positions of the children it models itself, so the ones added here name
# their own successors and slot in ahead of the first present.
_SUCCESSORS = {
    "w:rPr": {
        "w:bdr": ("w:shd", "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang"),
        "w:shd": ("w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang"),
    },
    "w:pPr": {
        "w:pBdr": ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"),
        "w:shd": ("w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"),
    },
    "w:tcPr": {
        "w:tcBorders": ("w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:vAlign"),
        "w:shd": ("w:noWrap", "w:tcMar", "w:textDirection", "w:vAlign"),
    },
}


def _border_size(width_px: float) -> int:
    return max(2, min(96, round(width_px * _EIGHTHS_PER_PX)))


def _fresh_child(props, tag: str, *, ordered_in: str | None = None):
    """An empty *tag* child of *props*, replacing any existing one.

    Replacing rather than appending keeps every helper idempotent: tracked
    runs are formatted, detached and re-formatted, and a grouping element can
    reach the same paragraph twice.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for existing in props.findall(qn(tag)):
        props.remove(existing)
    el = OxmlElement(tag)
    successors = _SUCCESSORS.get(ordered_in or "", {}).get(tag)
    if successors:
        props.insert_element_before(el, *successors)
    else:
        props.append(el)
    return el


def _set_shading(props, fill: str, *, ordered_in: str) -> None:
    from docx.oxml.ns import qn

    shd = _fresh_child(props, "w:shd", ordered_in=ordered_in)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_border(parent, tag: str, side: BorderSide, *, ordered_in: str | None = None) -> None:
    from docx.oxml.ns import qn

    el = _fresh_child(parent, tag, ordered_in=ordered_in)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(_border_size(side.width_px)))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), side.color)


def _set_run_border(rpr, side: BorderSide) -> None:
    """Word has ONE border for a whole run (`w:rPr/w:bdr`) — per-side borders
    exist only on paragraphs and cells. A side utility on an inline element
    therefore colours the whole box; documented degradation, not a bug."""
    _set_border(rpr, "w:bdr", side, ordered_in="w:rPr")


def _set_box(props, paint: Paint, *, props_tag: str, borders_tag: str) -> None:
    """Background and per-side borders onto a paragraph's `w:pPr` or a
    cell's `w:tcPr`."""
    if paint.borders:
        container = _fresh_child(props, borders_tag, ordered_in=props_tag)
        for name in ("top", "left", "bottom", "right"):  # schema order
            side = paint.borders.get(name)
            if side is not None:
                _set_border(container, f"w:{name}", side)
    if paint.background:
        _set_shading(props, paint.background, ordered_in=props_tag)


def _ink(paint: Paint) -> dict:
    """The run-spec fields for a leaf built without ``_runs_of``."""
    return {"color": paint.color} if paint.color else {}


def _paint_paragraph(paragraph, paint: Paint) -> None:
    if paint.background or paint.borders:
        _set_box(paragraph._p.get_or_add_pPr(), paint, props_tag="w:pPr", borders_tag="w:pBdr")


def _paint_cell(cell, paint: Paint) -> None:
    if paint.background or paint.borders:
        _set_box(cell._tc.get_or_add_tcPr(), paint, props_tag="w:tcPr", borders_tag="w:tcBorders")


# --------------------------------------------------------------------------
# tracked-changes OXML
class _Revisions:
    """Builds w:ins / w:del wrappers with stable ids and attribution."""

    def __init__(self):
        self._next = 1

    def _attrs(self, el, author: str, date: str) -> None:
        from docx.oxml.ns import qn

        el.set(qn("w:id"), str(self._next))
        el.set(qn("w:author"), author)
        el.set(qn("w:date"), date or "2026-01-01T00:00:00Z")
        self._next += 1

    def ins(self, paragraph, runs: list[dict], author: str, date: str) -> None:
        from docx.oxml import OxmlElement

        wrap = OxmlElement("w:ins")
        self._attrs(wrap, author, date)
        for spec in runs:
            wrap.append(self._make_run(paragraph, spec, deleted=False))
        paragraph._p.append(wrap)

    def dele(self, paragraph, runs: list[dict], author: str, date: str) -> None:
        from docx.oxml import OxmlElement

        wrap = OxmlElement("w:del")
        self._attrs(wrap, author, date)
        for spec in runs:
            wrap.append(self._make_run(paragraph, spec, deleted=True))
        paragraph._p.append(wrap)

    def replace_ins(
        self,
        paragraph,
        runs: list[dict],
        prior_author: str,
        prior_date: str,
        author: str,
        date: str,
    ) -> None:
        """Replace an earlier insertion with a later insertion of the same
        content. Rejecting the replacement restores the earlier revision;
        accepting it removes that copy and keeps the later one."""
        from docx.oxml import OxmlElement

        prior = OxmlElement("w:ins")
        self._attrs(prior, prior_author, prior_date)
        for spec in runs:
            prior.append(self._make_run(paragraph, spec, deleted=False))
        removed = OxmlElement("w:del")
        self._attrs(removed, author, date)
        removed.append(prior)
        paragraph._p.append(removed)
        self.ins(paragraph, runs, author, date)

    def row_ins(self, tr, author: str, date: str) -> None:
        self._row_change(tr, "w:ins", author, date)

    def row_dele(self, tr, author: str, date: str) -> None:
        self._row_change(tr, "w:del", author, date)

    def _row_change(self, tr, tag: str, author: str, date: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        props = tr.find(qn("w:trPr"))
        if props is None:
            props = OxmlElement("w:trPr")
            tr.insert(0, props)
        change = OxmlElement(tag)
        self._attrs(change, author, date)
        props.append(change)

    @staticmethod
    def _make_run(paragraph, spec: dict, *, deleted: bool):
        # build a real run for formatting, then detach and retag its text
        from docx.oxml.ns import qn

        run = paragraph.add_run(spec.get("text", ""))
        if spec.get("page_break"):
            from docx.enum.text import WD_BREAK

            run.add_break(WD_BREAK.PAGE)
        elif spec.get("break"):
            run.add_break()
        else:
            _format_run(run, spec)
        r = run._r
        r.getparent().remove(r)
        if deleted:
            for t in r.findall(qn("w:t")):
                t.tag = qn("w:delText")
                t.set(qn("xml:space"), "preserve")
        return r


def _actor_label(author) -> str:
    if author.type == "agent":
        return f"agent:{author.model or author.id or 'unknown'}"
    if author.type == "human":
        return author.id or "human"
    return author.id or "external"


def _style_for(tag: str) -> str | None:
    if tag in _HEADINGS:
        return f"Heading {tag[1]}"
    if tag == "blockquote":
        return "Quote"
    if tag == "li":
        return "List Bullet"
    return None


def _block_children(el: Element, paint: PaintResolver) -> list[Element]:
    """The block-level pieces of one chunk (a grouping block's children,
    recursively; a slide's children recursively, so a pending whole-slide
    add linearizes per block like an accepted slide; the element itself
    otherwise). Each block child of div/section/blockquote becomes its own
    paragraph — treating them as inline fused 'Quote one.Quote two.' into
    one paragraph — and direct inline content wraps in a synthetic block
    so it is never dropped."""
    if el.tag in ("section", "div"):
        out = _unpack_group(el, "p", paint)
    elif el.tag == "blockquote":
        out = _unpack_group(el, "blockquote", paint)
    elif el.tag == "aim-slide":
        out = []
        for child in el.elements():
            out.extend(_block_children(child, paint))
    else:
        return [el]
    _carry_group_borders(out, paint.of(el).borders, paint)
    return out


def _carry_group_borders(
    blocks: list[Element], borders: Mapping[str, BorderSide], paint: PaintResolver
) -> None:
    """Approximate one grouping box on every Word block it emits.

    Lists emit one paragraph per item and tables one box per cell; other
    block pieces map directly. A descendant's own side wins when Word cannot
    represent both the group and child border on that side.
    """
    for block in blocks:
        if block.tag in ("ul", "ol"):
            targets = block.find_all(lambda node: node.tag == "li")
        elif block.tag == "table":
            targets = block.find_all(lambda node: node.tag in ("td", "th"))
        else:
            targets = [block]
        for target in targets:
            paint.overlay_borders(target, borders)


def _unpack_group(el: Element, wrap_tag: str, paint: PaintResolver) -> list[Element]:
    """One block per block child of a grouping element; runs of direct
    inline content wrap in a synthetic *wrap_tag* element. Paragraphs
    inside a blockquote re-wrap as *wrap_tag* so they keep the Quote
    style."""
    out: list[Element] = []
    run: list = []

    def flush() -> None:
        if run:
            if any(isinstance(n, Element) or (isinstance(n, Text) and n.data.strip()) for n in run):
                out.append(_wrapped(wrap_tag, run, el, paint))
            run.clear()

    for child in el.children:
        if isinstance(child, Element) and child.tag not in _INLINE_TAGS:
            flush()
            if child.tag == "p" and wrap_tag != "p":
                out.append(_wrapped(wrap_tag, child.children, child, paint))
            else:
                out.extend(_block_children(child, paint))
        else:
            run.append(child)
    flush()
    return out


def _wrapped(tag: str, children: list, source: Element, paint: PaintResolver) -> Element:
    """A synthetic block wrapping loose inline content from a grouping element.

    ``source`` is the element the wrapper stands in for: the group itself for
    loose content, or the ``<p>`` being re-tagged. The wrapper adopts that
    element's computed paint, so paint declared on a group —
    ``<blockquote style="color:#ff69b4">Quote</blockquote>`` — is not lost
    before runs are built. Nothing is copied onto the SOURCE, and no attribute
    is copied onto the wrapper: an earlier version wrote resolved classes back
    into the tree and corrupted a document mid-export.
    """
    wrap = Element(tag)
    wrap.children = list(children)
    paint.adopt(wrap, source)
    return wrap


def _row_shape(tr: Element) -> list[tuple[int, int]]:
    """A table row's structural signature: per-cell (colspan, rowspan)."""
    return [
        (int(c.get("colspan") or 1), int(c.get("rowspan") or 1))
        for c in tr.elements()
        if c.tag in ("td", "th")
    ]


_PX_VALUE_RE = re.compile(r"^(\d+(?:\.\d+)?)px$")


def _style_px(el: Element, prop: str) -> float | None:
    """A px-valued inline-style property of *el*, if declared."""
    for piece in (el.get("style") or "").split(";"):
        if ":" not in piece:
            continue
        name, val = (s.strip() for s in piece.split(":", 1))
        if name == prop:
            m = _PX_VALUE_RE.match(val)
            return float(m.group(1)) if m else None
    return None


# --------------------------------------------------------------------------
class _Exporter:
    def __init__(self, doc: AimDocument, docx_mod):
        self.aim = doc
        self.docx = docx_mod
        # Resolved once per export. Held on the instance, not module state:
        # exports run in worker threads (the backend calls to_docx via
        # asyncio.to_thread), so a shared global would race between documents
        # with different themes.
        self.palette = _palette_of(doc)
        # Computed paint for each live construct, once. Every emitter reads
        # this instead of deriving colour from its own element — and it holds
        # records by object identity, so nothing is ever written back into the
        # document being exported. Structural body attributes stay outside
        # the addressable rendering state.
        self.paint = PaintResolver(self.palette)
        for construct in doc._state.constructs():
            self.paint.resolve(construct)
        self.out = docx_mod.Document()
        self.rev = _Revisions()
        # Set after a slide: True means accepted structure owns the next
        # break; a Proposal means the pending slide owns it.
        self._break_before_next: bool | Proposal = False
        self.pending_mod: dict[str, Proposal] = {}
        self.pending_del: dict[str, Proposal] = {}
        # adds keyed by (container, after) — every container, not just body
        self.adds_by_anchor: dict[tuple[str, str | None], list[Proposal]] = {}
        for p in doc.proposals:
            if p.action == "modify" and p.target and p.target not in ("aim:theme", "aim:doc"):
                self.pending_mod[p.target] = p
            elif p.action == "delete" and p.target:
                self.pending_del[p.target] = p
            elif p.action == "add":
                key = (p.anchor_container or "body", p.anchor_after)
                self.adds_by_anchor.setdefault(key, []).append(p)

    # -- top level -----------------------------------------------------------
    def run(self) -> None:
        title = self.aim.title
        if title:
            self.out.core_properties.title = title
        self._apply_theme_fonts()
        self._apply_page_setup()
        self._emit_anchored_adds("body", None)
        for construct in self.aim._state.constructs():
            self.emit_construct(construct)
            cid = construct.chunk_id or construct.container_id
            self._emit_anchored_adds("body", cid)
        # anything still unemitted (adds into containers that are themselves
        # pending-deleted) surfaces at the end rather than being silently
        # dropped
        for props in list(self.adds_by_anchor.values()):
            for prop in props:
                self._emit_add_paragraphs(prop)
        self.adds_by_anchor.clear()

    def _apply_theme_fonts(self) -> None:
        """The document's theme font-stack slots → the exported document's
        style fonts: ``--aim-font-body`` becomes Normal's face, and
        ``--aim-font-heading`` the face of every Heading/Title style. The
        first concrete family of each stack is used (Word names one face, not
        a CSS stack); colour slots are already handled by the paint resolver.
        """
        theme = dict(getattr(self.aim, "theme", None) or {})
        body = _first_family(theme.get("--aim-font-body"))
        heading = _first_family(theme.get("--aim-font-heading"))
        if not body and not heading:
            return
        styles = self.out.styles
        names = {s.name for s in styles}
        if body and "Normal" in names:
            styles["Normal"].font.name = body
        if heading:
            for name in names:
                if name.startswith("Heading") or name == "Title":
                    styles[name].font.name = heading

    def _apply_page_setup(self) -> None:
        """The document's page setup → Word section properties, from the
        same resolution the PDF's @page rule uses (aimformat.pagesetup)."""
        from docx.enum.section import WD_ORIENT
        from docx.shared import Mm

        setup = self.aim.page_setup
        section = self.out.sections[0]
        section.orientation = (
            WD_ORIENT.LANDSCAPE if setup.orientation == "landscape" else WD_ORIENT.PORTRAIT
        )
        section.page_width = Mm(setup.page_width_mm)
        section.page_height = Mm(setup.page_height_mm)
        m = setup.margins_mm
        section.top_margin = Mm(m["top"])
        section.right_margin = Mm(m["right"])
        section.bottom_margin = Mm(m["bottom"])
        section.left_margin = Mm(m["left"])

    def _pop_adds(self, container: str, after: str | None) -> list[Proposal]:
        return self.adds_by_anchor.pop((container, after), [])

    def _emit_anchored_adds(self, container: str, anchor: str | None) -> None:
        # reversed: resolution inserts every same-anchor add at index(anchor)+1,
        # so accept-all leaves the LAST-proposed sibling closest to the anchor
        for prop in reversed(self._pop_adds(container, anchor)):
            self._emit_add_paragraphs(prop)
            self._emit_anchored_adds(container, prop.id)  # chained adds anchor on this one

    def _emit_add_paragraphs(self, prop: Proposal, style: str | None = None) -> None:
        els = self._payload_elements(prop)
        slide_payload = any(el.tag == "aim-slide" for el in els)
        prior_break = self._break_before_next
        # a pending add anchored after a slide (or adding a slide) belongs to
        # the next page, exactly like accepted content. A pending slide owns
        # its added boundary breaks, so Word rejects them together.
        if prior_break or (slide_payload and self._has_content()):
            self._break_before_next = False
            if slide_payload and isinstance(prior_break, Proposal):
                # A blank pending slide ends with its own opening break. That
                # break must not suppress the separate boundary shared with
                # the following pending slide.
                self._shared_pending_slide_break(prior_break, prop)
            elif not self._ends_with_page_break():
                owner = prop if slide_payload else prior_break
                self._page_break(owner if isinstance(owner, Proposal) else None)
        for el in els:
            for block in _block_children(el, self.paint):
                if block.tag in ("ul", "ol"):
                    self._emit_added_list(block, prop)
                    continue
                if block.tag == "table":
                    self.emit_table(block, force="ins", prop=prop)
                    continue
                para = self.out.add_paragraph(
                    style=style or self._safe_style(_style_for(block.tag))
                )
                self.rev.ins(
                    para,
                    _tracked_block_runs(block, self.paint),
                    _actor_label(prop.author),
                    prop.at,
                )
        if slide_payload:
            # When an accepted slide already required a page boundary, move
            # that plain break past this insertion. Otherwise the trailing
            # boundary belongs to the pending slide too.
            self._break_before_next = True if prior_break is True else prop

    def _emit_added_list(self, el: Element, prop: Proposal) -> None:
        """A pending add of a whole list: one inserted paragraph per item —
        the ins half of ``emit_tracked_list_container``. Nested lists inside
        payload items flatten via ``_runs_of``, the same granularity the
        container-modify path has."""
        style = self._safe_style("List Bullet" if el.tag == "ul" else "List Number")
        label, date = _actor_label(prop.author), prop.at
        for li in el.elements():
            para = self.out.add_paragraph(style=style)
            self.rev.ins(para, _tracked_runs(li, self.paint), label, date)

    def _payload_elements(self, prop: Proposal) -> list[Element]:
        """A proposal's payload roots, with their paint already resolved.

        A payload is a separately parsed tree, so it has no ancestors to
        inherit from — it gets the paint and selector context of the exact
        parent it will land in."""
        els = [n for n in parse_fragment(prop.payload_html or "") if isinstance(n, Element)]
        parent = self._payload_parent_node(prop)
        context = self.paint.context_of(parent)
        ancestor_tags = self._payload_ancestor_tags(parent)
        for el in els:
            self.paint.resolve(el, inherited=context, ancestor_tags=ancestor_tags)
        return els

    def _payload_parent_node(self, prop: Proposal) -> Element | None:
        """The exact live parent whose context a payload will enter."""
        if prop.action == "add":
            container = prop.anchor_container or "body"
            parent = (
                self.aim._state.body
                if container == "body"
                else self.aim._state.container_node(container)
            )
            if parent is not None and parent.tag == "table" and prop.anchor_shell is not None:
                return next(
                    (child for child in parent.elements() if child.tag == prop.anchor_shell),
                    parent,
                )
            return parent
        if prop.target:
            try:
                return self.aim._state._target_elements(prop.target)[0][0]
            except Exception:  # malformed/dangling foreign proposal
                return None
        return None

    def _payload_ancestor_tags(self, parent: Element | None) -> tuple[str, ...]:
        """Type-selector ancestry a detached payload will have when applied."""
        tags: list[str] = []
        while parent is not None and parent is not self.aim._state.body:
            tags.append(parent.tag)
            parent = self.aim._state._parent_of(parent)
        return tuple(reversed(tags))

    def _safe_style(self, name: str | None) -> str | None:
        if name is None:
            return None
        return name if self._has_style(name) else None

    def _has_style(self, name: str) -> bool:
        try:
            self.out.styles[name]
            return True
        except KeyError:
            return False

    # -- constructs ------------------------------------------------------------
    def emit_construct(self, el: Element) -> None:
        if el.tag == "aim-slide":
            self.emit_slide(el)
            return
        if self._break_before_next:
            owner = self._break_before_next
            self._break_before_next = False
            self._page_break(owner if isinstance(owner, Proposal) else None)
        cid = el.chunk_id or el.container_id or ""
        prop = self.pending_del.get(cid) or self.pending_mod.get(cid)
        if el.container_id and el.tag in ("ul", "ol"):
            if prop is not None:
                self.emit_tracked_list_container(el, prop)
            else:
                self.emit_list(el)
        elif el.container_id and el.tag == "table":
            if prop is not None:
                self.emit_table(el, force="del", prop=prop)
                if prop.action == "modify":
                    # dispatch each payload root by its own tag: a legal
                    # table→list replacement has no <tr> and emit_table
                    # would drop it silently
                    self._emit_add_paragraphs(prop)
            else:
                self.emit_table(el)
        elif prop is not None:
            self.emit_tracked_chunk(el, prop)
        else:
            for block in _block_children(el, self.paint):
                self.emit_block(block, cid)

    # -- tracked replacements (exactly once per chunk, never per child) -------
    def emit_tracked_chunk(
        self, el: Element, prop: Proposal, style: str | None = None, *, payload: bool = True
    ) -> None:
        """``payload=False`` deletes *el* without re-emitting the modify
        payload — used for all but the last member of a run chunk, so the
        replacement lands exactly once per chunk id."""
        label, date = _actor_label(prop.author), prop.at
        for block in _block_children(el, self.paint):
            para = self.out.add_paragraph(style=style or self._safe_style(_style_for(block.tag)))
            self.rev.dele(para, _tracked_block_runs(block, self.paint), label, date)
        if payload and prop.action == "modify":
            for new_el in self._payload_elements(prop):
                for block in _block_children(new_el, self.paint):
                    para = self.out.add_paragraph(
                        style=style or self._safe_style(_style_for(block.tag))
                    )
                    self.rev.ins(para, _tracked_block_runs(block, self.paint), label, date)

    def emit_tracked_list_container(self, el: Element, prop: Proposal) -> None:
        label, date = _actor_label(prop.author), prop.at
        style = self._safe_style("List Bullet" if el.tag == "ul" else "List Number")
        for li in el.elements():
            para = self.out.add_paragraph(style=style)
            self.rev.dele(para, _tracked_runs(li, self.paint), label, date)
        if prop.action == "modify":
            for new_el in self._payload_elements(prop):
                for li in new_el.elements():
                    para = self.out.add_paragraph(style=style)
                    self.rev.ins(para, _tracked_runs(li, self.paint), label, date)

    # -- plain blocks -----------------------------------------------------------
    def emit_block(self, el: Element, cid: str, style: str | None = None) -> None:
        if el.tag == "figure":
            self.emit_figure(el)
            return
        if el.tag == "pre":
            self.emit_pre(el)
            return
        if el.tag == "table":  # atomic table chunk
            self.emit_table(el)
            return
        if el.tag == "hr":
            # Word's own rule is a paragraph border, but an UNPAINTED document
            # must gain no explicit border at all (the recipient's template
            # owns that), and an empty bordered paragraph is invisible without
            # one. So the em-dash rule stays and an authored border colour
            # paints its ink — a documented approximation.
            para = self.out.add_paragraph()
            runs = [{"text": "—" * 12}]
            border = self.paint.of(el).any_border
            if border is not None:
                runs = [{**runs[0], "color": border.color}]
            _apply_runs(para, runs)
            return
        if el.tag == "aim-page-break":
            from docx.enum.text import WD_BREAK

            para = self.out.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)
            return
        if el.tag in ("ul", "ol"):  # atomic list chunk / nested list content
            self.emit_list(el)
            return
        para = self.out.add_paragraph(style=style or self._safe_style(_style_for(el.tag)))
        align = _alignment_of(el)
        if align is not None:
            para.alignment = align
        runs, box = _clean_runs_and_box(el, self.paint)
        _apply_runs(para, runs)
        _paint_paragraph(para, box)

    def emit_pre(self, el: Element) -> None:
        """A code block, one run per (paint, line) segment.

        Flattening the whole block to a single run was what forced the old
        exporter to choose between painting a nested `<code>`'s sibling text
        and painting nothing at all; per-element paint removes the choice."""
        para = self.out.add_paragraph()
        runs, box = _clean_runs_and_box(el, self.paint)
        for segment in runs:
            if segment.get("break"):
                para.add_run().add_break()
                continue
            for i, line in enumerate(segment.get("text", "").split("\n")):
                if i:  # a newline inside the source text is a Word line break
                    para.add_run().add_break()
                if line:
                    _format_run(para.add_run(line), {**segment, "mono": True})
        _paint_paragraph(para, box)

    def emit_slide(self, el: Element) -> None:
        """Linearize a fixed-canvas page: a page break, then the slide's
        chunks as flowing blocks in reading order (see the module
        docstring). Chunk-level proposals inside the slide ride the same
        tracked/resolve machinery as body chunks; a proposal targeting the
        slide container itself exports as unchanged current content.
        """
        sid = el.container_id or ""
        self._break_before_next = False
        if self._has_content() and not self._ends_with_page_break():
            self._page_break()
        opened_at = len(self.out.element.body)
        if sid:
            self._emit_anchored_adds(sid, None)
        for child in el.elements():
            self.emit_construct(child)
            if sid:
                self._emit_anchored_adds(sid, child.chunk_id or child.container_id)
        if len(self.out.element.body) == opened_at:
            # a blank canvas is still a page (PDF prints it): without at
            # least one paragraph the slide leaves no mark for
            # _has_content(), the next slide resets _break_before_next,
            # and the page silently vanishes from the DOCX
            self.out.add_paragraph()
        # content following the slide belongs to the next page — mirror of
        # the print layer's page-break-after; nothing is emitted when the
        # slide is last, so the document gains no trailing blank page
        self._break_before_next = True

    def _has_content(self) -> bool:
        body = self.out.element.body
        return any(child.tag.endswith("}p") or child.tag.endswith("}tbl") for child in body)

    def _ends_with_page_break(self) -> bool:
        """True when the last emitted paragraph is nothing but a page break —
        an explicit ``aim-page-break`` right before a slide already paged, and
        a second break would print a blank page."""
        from docx.oxml.ns import qn

        paras = [c for c in self.out.element.body if c.tag.endswith("}p")]
        if not paras:
            return False
        last = paras[-1]
        breaks = [b for b in last.findall(".//" + qn("w:br")) if b.get(qn("w:type")) == "page"]
        text = "".join(t.text or "" for t in last.findall(".//" + qn("w:t")))
        return bool(breaks) and not text

    def _page_break(self, prop: Proposal | None = None) -> None:
        from docx.enum.text import WD_BREAK

        para = self.out.add_paragraph()
        if prop is None:
            para.add_run().add_break(WD_BREAK.PAGE)
        else:
            self.rev.ins(para, [{"page_break": True}], _actor_label(prop.author), prop.at)

    def _shared_pending_slide_break(self, prior: Proposal, current: Proposal) -> None:
        """Emit one boundary owned by either consecutive pending slide.

        The current slide replaces the prior slide's inserted break. Word can
        independently accept or reject both slides without losing the boundary
        or rendering two breaks when both are accepted.
        """
        para = self.out.add_paragraph()
        self.rev.replace_ins(
            para,
            [{"page_break": True}],
            _actor_label(prior.author),
            prior.at,
            _actor_label(current.author),
            current.at,
        )

    def emit_figure(self, el: Element) -> None:
        # emit_block routes figures here before its generic alignment code,
        # so the figure's own alignment class lands on the picture paragraph
        align = _alignment_of(el)
        img = el.find(lambda e: e.tag == "img")
        emitted = False
        if img is not None:
            src = img.get("src") or ""
            m = re.match(r"^data:image/[a-z+.-]+;base64,(.*)$", src, re.S)
            if m:
                try:
                    blob = base64.b64decode(m.group(1))
                    self.out.add_picture(io.BytesIO(blob), width=self._figure_width(el, img))
                    if align is not None:
                        self.out.paragraphs[-1].alignment = align
                    emitted = True
                except Exception:
                    emitted = False
            if not emitted:
                alt = img.get("alt") or "image"
                para = self.out.add_paragraph()
                if align is not None:
                    para.alignment = align
                spec = {"text": f"[image: {alt}]", "italic": True, **_ink(self.paint.of(img))}
                _apply_runs(para, [spec])
                emitted = True
        for child in el.elements():  # non-caption content first, then captions
            if child.tag not in ("figcaption", "img", "svg"):
                self.emit_block(child, el.chunk_id or "")
        for cap in el.elements():
            if cap.tag == "figcaption":
                # built through _runs_of rather than add_paragraph(text) so a
                # caption's own formatting — including colour — survives
                # (Codex aimformat#19)
                para = self.out.add_paragraph(style=self._safe_style("Caption"))
                runs, box = _clean_runs_and_box(cap, self.paint)
                _apply_runs(para, runs)
                _paint_paragraph(para, box)

    def _figure_width(self, fig: Element, img: Element):
        """The authored image width (inline style, CSS px at 96 dpi — the
        img's own, else the figure's), capped to the page content box; the
        historical 4.5 in when nothing is declared."""
        from docx.shared import Inches

        px = _style_px(img, "width")
        if px is None:
            px = _style_px(fig, "width")
        inches = px / 96.0 if px is not None else 4.5
        setup = self.aim.page_setup
        margins = setup.margins_mm
        avail = (setup.page_width_mm - margins["left"] - margins["right"]) / 25.4
        return Inches(max(0.25, min(inches, avail)))

    # -- lists ---------------------------------------------------------------------
    def emit_list(self, el: Element, level: int = 0) -> None:
        base = "List Bullet" if el.tag == "ul" else "List Number"
        style = base if level == 0 else f"{base} {min(level + 1, 3)}"
        if not self._has_style(style):
            style = base if self._has_style(base) else None
        container_id = el.container_id
        if container_id:
            self._emit_list_adds(container_id, None, style)
        items = el.elements()
        i = 0
        while i < len(items):
            cid = items[i].chunk_id or ""
            group = [items[i]]  # a run chunk = consecutive same-id items
            while cid and i + 1 < len(items) and items[i + 1].chunk_id == cid:
                i += 1
                group.append(items[i])
            i += 1
            prop = (self.pending_del.get(cid) or self.pending_mod.get(cid)) if cid else None
            for li in group:
                nested = [c for c in li.elements() if c.tag in ("ul", "ol")]
                content = Element("li")
                content.children = [
                    c for c in li.children if not (isinstance(c, Element) and c.tag in ("ul", "ol"))
                ]
                # the copy is synthetic, so it adopts the item's computed
                # paint rather than resolving as if it were in the tree — but
                # class/style must ride along: alignment and literal
                # typography are read off the element itself
                for attr in ("class", "style"):
                    value = li.get(attr)
                    if value:
                        content.set(attr, value)
                self.paint.adopt(content, li)
                if prop is not None:
                    self.emit_tracked_chunk(content, prop, style=style, payload=li is group[-1])
                else:
                    self.emit_block(content, cid, style=style)
                for sub in nested:
                    self.emit_list(sub, level + 1)
            if container_id and cid:
                self._emit_list_adds(container_id, cid, style)

    def _emit_list_adds(self, container: str, after: str | None, style: str | None) -> None:
        # reversed for the same reason as _emit_anchored_adds
        for prop in reversed(self._pop_adds(container, after)):
            self._emit_add_paragraphs(prop, style=style)
            self._emit_list_adds(container, prop.id, style)

    # -- tables ----------------------------------------------------------------------
    def emit_table(
        self, el: Element, *, force: str | None = None, prop: Proposal | None = None
    ) -> None:
        rows = el.find_all(lambda e: e.tag == "tr")
        if not rows:
            return
        # pending row-modifies: cell-level del+ins can express a replacement
        # only when the grid shape is unchanged (same row count, same per-cell
        # spans); a structural change becomes row-delete + inserted rows —
        # fusing surplus cells into the last old cell produced a grid no
        # accept/reject sequence in Word could turn into the AIM state
        cellwise_rows: dict[int, list[Element]] = {}  # ri -> payload row cells
        structural_ins: dict[int, list[Element]] = {}  # last-member ri -> payload rows
        structural_del: set[int] = set()
        if not force:
            member_rows: dict[str, list[int]] = {}
            for ri, row in enumerate(rows):
                if row.chunk_id and row.chunk_id in self.pending_mod:
                    member_rows.setdefault(row.chunk_id, []).append(ri)
            for rid, indices in member_rows.items():
                payload_rows = [
                    p for p in self._payload_elements(self.pending_mod[rid]) if p.tag == "tr"
                ]
                if len(payload_rows) == len(indices) and all(
                    _row_shape(payload_rows[k]) == _row_shape(rows[i])
                    for k, i in enumerate(indices)
                ):
                    for k, i in enumerate(indices):
                        cellwise_rows[i] = [
                            c for c in payload_rows[k].elements() if c.tag in ("td", "th")
                        ]
                else:
                    structural_ins[indices[-1]] = payload_rows
                    structural_del.update(indices)
        # true width of the live table: simulate the emit loop's cursor (spans
        # shift later cells right; a rowspan reaching below the last row is
        # clamped). Pending structural replacements build their own row shape
        # below and must not widen the original rows outside tracked revisions.
        ncols = 0
        spanned: set[tuple[int, int]] = set()
        for ri, row in enumerate(rows):
            ci = 0
            for c in row.elements():
                if c.tag not in ("td", "th"):
                    continue
                while (ri, ci) in spanned:
                    ci += 1
                cs = int(c.get("colspan") or 1)
                rs = min(int(c.get("rowspan") or 1), len(rows) - ri)
                if cs > 1 or rs > 1:
                    for dr in range(rs):
                        for dc in range(cs):
                            spanned.add((ri + dr, ci + dc))
                ci += cs
            ncols = max(ncols, ci)
        table = self.out.add_table(rows=len(rows), cols=ncols)
        table.style = "Table Grid" if self._has_style("Table Grid") else None
        occupied: set[tuple[int, int]] = set()
        container_id = el.container_id
        for ri, row in enumerate(rows):
            rid = row.chunk_id or ""
            prop_mod = self.pending_mod.get(rid) if not force else None
            prop_del = self.pending_del.get(rid) if not force else None
            if force == "del":
                prop_del = prop  # container-level: everything deleted
            new_cells = cellwise_rows.get(ri)  # structural members stay dele-only
            cells = [c for c in row.elements() if c.tag in ("td", "th")]
            ci = 0
            for orig_idx, cell_el in enumerate(cells):
                while (ri, ci) in occupied:
                    ci += 1
                colspan = min(int(cell_el.get("colspan") or 1), ncols - ci)
                rowspan = min(int(cell_el.get("rowspan") or 1), len(rows) - ri)
                cell = table.cell(ri, ci)
                if colspan > 1 or rowspan > 1:
                    cell = cell.merge(table.cell(ri + rowspan - 1, ci + colspan - 1))
                    for dr in range(rowspan):
                        for dc in range(colspan):
                            occupied.add((ri + dr, ci + dc))
                para = cell.paragraphs[0]
                tracked_cell = force is not None or prop_del is not None or prop_mod is not None
                if tracked_cell:
                    runs = _tracked_runs(cell_el, self.paint)
                    clean_box = None
                else:
                    runs, clean_box = _clean_runs_and_box(cell_el, self.paint)
                if cell_el.tag == "th":
                    runs = [{**r, "bold": True} for r in runs]
                if force == "ins" and prop is not None:
                    self.rev.ins(para, runs, _actor_label(prop.author), prop.at)
                elif prop_del is not None:
                    self.rev.dele(para, runs, _actor_label(prop_del.author), prop_del.at)
                elif prop_mod is not None:
                    self.rev.dele(para, runs, _actor_label(prop_mod.author), prop_mod.at)
                    if new_cells is not None:  # cellwise: shapes match 1:1
                        self.rev.ins(
                            para,
                            _tracked_runs(new_cells[orig_idx], self.paint),
                            _actor_label(prop_mod.author),
                            prop_mod.at,
                        )
                else:
                    _apply_runs(para, runs)
                # A tracked cell has one live w:tcPr outside both revisions.
                # Its old/new box paint rides the corresponding runs instead;
                # clean cells keep exact cell shading and per-side borders.
                if not tracked_cell:
                    assert clean_box is not None
                    _paint_cell(cell, clean_box)
                ci += colspan
        # row-adds only after the content loop: inserting rows mid-loop would
        # shift the (ri, ci) coordinates the loop and the merges rely on
        orig_trs = [r._tr for r in table.rows]  # 1:1 with the AIM rows
        if force and prop is not None:
            # container-level tracked change: the table STRUCTURE itself is
            # pending, not just its text — without trPr markers, accepting
            # the revisions in Word leaves an empty grid behind (del) and
            # rejecting an added table strands one (ins)
            mark = self.rev.row_dele if force == "del" else self.rev.row_ins
            for tr in orig_trs:
                mark(tr, _actor_label(prop.author), prop.at)
        for ri in sorted(structural_del):
            prop_mod = self.pending_mod[rows[ri].chunk_id or ""]
            self.rev.row_dele(orig_trs[ri], _actor_label(prop_mod.author), prop_mod.at)
        if container_id and not force:
            self._emit_row_adds(table, None, container_id, None, ncols, first=True)
            for ri, row in enumerate(rows):
                rid = row.chunk_id
                # anchor on the LAST member of a run chunk: draining on the
                # first member would insert the added row mid-run
                last_of_run = not (ri + 1 < len(rows) and rows[ri + 1].chunk_id == rid)
                if rid and last_of_run:
                    self._emit_row_adds(table, orig_trs[ri], container_id, rid, ncols)
        # structural replacements insert right after their (deleted) old rows —
        # closest to the anchor, ahead of any row-adds anchored on the chunk
        for ri, payload_rows in structural_ins.items():
            prop_mod = self.pending_mod[rows[ri].chunk_id or ""]
            label, date = _actor_label(prop_mod.author), prop_mod.at
            cur = orig_trs[ri]
            # grid col -> (continuation rows still owed, colspan) for rowspans
            # opened by an earlier payload row of THIS replacement
            vmerge: dict[int, tuple[int, int]] = {}
            for pi, p_row in enumerate(payload_rows):
                new_row = table.add_row()  # appended; repositioned below
                tr = new_row._tr
                for tc in list(tr.tc_lst):  # rebuilt cell by cell: the payload
                    tr.remove(tc)  # rows carry their own span structure
                ci = 0
                for c in p_row.elements():
                    if c.tag not in ("td", "th"):
                        continue
                    ci = self._continue_vmerges(table, tr, ci, vmerge)
                    cs = int(c.get("colspan") or 1)
                    rs = min(int(c.get("rowspan") or 1), len(payload_rows) - pi)
                    para = self._new_span_cell(
                        table, tr, colspan=cs, vmerge="restart" if rs > 1 else None
                    )
                    runs = _tracked_runs(c, self.paint)
                    if c.tag == "th":
                        runs = [{**run, "bold": True} for run in runs]
                    self.rev.ins(para, runs, label, date)
                    if rs > 1:
                        vmerge[ci] = (rs - 1, cs)
                    ci += cs
                self._continue_vmerges(table, tr, ci, vmerge, trailing=True)
                self.rev.row_ins(tr, label, date)
                cur.addnext(tr)
                cur = tr

    def _new_span_cell(self, table, tr, *, colspan: int, vmerge: str | None):
        """Append one cell to a rebuilt payload row with its span structure;
        return the cell's paragraph."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.table import _Cell

        tc = tr.add_tc()
        if colspan > 1:
            span = OxmlElement("w:gridSpan")
            span.set(qn("w:val"), str(colspan))
            tc.get_or_add_tcPr().append(span)
        if vmerge is not None:
            merge = OxmlElement("w:vMerge")
            if vmerge == "restart":
                merge.set(qn("w:val"), "restart")
            tc.get_or_add_tcPr().append(merge)
        cell = _Cell(tc, table)
        return cell.paragraphs[0]

    def _continue_vmerges(
        self, table, tr, ci: int, vmerge: dict[int, tuple[int, int]], *, trailing: bool = False
    ) -> int:
        """Emit the vertical-merge continuation cells owed at grid position
        *ci* (and, when *trailing*, any owed further right); returns the
        advanced position. Content lives in the restart cell — a
        continuation is an empty spanned cell."""
        while True:
            if ci in vmerge:
                left, cs = vmerge.pop(ci)
                self._new_span_cell(table, tr, colspan=cs, vmerge="continue")
                if left > 1:
                    vmerge[ci] = (left - 1, cs)
                ci += cs
            elif trailing and vmerge and min(vmerge) > ci:
                ci = min(vmerge)
            else:
                return ci

    def _emit_row_adds(
        self,
        table,
        anchor_tr,
        container: str,
        after: str | None,
        ncols: int,
        first: bool = False,
    ) -> None:
        """Insert pending row-adds as fully-inserted (w:ins) table rows,
        each positioned right after its anchor ``w:tr`` (the container start
        when ``first``)."""
        props = self._pop_adds(container, after)
        for prop in props:
            new_row = table.add_row()  # appended; repositioned below
            payload_cells = [
                c
                for p in self._payload_elements(prop)
                for c in p.elements()
                if c.tag in ("td", "th")
            ]
            for idx in range(min(len(payload_cells), ncols)):
                cell = new_row.cells[idx]
                self.rev.ins(
                    cell.paragraphs[0],
                    _tracked_runs(payload_cells[idx], self.paint),
                    _actor_label(prop.author),
                    prop.at,
                )
            if first:
                table.rows[0]._tr.addprevious(new_row._tr)
            else:
                anchor_tr.addnext(new_row._tr)
            # chained adds anchor on the row just inserted
            self._emit_row_adds(table, new_row._tr, container, prop.id, ncols)


# --------------------------------------------------------------------------
def _resolve_copy(doc: AimDocument, decision: str) -> AimDocument:
    clone = AimDocument.loads(doc.dumps())
    decider = external("docx-export")
    if decision == "accept-all":
        clone.accept_all(decided_by=decider)
    else:
        clone.reject_all(decided_by=decider)
    return clone


def to_docx(doc: AimDocument, path: str | Path, *, pending: str = "tracked") -> Path:
    """Write *doc* to *path* as a .docx file.

    ``pending`` — what to do with the pending lane: ``"tracked"`` emits Word
    revision markup, ``"accept-all"``/``"reject-all"`` resolve a throwaway
    copy first (the original document is never mutated).
    """
    if pending not in _PENDING_MODES:
        raise InvalidOperation(f"pending must be one of {_PENDING_MODES}, got {pending!r}")
    docx_mod = _require_docx()
    source = doc if pending == "tracked" else _resolve_copy(doc, pending)
    exporter = _Exporter(source, docx_mod)
    exporter.run()
    out = Path(path)
    exporter.out.save(str(out))
    return out
