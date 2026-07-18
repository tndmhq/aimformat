"""Page setup (aim:doc) and aim-page-break: grammars, ops, undo/verify,
hash coverage, proposals, and the export mappings (spec §3.6)."""

import pytest

import aimformat as aim
from aimformat.canonical import canonical_json
from aimformat.errors import InvalidOperation, ParseError
from aimformat.pagesetup import default_page_setup, page_css, page_setup_from_obj
from conftest import BOT, ME, ts


# --------------------------------------------------------------------------
class TestGrammar:
    def test_defaults_are_a4_15mm(self):
        s = default_page_setup()
        assert (s.size, s.orientation) == ("A4", "portrait")
        assert s.margins_mm == {"top": 15.0, "right": 15.0, "bottom": 15.0, "left": 15.0}
        assert (s.page_width_mm, s.page_height_mm) == (210.0, 297.0)
        assert (s.content_width_mm, s.content_height_mm) == (180.0, 267.0)

    def test_partial_object_falls_back_to_defaults(self):
        s = page_setup_from_obj({"size": "A5"})
        assert s.size == "A5" and s.orientation == "portrait"
        assert s.margins_mm["top"] == 15.0

    def test_landscape_swaps_dimensions(self):
        s = page_setup_from_obj({"size": "A4", "orientation": "landscape"})
        assert (s.page_width_mm, s.page_height_mm) == (297.0, 210.0)

    def test_unknown_size_is_d003(self):
        with pytest.raises(InvalidOperation) as exc:
            page_setup_from_obj({"size": "A0"})
        assert exc.value.lint_code == "D003"

    def test_unknown_orientation_is_d003(self):
        with pytest.raises(InvalidOperation) as exc:
            page_setup_from_obj({"orientation": "sideways"})
        assert exc.value.lint_code == "D003"

    def test_margin_grammar_is_d004(self):
        with pytest.raises(InvalidOperation) as exc:
            page_setup_from_obj({"margins": {"top": "15px"}})
        assert exc.value.lint_code == "D004"

    def test_margin_bound_is_d004(self):
        with pytest.raises(InvalidOperation) as exc:
            page_setup_from_obj({"margins": {"top": "250mm"}})
        assert exc.value.lint_code == "D004"

    def test_margins_must_leave_content_area(self):
        with pytest.raises(InvalidOperation) as exc:
            page_setup_from_obj({"size": "A5", "margins": {"left": "80mm", "right": "80mm"}})
        assert exc.value.lint_code == "D004"

    def test_non_object_is_d001(self):
        with pytest.raises(InvalidOperation) as exc:
            page_setup_from_obj(["A4"])
        assert exc.value.lint_code == "D001"

    def test_page_css(self):
        s = page_setup_from_obj(
            {
                "size": "Letter",
                "orientation": "landscape",
                "margins": {"top": "20mm", "right": "15mm", "bottom": "20mm", "left": "15mm"},
            }
        )
        assert page_css(s) == ("@page{size:279.4mm 215.9mm;margin:20mm 15mm 20mm 15mm}")

    def test_to_obj_round_trips(self):
        obj = {
            "size": "Legal",
            "orientation": "portrait",
            "margins": {"top": "12.5mm", "right": "15mm", "bottom": "12.5mm", "left": "15mm"},
        }
        assert (
            page_setup_from_obj(page_setup_from_obj(obj).to_obj()).to_obj()
            == page_setup_from_obj(obj).to_obj()
        )

    def test_fmt_mm_never_uses_exponent_notation(self):
        # %g would spell 0.000001 as "1e-06" — a valid float that violates
        # the margin grammar, so the writer would emit a block its own
        # linter rejects (Codex PR-4 review)
        from aimformat.pagesetup import _fmt_mm

        assert _fmt_mm(0.000001) == "0.000001mm"
        assert _fmt_mm(0.0) == "0mm"
        assert _fmt_mm(15.0) == "15mm"
        assert _fmt_mm(215.9) == "215.9mm"
        from aimformat.registry import REGISTRY

        for v in (0.000001, 1e-7, 0.1, 12.5, 100.0):
            assert REGISTRY.margin_pattern.match(_fmt_mm(v))

    def test_tiny_margin_round_trips_through_the_grammar(self):
        s = page_setup_from_obj({"margins": {"top": "0.000001mm"}})
        assert s.to_obj()["margins"]["top"] == "0.000001mm"
        assert page_setup_from_obj(s.to_obj()).margins_mm["top"] == 0.000001


# --------------------------------------------------------------------------
class TestSetPageSetup:
    def test_set_writes_block_and_event(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        assert basic_doc.page_setup.size == "A5"
        ev = basic_doc.history[-1]
        assert (ev.target, ev.action) == ("aim:doc", "modify")
        assert ev.get("before") is None
        text = basic_doc.dumps()
        assert 'type="application/aim-doc+json"' in text
        assert not [f for f in aim.lint_text(text) if f.level == "error"]

    def test_unchanged_setup_raises(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        with pytest.raises(InvalidOperation):
            basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(6))

    def test_undo_removes_introduced_block(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        basic_doc.undo(author=ME, at=ts(6))
        assert basic_doc.page_setup.size == "A4"
        assert basic_doc._state.script("doc") is None
        assert basic_doc.verify() == []
        basic_doc.redo(author=ME, at=ts(7))
        assert basic_doc.page_setup.size == "A5"
        assert basic_doc.verify() == []

    def test_second_set_carries_before_and_undoes_to_previous(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        basic_doc.set_page_setup({"size": "Letter"}, author=ME, at=ts(6))
        assert basic_doc.history[-1].get("before") is not None
        basic_doc.undo(author=ME, at=ts(7))
        assert basic_doc.page_setup.size == "A5"
        assert basic_doc.verify() == []

    def test_doc_hash_covers_settings(self, basic_doc):
        before = basic_doc.doc_hash
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        changed = basic_doc.doc_hash
        assert changed != before
        basic_doc.undo(author=ME, at=ts(6))
        assert basic_doc.doc_hash == before  # absent block = pre-0.2 hash

    def test_state_at_reconstructs_setup(self, basic_doc):
        seq_before = basic_doc.seq
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        assert basic_doc.state_at(seq_before).page_setup.size == "A4"
        assert basic_doc.state_at(basic_doc.seq).page_setup.size == "A5"

    def test_unknown_settings_fields_survive(self, empty_doc):
        empty_doc.flatten()
        text = empty_doc.dumps().replace(
            "</title>",
            '</title>\n<script type="application/aim-doc+json">\n'
            + canonical_json({"x_vendor": {"k": 1}})
            + "\n</script>",
        )
        doc = aim.loads(text)
        doc.set_page_setup({"size": "A5"}, author=ME, at=ts(1))
        assert doc.doc_settings["x_vendor"] == {"k": 1}
        assert doc.doc_settings["page"]["size"] == "A5"

    def test_duplicate_settings_blocks_are_d002(self, empty_doc):
        empty_doc.flatten()
        block = '<script type="application/aim-doc+json">\n{"page":{"size":"A5"}}\n</script>'
        text = empty_doc.dumps().replace("</title>", f"</title>\n{block}\n{block}")
        codes = {f.code for f in aim.lint_text(text) if f.level == "error"}
        assert codes == {"D002"}

    def test_malformed_block_is_parse_error(self, empty_doc):
        empty_doc.flatten()
        text = empty_doc.dumps().replace(
            "</title>", '</title>\n<script type="application/aim-doc+json">\n{not json]\n</script>'
        )
        doc = aim.loads(text)
        with pytest.raises(ParseError):
            _ = doc.doc_settings
        codes = {f.code for f in aim.lint_text(text)}
        assert "D001" in codes

    def test_invalid_setup_rejected_before_mutation(self, basic_doc):
        before = basic_doc.dumps()
        with pytest.raises(InvalidOperation):
            basic_doc.set_page_setup({"size": "A0"}, author=ME, at=ts(5))
        assert basic_doc.dumps() == before

    def test_tiny_margin_survives_a_full_write_read_cycle(self, basic_doc):
        # regression: "%g" serialized this as 1e-06mm — set_page_setup then
        # produced a document whose own re-read raised D004
        basic_doc.set_page_setup({"margins": {"top": "0.000001mm"}}, author=ME, at=ts(5))
        assert basic_doc.page_setup.margins_mm["top"] == 0.000001
        assert basic_doc.doc_settings["page"]["margins"]["top"] == "0.000001mm"
        assert not [f for f in aim.lint_text(basic_doc.dumps()) if f.level == "error"]
        assert basic_doc.verify() == []


# --------------------------------------------------------------------------
class TestPageBreakChunk:
    def test_add_move_delete_undo(self, basic_doc):
        pb = basic_doc.add_chunk(
            "<aim-page-break></aim-page-break>", author=ME, after="h1", at=ts(5)
        )
        serial = basic_doc._state.serial(pb.id)
        assert serial == (f'<aim-page-break data-aim="{pb.id}"></aim-page-break>')
        assert not [f for f in aim.lint_text(basic_doc.dumps()) if f.level == "error"]
        basic_doc.move_chunk(pb.id, author=ME, at=ts(6))  # to end
        assert basic_doc.body_ids[-1] == pb.id
        basic_doc.delete_chunk(pb.id, author=ME, at=ts(7))
        assert not basic_doc._state.exists(pb.id)
        basic_doc.undo(author=ME, at=ts(8))
        assert basic_doc._state.exists(pb.id)
        assert basic_doc.verify() == []

    def test_propose_and_accept_break(self, basic_doc):
        p = basic_doc.propose_add(
            "<aim-page-break></aim-page-break>",
            author=BOT,
            after="h1",
            at=ts(5),
            explanation="Start a fresh page.",
        )
        assert not [f for f in aim.lint_text(basic_doc.dumps()) if f.level == "error"]
        basic_doc.accept(p.id, decided_by=ME, at=ts(6))
        tags = [c.tag for c in basic_doc.chunks]
        assert "aim-page-break" in tags
        assert basic_doc.verify() == []

    def test_break_inside_container_is_d006(self, rich_doc):
        # the write path refuses illegal container members outright (AF-03);
        # the D006 diagnostic must still fire on hand-authored text
        with pytest.raises(InvalidOperation):
            rich_doc.add_chunk(
                "<aim-page-break></aim-page-break>", author=ME, container="list", at=ts(20)
            )
        text = rich_doc.dumps().replace(
            '<li data-aim="li1">First</li>',
            '<li data-aim="li1">First</li><aim-page-break data-aim="pbx"></aim-page-break>',
        )
        codes = {f.code for f in aim.lint_text(text)}
        assert "D006" in codes


# --------------------------------------------------------------------------
class TestPageSetupProposals:
    def test_propose_accept(self, basic_doc):
        p = basic_doc.propose_page_setup(
            {"size": "A5"}, author=BOT, at=ts(5), explanation="Booklet."
        )
        assert p.target == "aim:doc"
        assert not [f for f in aim.lint_text(basic_doc.dumps()) if f.level == "error"]
        basic_doc.accept(p.id, decided_by=ME, at=ts(6))
        assert basic_doc.page_setup.size == "A5"
        assert basic_doc.verify() == []

    def test_propose_reject_keeps_defaults(self, basic_doc):
        p = basic_doc.propose_page_setup({"size": "A5"}, author=BOT, at=ts(5))
        basic_doc.reject(p.id, decided_by=ME, at=ts(6))
        assert basic_doc.page_setup.size == "A4"
        assert basic_doc.verify() == []

    def test_accept_with_invalid_tweak_raises(self, basic_doc):
        p = basic_doc.propose_page_setup({"size": "A5"}, author=BOT, at=ts(5))
        with pytest.raises(InvalidOperation):
            basic_doc.accept(
                p.id,
                decided_by=ME,
                at=ts(6),
                applied='<script type="application/aim-doc+json">\n'
                '{"page":{"size":"A0"}}\n</script>',
            )

    def test_typed_script_in_ordinary_payload_still_v002(self, basic_doc):
        # the style/script vocabulary skip applies ONLY to whole-block
        # singleton payloads — inside a normal chunk payload a typed script
        # must stay a vocabulary error, not a smuggling hole
        basic_doc.propose_modify("intro", '<p data-aim="intro">ok</p>', author=BOT, at=ts(5))
        text = basic_doc.dumps().replace(
            '<template><p data-aim="intro">ok</p></template>',
            '<template><p data-aim="intro">ok<script '
            'type="application/aim-doc+json">{}</script></p></template>',
        )
        codes = {f.code for f in aim.lint_text(text) if f.level == "error"}
        assert "V002" in codes

    def test_self_closing_settings_payload_is_rejected(self, basic_doc):
        # in HTML a self-closed <script/> is still an open tag: the payload
        # would parse as defaults here but swallow markup in a browser
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(4))
        with pytest.raises(InvalidOperation) as exc:
            basic_doc.propose_modify(
                "aim:doc", '<script type="application/aim-doc+json"/>', author=BOT, at=ts(5)
            )
        assert exc.value.lint_code == "D001"

    def test_self_closed_live_settings_block_is_d001(self, empty_doc):
        empty_doc.flatten()
        text = empty_doc.dumps().replace(
            "</title>", '</title>\n<script type="application/aim-doc+json"/>'
        )
        codes = {f.code for f in aim.lint_text(text) if f.level == "error"}
        assert "D001" in codes

    def test_modify_via_propose_modify(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(4))
        p = basic_doc.propose_modify(
            "aim:doc",
            '<script type="application/aim-doc+json">\n'
            '{"page":{"margins":{"bottom":"15mm","left":"15mm",'
            '"right":"15mm","top":"15mm"},"orientation":"landscape",'
            '"size":"A5"}}\n</script>',
            author=BOT,
            at=ts(5),
        )
        basic_doc.accept(p.id, decided_by=ME, at=ts(6))
        assert basic_doc.page_setup.orientation == "landscape"
        assert basic_doc.verify() == []


# --------------------------------------------------------------------------
class TestReservedTargetGuards:
    """propose_delete("aim:doc") used to create a lint-clean pending delete
    that only exploded (TargetNotFound) at accept time — reserved heads have
    no body anchor. Delete/move must be rejected up front, for aim:theme
    alike (Codex PR-4 review)."""

    def test_propose_delete_doc_rejected_at_propose_time(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        with pytest.raises(InvalidOperation, match="reserved"):
            basic_doc.propose_delete("aim:doc", author=BOT, at=ts(6))
        assert basic_doc.proposals == []  # no half-created card

    def test_propose_move_doc_rejected_at_propose_time(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        with pytest.raises(InvalidOperation, match="reserved"):
            basic_doc.propose_move("aim:doc", author=BOT, container="body", at=ts(6))

    def test_direct_delete_and_move_rejected(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        with pytest.raises(InvalidOperation, match="reserved"):
            basic_doc.delete_chunk("aim:doc", author=ME, at=ts(6))
        with pytest.raises(InvalidOperation, match="reserved"):
            basic_doc.move_chunk("aim:doc", author=ME, at=ts(6))

    def test_theme_target_equally_guarded(self, basic_doc):
        with pytest.raises(InvalidOperation, match="reserved"):
            basic_doc.propose_delete("aim:theme", author=BOT, at=ts(5))
        with pytest.raises(InvalidOperation, match="reserved"):
            basic_doc.delete_chunk("aim:theme", author=ME, at=ts(5))

    def test_hand_authored_delete_card_fails_accept_with_intent(self, basic_doc):
        # a foreign tool can still write such a card: accept must fail
        # loudly as InvalidOperation, while reject stays available
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        p = basic_doc.propose_delete("intro", author=BOT, at=ts(6))
        doc = aim.loads(basic_doc.dumps().replace('data-for="intro"', 'data-for="aim:doc"'))
        with pytest.raises(InvalidOperation, match="reserved"):
            doc.accept(p.id, decided_by=ME, at=ts(7))
        doc.reject(p.id, decided_by=ME, at=ts(8))
        assert doc.proposals == []
        assert doc.verify() == []


# --------------------------------------------------------------------------
class TestDeclaredVersionStability:
    """The <html> open line is hashed state: an implicit version upgrade on
    save would break every checkpoint recorded under the old line (found via
    v0.1 documents 422-ing on re-save during the pagination work)."""

    def test_dumps_never_touches_the_declared_version(self, empty_doc):
        empty_doc.flatten()
        old = empty_doc.dumps().replace(
            f'data-aim-version="{aim.SPEC_VERSION}"', 'data-aim-version="0.1"'
        )
        round_tripped = aim.loads(old).dumps()
        assert 'data-aim-version="0.1"' in round_tripped
        findings = aim.lint_text(round_tripped)
        assert "S002" in {f.code for f in findings if f.level == "warning"}
        assert not [f for f in findings if f.level == "error"]

    def test_old_version_checkpoints_survive_a_resave(self, basic_doc):
        basic_doc.checkpoint("pinned", at=ts(5))
        text = basic_doc.dumps()
        # simulate the load→save cycle every backend write performs
        again = aim.loads(text)
        assert again.verify() == []
        assert aim.loads(again.dumps()).verify() == []


# --------------------------------------------------------------------------
class TestPrintCssAndPdfHtml:
    def test_aim_css_carries_print_pagination_layer(self):
        css = aim.generate_aim_css()
        assert "aim-page-break{display:block" in css
        assert "@media print{aim-page-break{border:0;margin:0;break-after:page}}" in css
        assert "@media print{[data-aim]{break-inside:avoid}}" in css
        assert "@media print{body{margin:0;padding:0;max-width:none}}" in css

    def test_print_html_splices_page_rule(self, basic_doc):
        from aimformat.convert._pdf_out import _print_html

        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        html = _print_html(basic_doc, "keep", "@font-face{font-family:X}")
        assert "@page{size:148mm 210mm;margin:15mm 15mm 15mm 15mm}" in html
        assert "@font-face{font-family:X}" in html
        assert html.index("@page{") < html.index("</head>")

    def test_print_html_resolves_pending_before_page_css(self, basic_doc):
        # regression: with pending="accept-all" the @page rule was computed
        # from the PRE-resolution document, printing the accepted A5 page
        # inside the old A4 geometry
        from aimformat.convert._pdf_out import _print_html

        basic_doc.propose_page_setup({"size": "A5"}, author=BOT, at=ts(5))
        html = _print_html(basic_doc, "accept-all", None)
        assert "@page{size:148mm 210mm" in html
        assert "@page{size:210mm 297mm" not in html
        assert len(basic_doc.proposals) == 1  # the copy was throwaway
        html = _print_html(basic_doc, "reject-all", None)
        assert "@page{size:210mm 297mm" in html
        with pytest.raises(InvalidOperation):
            _print_html(basic_doc, "bogus", None)

    def test_to_pdf_smoke_page_count(self, basic_doc, tmp_path):
        pytest.importorskip("playwright")
        basic_doc.add_chunk("<aim-page-break></aim-page-break>", author=ME, after="h1", at=ts(5))
        out = tmp_path / "smoke.pdf"
        try:
            aim.to_pdf(basic_doc, out)
        except RuntimeError as exc:
            if "Chromium is not installed" not in str(exc):
                raise  # a real to_pdf regression must fail, not skip
            pytest.skip(str(exc))
        data = out.read_bytes()
        assert data.startswith(b"%PDF")
        pages = data.count(b"/Type /Page") - data.count(b"/Type /Pages")
        assert pages >= 2  # the hard break yields at least two pages


# --------------------------------------------------------------------------
class TestDocxMappings:
    docx = pytest.importorskip("docx")

    def test_export_emits_section_and_break(self, basic_doc, tmp_path):
        from docx import Document

        basic_doc.set_page_setup(
            {
                "size": "Letter",
                "orientation": "landscape",
                "margins": {"top": "20mm", "right": "15mm", "bottom": "20mm", "left": "15mm"},
            },
            author=ME,
            at=ts(5),
        )
        basic_doc.add_chunk("<aim-page-break></aim-page-break>", author=ME, after="h1", at=ts(6))
        out = tmp_path / "out.docx"
        aim.to_docx(basic_doc, out)
        got = Document(str(out))
        section = got.sections[0]
        assert round(section.page_width / 36000, 1) == 279.4
        assert round(section.page_height / 36000, 1) == 215.9
        assert round(section.top_margin / 36000, 1) == 20.0
        from docx.oxml.ns import qn

        brs = [
            br
            for p in got.paragraphs
            for r in p.runs
            for br in r._r.findall(qn("w:br"))
            if br.get(qn("w:type")) == "page"
        ]
        assert len(brs) == 1

    def test_tracked_export_emits_pending_break_add_as_page_break(self, basic_doc, tmp_path):
        # regression: a pending aim-page-break add has no text runs, so the
        # tracked lane emitted an empty w:ins paragraph — the review DOCX
        # lost the proposed pagination change
        from docx import Document
        from docx.oxml.ns import qn

        basic_doc.propose_add("<aim-page-break></aim-page-break>", author=BOT, after="h1", at=ts(5))
        out = tmp_path / "tracked-break-add.docx"
        aim.to_docx(basic_doc, out)  # default pending="tracked"
        got = Document(str(out))
        brs = [
            br
            for ins in got.element.body.findall(".//" + qn("w:ins"))
            for br in ins.findall(".//" + qn("w:br"))
            if br.get(qn("w:type")) == "page"
        ]
        assert len(brs) == 1

    def test_tracked_export_emits_pending_break_delete_as_deleted_break(self, basic_doc, tmp_path):
        from docx import Document
        from docx.oxml.ns import qn

        pb = basic_doc.add_chunk(
            "<aim-page-break></aim-page-break>", author=ME, after="h1", at=ts(5)
        )
        basic_doc.propose_delete(pb.id, author=BOT, at=ts(6))
        out = tmp_path / "tracked-break-del.docx"
        aim.to_docx(basic_doc, out)
        got = Document(str(out))
        brs = [
            br
            for dele in got.element.body.findall(".//" + qn("w:del"))
            for br in dele.findall(".//" + qn("w:br"))
            if br.get(qn("w:type")) == "page"
        ]
        assert len(brs) == 1

    def test_pending_settings_proposal_not_misfiled_as_chunk_edit(self, basic_doc, tmp_path):
        from docx import Document
        from docx.oxml.ns import qn

        basic_doc.propose_page_setup({"size": "A5"}, author=BOT, at=ts(5))
        out = tmp_path / "pending.docx"
        aim.to_docx(basic_doc, out)  # default pending="tracked"
        got = Document(str(out))
        # the aim:doc proposal is not chunk content: no tracked-change runs
        ins = got.element.body.findall(".//" + qn("w:ins"))
        dels = got.element.body.findall(".//" + qn("w:del"))
        assert not ins and not dels

    def test_ingest_side_pass_reads_intent(self, tmp_path):
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Mm

        from aimformat.convert._docx_pages import (
            _read_break_anchors,
            _read_page_setup,
            apply_docx_pagination,
        )

        src = Document()
        section = src.sections[0]
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = section.bottom_margin = Mm(20)
        section.left_margin = section.right_margin = Mm(18)
        src.add_paragraph("Alpha paragraph.")
        beta = src.add_paragraph("Beta paragraph.")
        beta.add_run().add_break(WD_BREAK.PAGE)  # break after beta's text
        src.add_paragraph("Gamma paragraph.")
        delta = src.add_paragraph("Delta paragraph.")
        delta.paragraph_format.page_break_before = True
        path = tmp_path / "src.docx"
        src.save(str(path))

        assert _read_page_setup(Document(str(path))) == {
            "size": "A4",
            "orientation": "portrait",
            "margins": {"top": "20mm", "right": "18mm", "bottom": "20mm", "left": "18mm"},
        }
        assert _read_break_anchors(Document(str(path))) == [
            ("Beta paragraph.", 1),
            ("Gamma paragraph.", 1),
        ]

        doc = aim.new_document(title="Ingested")
        for i, txt in enumerate(
            ["Alpha paragraph.", "Beta paragraph.", "Gamma paragraph.", "Delta paragraph."]
        ):
            doc.add_chunk(f"<p>{txt}</p>", author=ME, at=ts(i))
        apply_docx_pagination(doc, path, author=ME)
        assert doc.page_setup.margins_mm["left"] == 18.0
        tags = [c.tag for c in doc.chunks]
        assert tags.count("aim-page-break") == 2
        texts = [(c.tag, c.text) for c in doc.chunks]
        assert texts.index(("aim-page-break", "")) == texts.index(("p", "Beta paragraph.")) + 1
        assert not [f for f in aim.lint_text(doc.dumps()) if f.level == "error"]
        assert doc.verify() == []

    def test_unmatched_hints_are_skipped(self, tmp_path):
        from docx import Document
        from docx.enum.text import WD_BREAK

        from aimformat.convert._docx_pages import apply_docx_pagination

        src = Document()
        src.add_paragraph("Only in the source.")
        src.add_paragraph("Also only here.").add_run().add_break(WD_BREAK.PAGE)
        path = tmp_path / "src.docx"
        src.save(str(path))
        doc = aim.new_document(title="Different content")
        doc.add_chunk("<p>Entirely different.</p>", author=ME, at=ts(0))
        apply_docx_pagination(doc, path, author=ME)
        assert "aim-page-break" not in [c.tag for c in doc.chunks]

    def test_duplicate_text_anchors_on_the_right_occurrence(self, tmp_path):
        # regression: first-match anchoring put the break after the FIRST
        # of two identical paragraphs when the source broke after the second
        from docx import Document
        from docx.enum.text import WD_BREAK

        from aimformat.convert._docx_pages import _read_break_anchors, apply_docx_pagination

        src = Document()
        src.add_paragraph("Repeat me.")
        src.add_paragraph("Repeat me.").add_run().add_break(WD_BREAK.PAGE)
        src.add_paragraph("After.")
        path = tmp_path / "dup.docx"
        src.save(str(path))
        assert _read_break_anchors(Document(str(path))) == [("Repeat me.", 2)]

        doc = aim.new_document(title="Dup")
        doc.add_chunk("<p>Repeat me.</p>", author=ME, at=ts(0))
        second = doc.add_chunk("<p>Repeat me.</p>", author=ME, at=ts(1))
        doc.add_chunk("<p>After.</p>", author=ME, at=ts(2))
        apply_docx_pagination(doc, path, author=ME)
        body = doc.body_ids
        breaks = [c.id for c in doc.chunks if c.tag == "aim-page-break"]
        assert len(breaks) == 1
        assert body.index(breaks[0]) == body.index(second.id) + 1

    def test_same_run_text_before_break_is_the_anchor(self, tmp_path):
        # regression: WordprocessingML allows <w:t>Beta</w:t><w:br/> in ONE
        # run; the prefix bookkeeping missed it and fell back to the
        # previous paragraph, importing the break one construct too early
        from docx import Document
        from docx.enum.text import WD_BREAK

        from aimformat.convert._docx_pages import _read_break_anchors

        src = Document()
        src.add_paragraph("Alpha.")
        run = src.add_paragraph().add_run("Beta.")
        run.add_break(WD_BREAK.PAGE)  # text and break share one run
        src.add_paragraph("Gamma.")
        path = tmp_path / "samerun.docx"
        src.save(str(path))
        assert _read_break_anchors(Document(str(path))) == [("Beta.", 1)]
