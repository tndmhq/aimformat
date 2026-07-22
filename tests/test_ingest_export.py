"""The document loop: DoclingDocument → .aim → DOCX.

docling-core builds faithful fixture documents (the same dict shape the full
docling converter emits); python-docx reads back what the exporter wrote.
Both are dev/test dependencies only — the package itself stays stdlib-only.
"""

import pytest

import aimformat as aim
from aimformat.errors import InvalidOperation
from conftest import BOT, ME, ts

docling_core = pytest.importorskip("docling_core")
docx = pytest.importorskip("docx")

from docling_core.types.doc import GroupLabel, TableCell, TableData  # noqa: E402
from docling_core.types.doc.document import DoclingDocument  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


def build_docling_fixture() -> DoclingDocument:
    d = DoclingDocument(name="pilot-report")
    d.add_title(text="Pilot report")
    d.add_text(label="text", text="We ran the pilot for six weeks & learned a lot.")
    d.add_heading(text="What worked", level=1)
    lst = d.add_group(label=GroupLabel.LIST)
    d.add_list_item(text="Review speed doubled", parent=lst)
    d.add_list_item(text="Fewer copy-paste errors", parent=lst)
    olst = d.add_group(label=GroupLabel.ORDERED_LIST)
    d.add_list_item(text="Roll out to briefs", parent=olst, enumerated=True)
    d.add_list_item(text="Then contracts", parent=olst, enumerated=True)
    d.add_heading(text="Numbers", level=2)
    td = TableData(
        num_rows=2,
        num_cols=2,
        table_cells=[
            TableCell(
                text="Metric",
                start_row_offset_idx=0,
                end_row_offset_idx=1,
                start_col_offset_idx=0,
                end_col_offset_idx=1,
                column_header=True,
            ),
            TableCell(
                text="Value",
                start_row_offset_idx=0,
                end_row_offset_idx=1,
                start_col_offset_idx=1,
                end_col_offset_idx=2,
                column_header=True,
            ),
            TableCell(
                text="Accepted edits",
                start_row_offset_idx=1,
                end_row_offset_idx=2,
                start_col_offset_idx=0,
                end_col_offset_idx=1,
            ),
            TableCell(
                text="38%",
                start_row_offset_idx=1,
                end_row_offset_idx=2,
                start_col_offset_idx=1,
                end_col_offset_idx=2,
            ),
        ],
    )
    d.add_table(data=td)
    d.add_text(label="code", text="aim lint report.aim")
    return d


@pytest.fixture
def ingested() -> aim.AimDocument:
    return aim.from_docling(build_docling_fixture())


class TestIngest:
    def test_accepts_object_and_dict_equivalently(self):
        src = build_docling_fixture()
        a = aim.from_docling(src)
        b = aim.from_docling(src.export_to_dict())
        assert [c.html.replace(c.id, "X") for c in a.chunks] == [
            c.html.replace(c.id, "X") for c in b.chunks
        ]

    def test_rejects_wrong_type(self):
        with pytest.raises(TypeError):
            aim.from_docling(42)

    def test_title_becomes_h1_and_doc_title(self, ingested):
        assert ingested.title == "Pilot report"
        assert ingested.chunks[0].tag == "h1"
        assert ingested.chunks[0].text == "Pilot report"

    def test_heading_levels_shift_below_title(self, ingested):
        tags = [c.tag for c in ingested.chunks]
        assert "h2" in tags and "h3" in tags  # docling levels 1,2 -> h2,h3

    def test_entities_escaped(self, ingested):
        para = next(c for c in ingested.chunks if c.tag == "p")
        assert "&amp;" in para.html and "&" in para.text

    def test_lists_become_containers_with_item_chunks(self, ingested):
        assert len(ingested.containers) == 3  # ul, ol, table
        lis = [c for c in ingested.chunks if c.tag == "li"]
        assert len(lis) == 4
        assert all(c.container in ingested.containers for c in lis)

    def test_ordered_list_is_ol(self, ingested):
        text = ingested.dumps()
        assert "<ol data-aim-container=" in text
        assert "<ul data-aim-container=" in text

    def test_table_grid_with_header(self, ingested):
        rows = [c for c in ingested.chunks if c.tag == "tr"]
        assert len(rows) == 2
        text = ingested.dumps()
        assert "<thead><tr" in text and "<th>Metric</th>" in text
        assert "<td>38%</td>" in text

    def test_code_becomes_pre(self, ingested):
        assert any(c.tag == "pre" for c in ingested.chunks)

    def test_ingestion_is_recorded_history(self, ingested):
        events = ingested.history
        assert events and all(e.action == "add" for e in events)
        assert all(e.author.type == "external" for e in events)
        assert len({e.batch for e in events}) == 1  # one ingestion batch

    def test_ingested_doc_lints_clean_and_verifies(self, ingested):
        assert ingested.verify() == []
        assert not [f for f in aim.lint_text(ingested.dumps()) if f.level == "error"]

    def test_picture_with_caption(self):
        d = DoclingDocument(name="pics")
        pic = d.add_picture()
        d.add_text(label="caption", text="Figure 1: the setup", parent=pic)
        obj = d.export_to_dict()
        # wire the caption ref the way docling emits it
        obj["pictures"][0]["captions"] = [{"$ref": obj["texts"][0]["self_ref"]}]
        doc = aim.from_docling(obj)
        fig = next(c for c in doc.chunks if c.tag == "figure")
        assert "figcaption" in fig.html and "Figure 1" in fig.text

    def test_colspan_survives(self):
        d = DoclingDocument(name="spans")
        td = TableData(
            num_rows=2,
            num_cols=2,
            table_cells=[
                TableCell(
                    text="Wide",
                    start_row_offset_idx=0,
                    end_row_offset_idx=1,
                    start_col_offset_idx=0,
                    end_col_offset_idx=2,
                    col_span=2,
                    column_header=True,
                ),
                TableCell(
                    text="a",
                    start_row_offset_idx=1,
                    end_row_offset_idx=2,
                    start_col_offset_idx=0,
                    end_col_offset_idx=1,
                ),
                TableCell(
                    text="b",
                    start_row_offset_idx=1,
                    end_row_offset_idx=2,
                    start_col_offset_idx=1,
                    end_col_offset_idx=2,
                ),
            ],
        )
        d.add_table(data=td)
        doc = aim.from_docling(d)
        assert 'colspan="2"' in doc.dumps()

    def test_furniture_skipped(self):
        d = DoclingDocument(name="f")
        from docling_core.types.doc.document import ContentLayer

        d.add_text(label="text", text="page header", content_layer=ContentLayer.FURNITURE)
        d.add_text(label="text", text="real content")
        doc = aim.from_docling(d)
        texts = [c.text for c in doc.chunks]
        assert texts == ["real content"]


def _docx_paragraphs(path):
    return [(p.style.name, p.text) for p in docx.Document(str(path)).paragraphs]


def _paragraph_sequence(path):
    """Body paragraphs in order: 'BREAK' for page-break paragraphs, else the
    paragraph's full text (including runs inside w:ins revisions)."""
    seq = []
    for p in docx.Document(str(path)).paragraphs:
        if p._p.findall(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']"):
            seq.append("BREAK")
            continue
        text = "".join(t.text or "" for t in p._p.iter(qn("w:t")))
        if text:
            seq.append(text)
    return seq


def _revision_authors(path, tag):
    d = docx.Document(str(path))
    return [el.get(qn("w:author")) for el in d.element.body.iter(qn(f"w:{tag}"))]


def _table_row_texts(path):
    """Per-table, per-row text in XML order (runs inside w:ins included —
    python-docx cell.text does not see them)."""
    d = docx.Document(str(path))
    return [
        ["".join(t.text or "" for t in tr.iter(qn("w:t"))) for tr in tbl._tbl.findall(qn("w:tr"))]
        for tbl in d.tables
    ]


def _ins_del_para_texts(path):
    """Per-paragraph inserted (w:ins) and deleted (w:del) text, in order."""
    ins, dele = [], []
    for p in docx.Document(str(path)).paragraphs:
        it = "".join(t.text or "" for w in p._p.iter(qn("w:ins")) for t in w.iter(qn("w:t")))
        dt = "".join(t.text or "" for w in p._p.iter(qn("w:del")) for t in w.iter(qn("w:delText")))
        if it:
            ins.append(it)
        if dt:
            dele.append(dt)
    return ins, dele


class TestExportDocx:
    def test_structure_maps_to_word_styles(self, ingested, tmp_path):
        out = aim.to_docx(ingested, tmp_path / "out.docx")
        paras = _docx_paragraphs(out)
        assert ("Heading 1", "Pilot report") in paras
        assert ("Heading 2", "What worked") in paras
        assert ("List Bullet", "Review speed doubled") in paras
        assert ("List Number", "Roll out to briefs") in paras

    def test_table_content_and_header_bold(self, ingested, tmp_path):
        out = aim.to_docx(ingested, tmp_path / "t.docx")
        tables = docx.Document(str(out)).tables
        assert len(tables) == 1
        assert tables[0].cell(0, 0).text == "Metric"
        assert tables[0].cell(1, 1).text == "38%"
        head_run = tables[0].cell(0, 0).paragraphs[0].runs[0]
        assert head_run.bold

    def test_inline_marks(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<p data-aim="p1">plain <strong>bold</strong> <em>ital</em> <code>mono</code></p>',
            author=BOT,
            at=ts(0),
        )
        out = aim.to_docx(doc, tmp_path / "m.docx")
        para = docx.Document(str(out)).paragraphs[0]
        runs = {r.text: r for r in para.runs}
        assert runs["bold"].bold and runs["ital"].italic
        assert runs["mono"].font.name == "Consolas"

    def test_pending_tracked_emits_ins_and_del(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">The old wording.</p>', author=ME, at=ts(0))
        doc.propose_modify(
            "p1",
            '<p data-aim="p1">The new wording.</p>',
            author=BOT,
            explanation="better",
            at=ts(1),
        )
        out = aim.to_docx(doc, tmp_path / "tracked.docx", pending="tracked")
        ins, dele = _revision_authors(out, "ins"), _revision_authors(out, "del")
        assert ins and dele
        assert ins[0].startswith("agent:") and dele[0].startswith("agent:")
        # original document untouched
        assert doc.proposals and doc.chunk("p1").text == "The old wording."

    def test_pending_add_tracked_as_inserted_paragraph(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Anchor.</p>', author=ME, at=ts(0))
        doc.propose_add("<p>Brand new paragraph.</p>", author=BOT, after="p1", at=ts(1))
        out = aim.to_docx(doc, tmp_path / "add.docx")
        d = docx.Document(str(out))
        # runs inside w:ins are invisible to Paragraph.text — check the XML
        xml = d.element.body.xml
        assert "Brand new paragraph." in xml
        assert _revision_authors(out, "ins")

    @pytest.mark.parametrize("tag,style", [("ul", "List Bullet"), ("ol", "List Number")])
    def test_pending_list_add_tracked_keeps_items(self, tmp_path, tag, style):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">base</p>', author=ME, at=ts(0))
        doc.propose_add(
            f"<{tag}><li>Item one</li><li>Item two</li><li>Item three</li></{tag}>",
            author=BOT,
            after="p1",
            at=ts(1),
        )
        out = aim.to_docx(doc, tmp_path / "add-list.docx", pending="tracked")
        # one inserted paragraph per item, never a concatenated flattening
        assert _paragraph_sequence(out) == ["base", "Item one", "Item two", "Item three"]
        assert "Item oneItem two" not in docx.Document(str(out)).element.body.xml
        item_styles = [
            p.style.name
            for p in docx.Document(str(out)).paragraphs
            if "Item" in "".join(t.text or "" for t in p._p.iter(qn("w:t")))
        ]
        assert item_styles == [style] * 3
        authors = _revision_authors(out, "ins")
        assert authors and all(a.startswith("agent:") for a in authors)

    def test_pending_table_add_tracked_emits_table(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">base</p>', author=ME, at=ts(0))
        doc.propose_add(
            "<table><tr><td>C1</td><td>C2</td></tr><tr><td>C3</td><td>C4</td></tr></table>",
            author=BOT,
            after="p1",
            at=ts(1),
        )
        out = aim.to_docx(doc, tmp_path / "add-table.docx", pending="tracked")
        d = docx.Document(str(out))
        assert len(d.tables) == 1
        for (r, c), text in {(0, 0): "C1", (0, 1): "C2", (1, 0): "C3", (1, 1): "C4"}.items():
            cell_xml = d.tables[0].cell(r, c)._tc.xml
            assert "w:ins" in cell_xml and text in cell_xml
        # nothing leaks into flattened body paragraphs
        assert _paragraph_sequence(out) == ["base"]

    def test_pending_container_add_matches_accept_all_texts(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">base</p>', author=ME, at=ts(0))
        p = doc.propose_add(
            "<ul><li>Item one</li><li>Item two</li></ul>", author=BOT, after="p1", at=ts(1)
        )
        doc.propose_add(
            "<table><tr><td>C1</td><td>C2</td></tr></table>", author=BOT, after=p.id, at=ts(2)
        )
        tracked = aim.to_docx(doc, tmp_path / "t.docx", pending="tracked")
        accepted = aim.to_docx(doc, tmp_path / "a.docx", pending="accept-all")
        assert _paragraph_sequence(tracked) == _paragraph_sequence(accepted)
        assert _table_row_texts(tracked) == _table_row_texts(accepted)

    def test_pending_tracked_preserves_br(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Old A<br>Old B</p>', author=ME, at=ts(0))
        doc.add_chunk('<p data-aim="p2">Doomed X<br>Doomed Y</p>', author=ME, at=ts(1))
        doc.propose_modify(
            "p1", '<p data-aim="p1">New line A<br>New line B</p>', author=BOT, at=ts(2)
        )
        doc.propose_delete("p2", author=BOT, at=ts(3))
        out = aim.to_docx(doc, tmp_path / "br.docx", pending="tracked")
        body = docx.Document(str(out)).element.body

        def run_items(wrap):
            items = []
            for r in wrap.iter(qn("w:r")):
                for child in r:
                    if child.tag in (qn("w:t"), qn("w:delText")):
                        items.append(child.text)
                    elif child.tag == qn("w:br"):
                        items.append("BR")
            return items

        ins_seqs = [run_items(w) for w in body.iter(qn("w:ins"))]
        del_seqs = [run_items(w) for w in body.iter(qn("w:del"))]
        assert ["New line A", "BR", "New line B"] in ins_seqs
        assert ["Doomed X", "BR", "Doomed Y"] in del_seqs

    @pytest.mark.parametrize(
        "shape", ["single-mid", "chained-mid", "siblings-mid", "chained-last", "chained-start"]
    )
    def test_pending_row_adds_keep_accept_all_order(self, tmp_path, shape):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl">'
            '<tr data-aim="r1"><td>one</td><td>1</td></tr>'
            '<tr data-aim="r2"><td>two</td><td>2</td></tr></table>',
            author=ME,
            at=ts(0),
        )

        def add(after, text, i):
            return doc.propose_add(
                f"<tr><td>{text}</td><td>{text.lower()}</td></tr>",
                author=BOT,
                container="tbl",
                after=after,
                at=ts(i),
            )

        if shape == "single-mid":
            add("r1", "A", 1)
        elif shape == "chained-mid":
            add(add("r1", "A", 1).id, "B", 2)
        elif shape == "siblings-mid":
            add("r1", "A", 1)
            add("r1", "B", 2)
        elif shape == "chained-last":
            add(add("r2", "A", 1).id, "B", 2)
        else:  # chained on a container-start add
            add(add(None, "A", 1).id, "B", 2)
        tracked = aim.to_docx(doc, tmp_path / "t.docx", pending="tracked")
        accepted = aim.to_docx(doc, tmp_path / "a.docx", pending="accept-all")
        # tracked row order must match what accepting the changes yields
        assert _table_row_texts(tracked) == _table_row_texts(accepted)
        trs = docx.Document(str(tracked)).tables[0]._tbl.findall(qn("w:tr"))
        added = 0
        for tr in trs:
            ins_ts = {id(t) for w in tr.iter(qn("w:ins")) for t in w.iter(qn("w:t"))}
            all_ts = list(tr.iter(qn("w:t")))
            assert all_ts  # no empty orphan rows
            inside = [t for t in all_ts if id(t) in ins_ts]
            assert len(inside) in (0, len(all_ts))  # never mixes inserted and original
            added += bool(inside)
        assert added == (1 if shape == "single-mid" else 2)

    def test_run_chunk_modify_tracked_emits_payload_once(self, rich_doc, tmp_path):
        rich_doc.propose_modify("li2", '<li data-aim="li2">Rewritten</li>', author=BOT, at=ts(9))
        out = aim.to_docx(rich_doc, tmp_path / "run.docx", pending="tracked")
        ins, dele = _ins_del_para_texts(out)
        assert ins == ["Rewritten"]
        assert dele == ["Second, part one…", "…second, part two"]
        acc = aim.to_docx(rich_doc, tmp_path / "acc.docx", pending="accept-all")
        texts = [p.text for p in docx.Document(str(acc)).paragraphs]
        assert texts.count("Rewritten") == 1

    def test_run_chunk_delete_tracked_stays_delete_only(self, rich_doc, tmp_path):
        rich_doc.propose_delete("li2", author=BOT, at=ts(9))
        out = aim.to_docx(rich_doc, tmp_path / "rundel.docx", pending="tracked")
        ins, dele = _ins_del_para_texts(out)
        assert ins == []
        assert dele == ["Second, part one…", "…second, part two"]

    def test_run_chunk_row_modify_tracked_emits_payload_once(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl">'
            '<tr data-aim="r1"><td>one</td></tr>'
            '<tr data-aim="rr"><td>two-a</td></tr>'
            '<tr data-aim="rr"><td>two-b</td></tr></table>',
            author=ME,
            at=ts(0),
        )
        doc.propose_modify("rr", '<tr data-aim="rr"><td>NEWCELL</td></tr>', author=BOT, at=ts(1))
        out = aim.to_docx(doc, tmp_path / "trrun.docx", pending="tracked")
        body = docx.Document(str(out)).element.body
        ins_texts = [
            "".join(t.text or "" for t in w.iter(qn("w:t"))) for w in body.iter(qn("w:ins"))
        ]
        assert ins_texts.count("NEWCELL") == 1

    def test_pending_delete_tracked(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Doomed.</p>', author=ME, at=ts(0))
        doc.propose_delete("p1", author=BOT, at=ts(1))
        out = aim.to_docx(doc, tmp_path / "del.docx")
        assert _revision_authors(out, "del")

    def test_accept_all_resolves_on_copy(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Old.</p>', author=ME, at=ts(0))
        doc.propose_modify("p1", '<p data-aim="p1">New.</p>', author=BOT, at=ts(1))
        out = aim.to_docx(doc, tmp_path / "acc.docx", pending="accept-all")
        texts = [p.text for p in docx.Document(str(out)).paragraphs]
        assert "New." in texts and "Old." not in texts
        assert doc.proposals  # original untouched

    def test_reject_all(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Old.</p>', author=ME, at=ts(0))
        doc.propose_modify("p1", '<p data-aim="p1">New.</p>', author=BOT, at=ts(1))
        out = aim.to_docx(doc, tmp_path / "rej.docx", pending="reject-all")
        texts = [p.text for p in docx.Document(str(out)).paragraphs]
        assert "Old." in texts and "New." not in texts

    def test_accept_all_resolves_chained_adds_in_order(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Anchor.</p>', author=ME, at=ts(0))
        p1 = doc.propose_add('<p data-aim="n1">One.</p>', author=BOT, after="p1", at=ts(1))
        doc.propose_add('<p data-aim="n2">Two.</p>', author=BOT, after=p1.id, at=ts(2))
        out = aim.to_docx(doc, tmp_path / "chain.docx", pending="accept-all")
        texts = [p.text for p in docx.Document(str(out)).paragraphs]
        assert texts.index("One.") < texts.index("Two.")

    def test_row_modify_tracked_in_cells(self, rich_doc, tmp_path):
        rich_doc.propose_modify(
            "row1", '<tr data-aim="row1"><td>alpha</td><td>99</td></tr>', author=BOT, at=ts(20)
        )
        out = aim.to_docx(rich_doc, tmp_path / "row.docx")
        assert _revision_authors(out, "ins") and _revision_authors(out, "del")

    def test_invalid_pending_mode(self, ingested, tmp_path):
        with pytest.raises(InvalidOperation):
            aim.to_docx(ingested, tmp_path / "x.docx", pending="merge")

    def test_full_loop_docling_to_aim_to_docx(self, ingested, tmp_path):
        p = ingested.propose_modify(
            next(c.id for c in ingested.chunks if c.tag == "p"),
            "<p>We ran the pilot for six weeks and learned plenty.</p>",
            author=BOT,
            explanation="Tighter.",
            at=ts(30),
        )
        ingested.accept(p.id, decided_by=ME, at=ts(31))
        assert ingested.verify() == []
        out = aim.to_docx(ingested, tmp_path / "loop.docx", pending="accept-all")
        text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
        assert "learned plenty" in text
        assert "Pilot report" in text

    def test_figure_caption_and_hr(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<figure data-aim="f"><img alt="chart" '
            'src="https://example.org/x.png">'
            "<figcaption>The chart.</figcaption></figure>",
            author=BOT,
            at=ts(0),
        )
        doc.add_chunk('<hr data-aim="rule">', author=BOT, at=ts(1))
        out = aim.to_docx(doc, tmp_path / "fig.docx")
        text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
        assert "[image: chart]" in text and "The chart." in text


_PNG_1PX = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4nGNgYGBg"
    "AAAABQABh6FO1AAAAABJRU5ErkJggg=="
)


def _deck(with_flow=True):
    doc = aim.new_document(title="Deck")
    if with_flow:
        doc.add_chunk('<p data-aim="intro">Before the pages.</p>', author=ME, at=ts(0))
    doc.add_chunk(
        '<aim-slide data-aim-container="s1" style="width:420px; height:595px">'
        '<h2 data-aim="t1" style="left:10px; top:10px; width:300px">Page one</h2>'
        '<p data-aim="b1" style="left:10px; top:60px; width:300px">First body.</p>'
        "</aim-slide>",
        author=ME,
        at=ts(1),
    )
    doc.add_chunk(
        '<aim-slide data-aim-container="s2" style="width:420px; height:595px">'
        '<p data-aim="b2" style="left:10px; top:10px; width:300px">Second body.</p>'
        "</aim-slide>",
        author=ME,
        at=ts(2),
    )
    return doc


class TestExportDocxSlides:
    def test_slides_linearize_instead_of_dropping(self, tmp_path):
        out = aim.to_docx(_deck(), tmp_path / "deck.docx")
        paras = _docx_paragraphs(out)
        assert ("Heading 2", "Page one") in paras
        texts = [t for _, t in paras]
        assert "First body." in texts and "Second body." in texts
        # reading order preserved across the linearization
        assert texts.index("Page one") < texts.index("First body.") < texts.index("Second body.")

    def test_page_breaks_frame_slides(self, tmp_path):
        out = aim.to_docx(_deck(), tmp_path / "deck.docx")
        d = docx.Document(str(out))
        breaks = d.element.body.findall(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']")
        # one before each slide (after the intro / between slides), none trailing
        assert len(breaks) == 2

    def test_leading_slide_gets_no_blank_first_page(self, tmp_path):
        out = aim.to_docx(_deck(with_flow=False), tmp_path / "deck.docx")
        d = docx.Document(str(out))
        breaks = d.element.body.findall(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']")
        assert len(breaks) == 1  # only between the two slides

    def test_flow_after_slide_starts_on_fresh_page(self, tmp_path):
        doc = _deck(with_flow=False)
        doc.add_chunk('<p data-aim="after">Back to flow.</p>', author=ME, at=ts(3))
        out = aim.to_docx(doc, tmp_path / "deck.docx")
        d = docx.Document(str(out))
        breaks = d.element.body.findall(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']")
        assert len(breaks) == 2  # s1→s2 and s2→flow

    def test_in_slide_pending_modify_rides_tracked_changes(self, tmp_path):
        doc = _deck()
        doc.propose_modify(
            "b1",
            '<p data-aim="b1" style="left:10px; top:60px; width:300px">Sharper body.</p>',
            author=BOT,
            at=ts(3),
        )
        out = aim.to_docx(doc, tmp_path / "deck.docx", pending="tracked")
        assert _revision_authors(out, "ins") and _revision_authors(out, "del")
        xml = docx.Document(str(out)).element.body.xml
        assert "Sharper body." in xml and "First body." in xml

    def test_in_slide_pending_add_emits_in_place(self, tmp_path):
        doc = _deck()
        doc.propose_add(
            '<p style="left:10px; top:110px; width:300px">Added into the page.</p>',
            container="s1",
            after="b1",
            author=BOT,
            at=ts(3),
        )
        out = aim.to_docx(doc, tmp_path / "deck.docx", pending="tracked")
        xml = docx.Document(str(out)).element.body.xml
        assert "Added into the page." in xml
        texts = [t for _, t in _docx_paragraphs(out)]
        joined = "\n".join(xml.split("Second body.")[0:1])
        assert "Added into the page." in joined  # lands inside s1, not at the end
        assert _revision_authors(out, "ins")
        assert texts  # document is non-empty for the plain reader too

    def test_body_add_after_slide_lands_on_next_page(self, tmp_path):
        doc = _deck(with_flow=False)
        doc.propose_add("<p>Between the pages.</p>", after="s1", author=BOT, at=ts(3))
        out = aim.to_docx(doc, tmp_path / "deck.docx", pending="tracked")
        kinds = _paragraph_sequence(out)
        i_first = kinds.index("First body.")
        i_ins = kinds.index("Between the pages.")
        i_second = kinds.index("Second body.")
        assert i_first < i_ins < i_second
        # the page break comes BEFORE the tracked insertion (the add belongs
        # to the page after the slide, like accepted content) and slide two
        # still opens its own page
        assert "BREAK" in kinds[i_first + 1 : i_ins]
        assert "BREAK" in kinds[i_ins + 1 : i_second]

    def test_explicit_page_break_before_slide_stays_single(self, tmp_path):
        doc = aim.new_document(title="Breaks")
        doc.add_chunk("<p>Before.</p>", author=ME, at=ts(0))
        doc.add_chunk("<aim-page-break></aim-page-break>", author=ME, at=ts(1))
        doc.add_chunk(
            '<aim-slide style="width:420px; height:595px">'
            '<p style="left:10px; top:10px; width:300px">Page.</p>'
            "</aim-slide>",
            author=ME,
            at=ts(2),
        )
        out = aim.to_docx(doc, tmp_path / "breaks.docx")
        d = docx.Document(str(out))
        breaks = d.element.body.findall(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']")
        # the explicit break already pages: the slide must not add a second
        # one (two consecutive breaks render a blank page in Word)
        assert len(breaks) == 1

    def test_pending_whole_slide_add_linearizes_per_block(self, tmp_path):
        doc = _deck(with_flow=False)
        doc.propose_add(
            '<aim-slide style="width:420px; height:595px">'
            '<h2 style="left:10px; top:10px; width:300px">New page</h2>'
            '<p style="left:10px; top:60px; width:300px">New body.</p>'
            "</aim-slide>",
            after="s2",
            author=BOT,
            at=ts(3),
        )
        out = aim.to_docx(doc, tmp_path / "deck.docx", pending="tracked")
        d = docx.Document(str(out))
        styled = [
            (p.style.name, "".join(t.text or "" for t in p._p.iter(qn("w:t"))))
            for p in d.paragraphs
        ]
        # two separate inserted blocks — heading styled — not one collapsed blob
        assert ("Heading 2", "New page") in styled
        texts = [t for _, t in styled]
        assert "New body." in texts
        assert not any("New page" in t and "New body." in t for t in texts)
        assert _revision_authors(out, "ins")
        # the inserted page opens after slide two's page: s1→s2 and s2→add
        kinds = _paragraph_sequence(out)
        assert kinds.count("BREAK") == 2
        assert kinds.index("New page") > kinds.index("Second body.")

    def test_accept_all_deck_exports_clean(self, tmp_path):
        doc = _deck()
        doc.propose_modify(
            "b2",
            '<p data-aim="b2" style="left:10px; top:10px; width:300px">Final body.</p>',
            author=BOT,
            at=ts(3),
        )
        out = aim.to_docx(doc, tmp_path / "deck.docx", pending="accept-all")
        texts = [t for _, t in _docx_paragraphs(out)]
        assert "Final body." in texts and "Second body." not in texts


class TestExportDocxFigureWidth:
    def _fig_doc(self, style: str | None):
        doc = aim.new_document(title="T")
        s = f' style="{style}"' if style else ""
        doc.add_chunk(
            f'<figure data-aim="f"><img alt="dot"{s} '
            f'src="data:image/png;base64,{_PNG_1PX}"></figure>',
            author=ME,
            at=ts(0),
        )
        return doc

    def _picture_width_emu(self, path):
        d = docx.Document(str(path))
        assert d.inline_shapes, "no picture was embedded"
        return d.inline_shapes[0].width

    def test_authored_width_is_honored(self, tmp_path):
        out = aim.to_docx(self._fig_doc("width:300px"), tmp_path / "f.docx")
        emu = self._picture_width_emu(out)
        assert abs(emu - int(914400 * 300 / 96)) <= 914400 // 96  # 3.125in ±1px

    def test_default_width_when_unspecified(self, tmp_path):
        out = aim.to_docx(self._fig_doc(None), tmp_path / "f.docx")
        assert abs(self._picture_width_emu(out) - int(4.5 * 914400)) <= 9144

    def test_oversized_width_clamps_to_content_box(self, tmp_path):
        out = aim.to_docx(self._fig_doc("width:2000px"), tmp_path / "f.docx")
        # A4 portrait, 15mm margins → 180mm content ≈ 7.09in
        assert self._picture_width_emu(out) <= int(7.1 * 914400)


class TestDocxTextColour:
    """DOCX dropped text colour entirely — export_docx emitted no w:color at
    all, so a document whose title had been branded (or given an inline colour)
    exported in Word's default ink.

    Reported from production: "the title is not in pink when downloaded in
    docx" (2026-07-21).
    """

    @staticmethod
    def _colours(path):
        import re
        import zipfile

        xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode()
        return sorted(set(re.findall(r'<w:color w:val="([0-9A-Fa-f]{6})"', xml)))

    @staticmethod
    def _colour_count(path):
        import re
        import zipfile

        xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode()
        return len(re.findall(r'<w:color w:val="[0-9A-Fa-f]{6}"', xml))

    def _export(self, markup, tmp_path, theme=None):
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk(markup, author=aim.external("test"))
        text = doc.dumps()
        if theme:
            text = text.replace("</head>", f"<style data-aim-theme>:root{{{theme}}}</style></head>")
        out = tmp_path / "out.docx"
        aim.to_docx(aim.loads(text), out, pending="accept-all")
        return out

    def test_a_brand_class_reaches_word_as_a_run_colour(self, tmp_path):
        out = self._export('<h1 class="text-brand-1">Title</h1>', tmp_path)
        assert "1D4ED8" in self._colours(out)  # the default brand-1

    def test_the_documents_theme_wins_over_the_default(self, tmp_path):
        """The reported case end to end: the model sets the brand slot to pink
        and puts the class on the heading. Word must get the pink, not the
        stylesheet default."""
        out = self._export(
            '<h1 class="text-brand-1">Quarterly Review</h1>',
            tmp_path,
            theme="--aim-brand-1:#ff1493",
        )
        assert self._colours(out) == ["FF1493"]

    def test_a_registered_palette_class_resolves(self, tmp_path):
        out = self._export('<p class="text-red-600">Red</p>', tmp_path)
        assert "DC2626" in self._colours(out)  # registry palette red-600

    def test_an_rgb_theme_value_resolves(self, tmp_path):
        """registry.json permits rgb() for a colour slot, so the exporter must
        convert it rather than silently drop the colour (Codex aimformat#19)."""
        out = self._export(
            '<h1 class="text-brand-1">Title</h1>',
            tmp_path,
            theme="--aim-brand-1:rgb(255, 20, 147)",
        )
        assert self._colours(out) == ["FF1493"]

    def test_inline_paint_reaches_word(self, tmp_path):
        """The literal is a first-class value since 0.3 — the point of the
        change. It also has to be lint-clean, or the exporter would render
        markup the format does not accept."""
        markup = '<p style="color:#ff69b4">Pink</p>'
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk(markup, author=aim.external("test"))
        assert [f for f in aim.lint_text(doc.dumps()) if f.level == "error"] == []
        assert self._colours(self._export(markup, tmp_path)) == ["FF69B4"]

    def test_inline_paint_outranks_a_conflicting_class(self, tmp_path):
        out = self._export(
            '<p class="text-brand-1 text-red-600" style="color:#ff69b4">Pink</p>', tmp_path
        )
        assert self._colours(out) == ["FF69B4"]

    def test_an_uncoloured_document_gains_no_colour_runs(self, tmp_path):
        """No colour must mean no w:color — writing a default would override
        the Word theme the user's template supplies."""
        out = self._export("<p>Plain</p>", tmp_path)
        assert self._colours(out) == []

    def test_a_colour_we_cannot_resolve_is_left_alone(self, tmp_path):
        """Named colours are outside the paint grammar (V008) and are not
        guessed at either: putting the WRONG colour in the file is worse than
        leaving Word's default."""
        markup = '<p style="color:rebeccapurple">Named</p>'
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk(markup, author=aim.external("test"))
        assert "V008" in {f.code for f in aim.lint_text(doc.dumps())}
        assert self._colours(self._export(markup, tmp_path)) == []

    def test_export_never_mutates_the_source(self, tmp_path):
        """A tracked export uses the source document directly, so nothing in
        the exporter may write back into it (Codex aimformat#19)."""
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk(
            '<div class="text-brand-1"><p data-aim="p1">Child</p></div>',
            author=aim.external("test"),
        )
        before = doc.dumps()
        aim.to_docx(doc, tmp_path / "o.docx", pending="tracked")
        assert doc.dumps() == before

    def test_colour_inherits_from_a_wrapper(self, tmp_path):
        """CSS inherits colour, so a browser renders this child red and DOCX
        must agree. Every leaf emitter reads one resolver rather than each
        re-deriving the colour from its own element (the previous shape, where
        three review rounds each found a different leaf unthreaded)."""
        out = self._export('<div class="text-red-600"><p>Child</p></div>', tmp_path)
        assert self._colours(out) == ["DC2626"]

    def test_unversioned_body_paint_is_not_exported(self, tmp_path):
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk('<p data-aim="p1">Child</p>', author=aim.external("test"))
        painted = aim.loads(doc.dumps().replace("<body>", '<body style="color:#ff69b4">'))
        out = tmp_path / "body-paint.docx"
        aim.to_docx(painted, out, pending="accept-all")
        assert self._colours(out) == []

    def test_inherited_colour_reaches_every_leaf_family(self, tmp_path):
        out = self._export(
            '<section style="color:#ff69b4">'
            "<p>Block</p>"
            "<ul><li>Item</li></ul>"
            "<pre><code>code</code></pre>"
            "<table><tbody><tr><td>Cell</td></tr></tbody></table>"
            "</section>",
            tmp_path,
        )
        assert self._colours(out) == ["FF69B4"]
        assert self._colour_count(out) == 4  # one per leaf run, none skipped

    def test_a_leaf_whose_base_layer_sets_its_own_colour_does_not_inherit(self, tmp_path):
        """`blockquote{color:#4b5563}` and `figcaption{color:#6b7280}` beat an
        inherited value in a browser too, so neither is pink there. We paint
        nothing and Word's Quote/Caption styles keep their own ink."""
        out = self._export(
            '<section style="color:#ff69b4">'
            "<blockquote>Quoted</blockquote>"
            '<figure><img src="https://e.example/i.png" alt="x">'
            "<figcaption>Cap</figcaption></figure>"
            "</section>",
            tmp_path,
        )
        assert self._colour_count(out) == 1  # the image placeholder only

    def test_a_group_paints_its_own_blockquote_when_asked_directly(self, tmp_path):
        out = self._export(
            '<section><blockquote style="color:#ff69b4">Quoted</blockquote></section>', tmp_path
        )
        assert self._colours(out) == ["FF69B4"]

    def test_a_links_own_base_layer_colour_is_left_to_word(self, tmp_path):
        """`a{color:var(--aim-brand-1)}` is a base-layer rule: the link is not
        red in a browser either, and Word's template owns hyperlink ink. A
        fragment href emits no trailing URL run, so the count is exactly the
        surrounding text."""
        out = self._export('<p class="text-red-600">see <a href="#x">link</a></p>', tmp_path)
        assert self._colours(out) == ["DC2626"]
        assert self._colour_count(out) == 1

    def test_the_cascade_winner_is_the_last_sorted_class(self, tmp_path):
        """generate_aim_css sorts class rules by name and CSS is last-wins, so
        `text-brand-1 text-red-600` renders RED whichever is written first.
        DOCX must agree (Codex aimformat#19)."""
        assert self._colours(
            self._export('<p class="text-brand-1 text-red-600">X</p>', tmp_path)
        ) == ["DC2626"]
        assert self._colours(
            self._export('<p class="text-red-600 text-brand-1">X</p>', tmp_path)
        ) == ["DC2626"]

    def test_an_out_of_range_rgb_theme_value_is_clamped(self, tmp_path):
        out = self._export(
            '<h1 class="text-brand-1">T</h1>', tmp_path, theme="--aim-brand-1:rgb(999, 0, 0)"
        )
        assert self._colours(out) == ["FF0000"]

    def test_a_text_size_utility_is_not_treated_as_a_colour(self, tmp_path):
        out = self._export('<p class="text-xl">Plain</p>', tmp_path)
        assert self._colours(out) == []

    def test_text_white_resolves(self, tmp_path):
        """text-white lives in classes.singles, not classes.palette, so an
        enumeration of brand slots plus palette families missed it (Codex
        aimformat#19). Colours are read from the generated declarations now."""
        out = self._export('<p class="text-white">On dark</p>', tmp_path)
        assert self._colours(out) == ["FFFFFF"]

    def test_a_code_block_carries_its_own_colour(self, tmp_path):
        """emit_pre builds mono runs directly, bypassing the colour path, so a
        class applied straight to the <pre> was ignored (Codex aimformat#19)."""
        out = self._export('<pre class="text-red-600">Error</pre>', tmp_path)
        assert "DC2626" in self._colours(out)

    def test_a_list_items_own_colour_survives(self, tmp_path):
        """emit_list rebuilds each item as a synthetic <li> carrying only its
        children, which dropped the item's own class (Codex aimformat#19)."""
        out = self._export('<ul><li class="text-red-600">Red</li></ul>', tmp_path)
        assert "DC2626" in self._colours(out)

    def test_colour_on_nested_code_in_a_pre(self, tmp_path):
        """class is legal on <code> and browsers colour the code text from it,
        but emit_pre only checked the outer <pre> (Codex aimformat#19)."""
        out = self._export('<pre><code class="text-red-600">Error</code></pre>', tmp_path)
        assert "DC2626" in self._colours(out)

    def test_a_figure_caption_keeps_its_colour(self, tmp_path):
        """Captions were written with add_paragraph(text), bypassing runs
        entirely, so no formatting reached them (Codex aimformat#19)."""
        out = self._export(
            '<figure><img src="https://e.example/i.png" alt="x">'
            '<figcaption class="text-red-600">Cap</figcaption></figure>',
            tmp_path,
        )
        assert "DC2626" in self._colours(out)

    def test_a_grouping_elements_direct_colour_survives(self, tmp_path):
        """_block_children rewraps a group's loose text into a fresh, bare
        element, dropping a colour declared on the group (Codex #19)."""
        out = self._export('<blockquote class="text-red-600">Quoted</blockquote>', tmp_path)
        assert "DC2626" in self._colours(out)

    def test_a_mixed_pre_paints_only_the_coloured_child(self, tmp_path):
        """emit_pre used to flatten the whole block to one run, so it could
        either paint the sibling text too or paint nothing; it chose nothing
        (Codex #19). Per-element paint closes that hole: the run carries its
        own colour and the block guesses nothing."""
        out = self._export('<pre>plain <code class="text-red-600">x</code></pre>', tmp_path)
        assert self._colours(out) == ["DC2626"]
        assert self._colour_count(out) == 1  # the <code> run, not the sibling text
        only = self._export('<pre><code class="text-red-600">x</code></pre>', tmp_path)
        assert "DC2626" in self._colours(only)

    def test_a_nested_span_paints_only_its_own_run(self, tmp_path):
        out = self._export('<p>plain <span style="color:#ff69b4">pink</span> tail</p>', tmp_path)
        assert self._colours(out) == ["FF69B4"]
        assert self._colour_count(out) == 1

    def test_a_slide_child_keeps_its_paint_after_linearization(self, tmp_path):
        out = self._export(
            '<aim-slide style="width:960px; height:540px">'
            '<h2 style="left:48px; top:32px; width:450px; color:#ff69b4">Slide title</h2>'
            "</aim-slide>",
            tmp_path,
        )
        assert self._colours(out) == ["FF69B4"]

    def test_a_table_cells_own_paint_survives(self, tmp_path):
        out = self._export(
            "<table><tbody><tr>"
            '<td style="color:#ff69b4">Pink</td><td>Plain</td>'
            "</tr></tbody></table>",
            tmp_path,
        )
        assert self._colours(out) == ["FF69B4"]
        assert self._colour_count(out) == 1


class TestDocxPaintBoxes:
    """Backgrounds and borders. Word has no CSS box, so these are honest
    approximations with a documented degradation contract — asserted on the
    OOXML, not on python-docx's high-level view."""

    @staticmethod
    def _xml(path) -> str:
        import zipfile

        return zipfile.ZipFile(str(path)).read("word/document.xml").decode()

    def _export(self, markup, tmp_path, *, pending="accept-all"):
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk(markup, author=aim.external("test"))
        out = tmp_path / "out.docx"
        aim.to_docx(doc, out, pending=pending)
        return out

    def _fills(self, path) -> list[str]:
        import re

        return sorted(set(re.findall(r'<w:shd [^>]*w:fill="([0-9A-Fa-f]{6})"', self._xml(path))))

    def _borders(self, path, kind: str) -> dict[str, str]:
        """{side: RRGGBB} for one border container (pBdr / tcBorders / bdr)."""
        import re

        xml = self._xml(path)
        if kind == "bdr":  # run border: one element, no sides
            hits = re.findall(r'<w:bdr [^>]*w:color="([0-9A-Fa-f]{6})"', xml)
            return {"run": hits[0]} if hits else {}
        block = re.search(rf"<w:{kind}>(.*?)</w:{kind}>", xml, re.S)
        if block is None:
            return {}
        return {
            side: colour
            for side, colour in re.findall(
                r'<w:(top|left|bottom|right) [^>]*w:color="([0-9A-Fa-f]{6})"', block.group(1)
            )
        }

    # -- backgrounds ---------------------------------------------------------
    def test_a_block_background_becomes_paragraph_shading(self, tmp_path):
        out = self._export('<p style="background-color:#fff1f7">Tinted</p>', tmp_path)
        assert self._fills(out) == ["FFF1F7"]
        assert "<w:pPr>" in self._xml(out) and "w:shd" in self._xml(out)

    def test_an_inline_background_becomes_run_shading_not_a_highlight(self, tmp_path):
        """Word's highlight is a 16-value enum; shading takes the real RGB."""
        out = self._export('<p>a <span style="background-color:#fff1f7">b</span></p>', tmp_path)
        assert self._fills(out) == ["FFF1F7"]
        assert "w:highlight" not in self._xml(out)

    def test_a_list_items_background_shades_its_paragraph(self, tmp_path):
        out = self._export('<ul><li style="background-color:#fff1f7">Item</li></ul>', tmp_path)
        assert self._fills(out) == ["FFF1F7"]

    def test_a_cell_background_becomes_cell_shading(self, tmp_path):
        out = self._export(
            '<table><tbody><tr><td style="background-color:#fff1f7">c</td></tr></tbody></table>',
            tmp_path,
        )
        assert self._fills(out) == ["FFF1F7"]
        assert "w:tcPr" in self._xml(out)

    def test_a_grouping_background_shades_its_descendant_paragraphs(self, tmp_path):
        """Degradation contract: Word gets no single contiguous box, so every
        emitted descendant whose own background is transparent is shaded."""
        out = self._export(
            '<section style="background-color:#fff1f7"><p>one</p><p>two</p></section>', tmp_path
        )
        import re

        assert self._fills(out) == ["FFF1F7"]
        assert len(re.findall(r'<w:shd [^>]*w:fill="FFF1F7"', self._xml(out))) == 2

    def test_a_descendants_own_background_wins_over_the_group(self, tmp_path):
        out = self._export(
            '<section style="background-color:#fff1f7">'
            '<p style="background-color:#eeeeee">own</p></section>',
            tmp_path,
        )
        assert self._fills(out) == ["EEEEEE"]

    def test_a_descendant_base_background_hides_the_group(self, tmp_path):
        import re

        out = self._export(
            '<table style="background-color:#fff1f7"><thead><tr><th>Head</th></tr></thead>'
            "<tbody><tr><td>Body</td></tr></tbody></table>",
            tmp_path,
        )
        assert self._fills(out) == ["FFF1F7"]
        assert len(re.findall(r'<w:shd [^>]*w:fill="FFF1F7"', self._xml(out))) == 1

    def test_an_unpainted_document_gains_no_shading(self, tmp_path):
        assert self._fills(self._export("<p>Plain</p>", tmp_path)) == []

    # -- borders -------------------------------------------------------------
    def test_colour_without_a_border_emits_nothing(self, tmp_path):
        out = self._export('<p style="border-color:#ff69b4">x</p>', tmp_path)
        assert self._borders(out, "pBdr") == {}

    def test_a_full_border_recolours_every_side(self, tmp_path):
        out = self._export('<p class="border" style="border-color:#ff69b4">x</p>', tmp_path)
        assert self._borders(out, "pBdr") == {
            "top": "FF69B4",
            "left": "FF69B4",
            "bottom": "FF69B4",
            "right": "FF69B4",
        }

    def test_a_top_only_border_recolours_only_the_top(self, tmp_path):
        out = self._export('<p class="border-t" style="border-color:#ff69b4">x</p>', tmp_path)
        assert self._borders(out, "pBdr") == {"top": "FF69B4"}

    def test_a_bottom_only_border_recolours_only_the_bottom(self, tmp_path):
        out = self._export('<p class="border-b" style="border-color:#ff69b4">x</p>', tmp_path)
        assert self._borders(out, "pBdr") == {"bottom": "FF69B4"}

    def test_a_blockquotes_base_stylesheet_border_is_recoloured(self, tmp_path):
        out = self._export('<blockquote style="border-color:#ff69b4">q</blockquote>', tmp_path)
        assert self._borders(out, "pBdr") == {"left": "FF69B4"}

    def test_a_grouping_border_is_carried_by_each_emitted_child_block(self, tmp_path):
        import re

        out = self._export(
            '<blockquote style="border-color:#ff69b4"><p>One</p><p>Two</p></blockquote>',
            tmp_path,
        )
        assert len(re.findall(r'<w:left [^>]*w:color="FF69B4"', self._xml(out))) == 2

    def test_a_class_border_colour_matches_the_browsers_shorthand_reset(self, tmp_path):
        """`.border-t` sorts AFTER `.border-red-600` in the generated
        stylesheet, so its `border-top:1px solid #e5e7eb` shorthand resets the
        colour and a browser renders GREY. DOCX must emit that computed grey,
        not the losing red and not an absent border."""
        out = self._export('<p class="border-t border-red-600">x</p>', tmp_path)
        assert self._borders(out, "pBdr") == {"top": "E5E7EB"}

    def test_a_cell_border_colour_reaches_the_cell(self, tmp_path):
        out = self._export(
            '<table><tbody><tr><td class="border" style="border-color:#ff69b4">c</td>'
            "</tr></tbody></table>",
            tmp_path,
        )
        assert self._borders(out, "tcBorders") == {
            "top": "FF69B4",
            "left": "FF69B4",
            "bottom": "FF69B4",
            "right": "FF69B4",
        }

    def test_an_inline_border_is_a_whole_run_border(self, tmp_path):
        """Documented degradation: `w:rPr/w:bdr` is ONE border for the whole
        run — Word has no per-side border on a run — so a side utility on an
        inline element colours the whole box."""
        out = self._export(
            '<p>a <span class="border-t" style="border-color:#ff69b4">b</span></p>', tmp_path
        )
        assert self._borders(out, "bdr") == {"run": "FF69B4"}

    def test_word_property_children_stay_in_schema_order(self, tmp_path):
        """WordprocessingML property children are a SEQUENCE. python-docx only
        knows the positions of the elements it models, so an appended
        `w:shd`/`w:pBdr` would land after siblings that must follow it and
        make the file invalid."""
        import re

        out = self._export(
            '<p class="border" style="background-color:#fff1f7; border-color:#ff69b4">Boxed '
            '<span class="border-t" style="background-color:#eeeeee; '
            'border-color:#0f766e">run</span></p>',
            tmp_path,
        )
        xml = self._xml(out)
        for tag, want in (("w:pPr", ["w:pBdr", "w:shd"]), ("w:rPr", ["w:bdr", "w:shd"])):
            blocks = re.findall(rf"<{tag}>(.*?)</{tag}>", xml, re.S)
            painted = [[k for k in re.findall(r"<(w:[a-zA-Z]+)", b) if k in want] for b in blocks]
            assert want in painted, f"{tag}: {painted}"

    def test_a_cells_property_order_survives_the_span_structure(self, tmp_path):
        import re

        out = self._export(
            '<table><tbody><tr><td class="border" '
            'style="background-color:#fff1f7; border-color:#ff69b4">c</td></tr></tbody></table>',
            tmp_path,
        )
        block = re.search(r"<w:tcPr>(.*?)</w:tcPr>", self._xml(out), re.S)
        assert block is not None
        want = ("w:tcBorders", "w:shd")
        kids = [k for k in re.findall(r"<(w:[a-zA-Z]+)", block.group(1)) if k in want]
        assert kids == list(want)

    def test_an_unpainted_document_gains_no_borders(self, tmp_path):
        xml = self._xml(self._export('<p class="border">x</p>', tmp_path))
        assert "w:pBdr" not in xml


class TestDocxPaintPendingModes:
    """Paint has to survive the pending lane: tracked revisions carry it in
    the revision run properties, and the resolved exports carry exactly the
    chosen paint."""

    @staticmethod
    def _xml(path) -> str:
        import zipfile

        return zipfile.ZipFile(str(path)).read("word/document.xml").decode()

    def _doc(self):
        doc = aim.from_text("placeholder", title="t")
        doc.add_chunk('<p data-aim="p1">Plain</p>', author=BOT, at=ts(0))
        doc.propose_modify(
            "p1",
            '<p data-aim="p1" class="border" '
            'style="color:#ff69b4; background-color:#fff1f7; border-color:#123456">'
            "Pink</p>",
            author=BOT,
            explanation="Recolour.",
            at=ts(1),
        )
        return doc

    def _out(self, tmp_path, pending):
        out = tmp_path / f"{pending}.docx"
        aim.to_docx(self._doc(), out, pending=pending)
        return out

    def test_tracked_insertions_carry_the_proposed_paint(self, tmp_path):
        """Text and the run-level approximation of the new block box ride the
        inserted revision. Rejecting it must not leave paint on the empty
        paragraph that remains."""
        import re

        xml = self._xml(self._out(tmp_path, "tracked"))
        inserted = next(p for p in re.findall(r"<w:p>.*?</w:p>", xml, re.S) if "<w:ins " in p)
        ins = re.search(r"<w:ins .*?</w:ins>", inserted, re.S)
        assert ins is not None and 'w:val="FF69B4"' in ins.group(0)
        assert 'w:fill="FFF1F7"' in ins.group(0)
        assert '<w:bdr w:val="single"' in ins.group(0)
        assert 'w:color="123456"' in ins.group(0)
        assert 'w:fill="FFF1F7"' not in inserted.split("<w:ins ")[0]
        assert "w:pBdr" not in inserted.split("<w:ins ")[0]

    def test_a_tracked_deletion_keeps_the_paint_it_is_removing(self, tmp_path):
        import re

        xml = self._xml(self._out(tmp_path, "tracked"))
        deleted = next(p for p in re.findall(r"<w:p>.*?</w:p>", xml, re.S) if "<w:del " in p)
        assert "FF69B4" not in deleted and "FFF1F7" not in deleted

    def test_a_tracked_cell_replacement_carries_both_box_colours_in_the_revisions(self, tmp_path):
        import re

        doc = aim.new_document(title="cell paint")
        doc.add_chunk(
            '<table data-aim-container="table"><tbody><tr data-aim="row">'
            '<td style="background-color:#111111">Old</td></tr></tbody></table>',
            author=BOT,
            at=ts(0),
        )
        doc.propose_modify(
            "row",
            '<tr data-aim="row"><td style="background-color:#ff69b4">New</td></tr>',
            author=BOT,
            at=ts(1),
        )
        out = tmp_path / "tracked-cell.docx"
        aim.to_docx(doc, out, pending="tracked")
        xml = self._xml(out)
        cell = re.search(r"<w:tc>.*?</w:tc>", xml, re.S)
        assert cell is not None
        deleted = re.search(r"<w:del .*?</w:del>", cell.group(0), re.S)
        inserted = re.search(r"<w:ins .*?</w:ins>", cell.group(0), re.S)
        assert deleted is not None and 'w:fill="111111"' in deleted.group(0)
        assert inserted is not None and 'w:fill="FF69B4"' in inserted.group(0)
        tc_props = re.search(r"<w:tcPr>.*?</w:tcPr>", cell.group(0), re.S)
        assert tc_props is not None and "w:shd" not in tc_props.group(0)

    def test_a_pending_container_replacement_inherits_from_its_actual_parent(self, tmp_path):
        import re

        doc = aim.new_document(title="pending inheritance")
        doc.add_chunk(
            '<aim-slide data-aim-container="slide" style="color:#ff69b4">'
            '<ul data-aim-container="list"><li data-aim="old">Old</li></ul></aim-slide>',
            author=BOT,
            at=ts(0),
        )
        doc.propose_modify(
            "list",
            '<ul data-aim-container="list"><li data-aim="new">New</li></ul>',
            author=BOT,
            at=ts(1),
        )
        out = tmp_path / "pending-parent.docx"
        aim.to_docx(doc, out, pending="tracked")
        inserted = re.search(r"<w:ins .*?</w:ins>", self._xml(out), re.S)
        assert inserted is not None and 'w:val="FF69B4"' in inserted.group(0)

    def test_accept_all_carries_exactly_the_proposed_paint(self, tmp_path):
        xml = self._xml(self._out(tmp_path, "accept-all"))
        assert 'w:val="FF69B4"' in xml and 'w:color="123456"' in xml
        assert "w:pBdr" in xml and "w:ins" not in xml

    def test_reject_all_carries_no_paint(self, tmp_path):
        xml = self._xml(self._out(tmp_path, "reject-all"))
        assert "FF69B4" not in xml and "FFF1F7" not in xml and "123456" not in xml

    @pytest.mark.parametrize("pending", ["tracked", "accept-all", "reject-all"])
    def test_no_pending_mode_mutates_the_source(self, tmp_path, pending):
        doc = self._doc()
        before = doc.dumps()
        aim.to_docx(doc, tmp_path / f"{pending}.docx", pending=pending)
        assert doc.dumps() == before
