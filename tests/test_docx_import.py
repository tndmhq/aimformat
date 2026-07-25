"""The native DOCX importer (extra ``docx``): styling fidelity, the
rhythm-vs-local-intent doctrine, structure, pagination, theme derivation.

Fixtures here are built with python-docx (a dev/test dependency), same
pattern as the pagination tests: one feature per document, so a failure
names the feature. Whole real-world documents — the combinations no
synthetic fixture produces — live as binaries in ``tests/fixtures/docxs/``
and are asserted in ``test_docx_real_documents.py``.
"""

from __future__ import annotations

import io
import re
import zipfile

import pytest

import aimformat as aim

pytest.importorskip("docx_parser_converter")
docx = pytest.importorskip("docx")

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX  # noqa: E402
from docx.oxml import OxmlElement, parse_xml  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Emu, Inches, Pt, RGBColor  # noqa: E402
from lxml import etree  # noqa: E402

from aimformat.convert import from_docx  # noqa: E402
from aimformat.registry import REGISTRY  # noqa: E402

# a valid 1×1 red PNG
_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
    "53de0000000c4944415408d763f8cfc000000301010018dd8db00000000049"
    "454e44ae426082"
)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _styled_docx() -> io.BytesIO:
    doc = Document()
    doc.add_heading("Heading One Alpha", level=1)
    doc.add_heading("Heading Two Bravo", level=2)

    p = doc.add_paragraph()
    r = p.add_run("GeorgiaRun ")
    r.font.name = "Georgia"
    r.font.size = Pt(18)
    r2 = p.add_run("CourierRun")
    r2.font.name = "Courier New"
    r2.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run("BoldRun").bold = True
    p.add_run(" plain ")
    p.add_run("ItalicRun").italic = True
    p.add_run(" plain ")
    p.add_run("UnderlineRun").underline = True
    p.add_run(" plain ")
    p.add_run("StrikeRun").font.strike = True

    p = doc.add_paragraph()
    rc = p.add_run("RedRun")
    rc.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.add_run(" and ")
    rh = p.add_run("HighlightRun")
    rh.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run(" and E=mc")
    p.add_run("2").font.superscript = True

    pc = doc.add_paragraph("Centered text.")
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pj = doc.add_paragraph("Justified text that is long enough to wrap.")
    pj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for item in ("BulletOne", "BulletTwo"):
        doc.add_paragraph(item, style="List Bullet")
    for item in ("NumberOne", "NumberTwo"):
        doc.add_paragraph(item, style="List Number")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "CellA"
    table.cell(0, 1).text = "CellB"
    table.cell(1, 0).text = "CellC"
    cell = table.cell(1, 1)
    cell.text = ""
    bold_run = cell.paragraphs[0].add_run("BoldCell")
    bold_run.bold = True

    p = doc.add_paragraph("Visit ")
    _add_hyperlink(p, "https://example.com/linktarget", "LinkText")
    p.add_run(" now.")

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("SecondPage.")

    doc.add_picture(io.BytesIO(_PNG), width=Emu(914400))  # 1 inch

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


@pytest.fixture(scope="module")
def imported() -> aim.AimDocument:
    from aimformat.convert._docx_in import convert_docx

    return convert_docx(_styled_docx())


def test_public_entry_point_takes_a_path(tmp_path):
    target = tmp_path / "styled.docx"
    target.write_bytes(_styled_docx().read())
    doc = from_docx(target)
    assert doc.title == "Heading One Alpha"
    assert any(e.author.id == "docx-import" for e in doc.history if e.author)


@pytest.fixture(scope="module")
def html(imported) -> str:
    return "\n".join(c.html for c in imported.chunks)


class TestStyling:
    def test_fonts_and_sizes_become_literal_typography(self, html):
        assert '<span style="font-size:18pt; font-family:Georgia">' in html
        assert '<span style="font-size:9pt; font-family:Courier New">' in html

    def test_color_becomes_literal_paint(self, html):
        assert '<span style="color:#ff0000">RedRun</span>' in html

    def test_highlight_becomes_mark(self, html):
        assert "<mark>HighlightRun</mark>" in html

    def test_classic_marks(self, html):
        assert "<strong>BoldRun</strong>" in html
        assert "<em>ItalicRun</em>" in html
        assert "<u>UnderlineRun</u>" in html
        assert "<s>StrikeRun</s>" in html
        assert "<sup>2</sup>" in html

    def test_alignment_becomes_classes(self, imported):
        tags = {c.html.split(">", 1)[0] for c in imported.chunks}
        assert any('class="text-center"' in t for t in tags)
        assert any('class="text-justify"' in t for t in tags)

    def test_list_item_alignment_becomes_a_class(self):
        # a centered bullet is visible structure, same as a centered heading
        doc = Document()
        p = doc.add_paragraph("centered bullet", style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("plain bullet", style="List Bullet")
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        body = convert_docx(out).dumps()
        assert '<li data-aim="' in body
        assert re.search(r'<li[^>]*class="text-center"[^>]*>centered bullet</li>', body), body[:400]
        assert re.search(r"<li[^>]*>plain bullet</li>", body)

    def test_caps_combines_with_other_run_styling(self):
        # all-caps + colour on one run: one span carrying BOTH the uppercase
        # class and the literal paint — neither silently dropped
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Shouted")
        run.font.all_caps = True
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        lone = p.add_run(" and just caps")
        lone.font.all_caps = True
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        html = "\n".join(c.html for c in convert_docx(out).chunks)
        assert '<span class="uppercase" style="color:#1d4ed8">Shouted</span>' in html
        assert '<span class="uppercase"> and just caps</span>' in html

    def test_style_driven_bold_is_suppressed_on_headings(self, imported):
        h1 = next(c for c in imported.chunks if c.tag == "h1")
        assert h1.html == f'<h1 data-aim="{h1.id}">Heading One Alpha</h1>'

    def test_hyperlink_character_style_is_suppressed(self, html):
        assert '<a href="https://example.com/linktarget">LinkText</a>' in html


class TestStructure:
    def test_heading_levels(self, imported):
        assert [c.tag for c in imported.chunks[:2]] == ["h1", "h2"]

    def test_lists_split_by_numbering(self, imported):
        text = imported.dumps()
        assert "<ul data-aim-container=" in text and "<ol data-aim-container=" in text

    def test_table_with_cell_formatting(self, html):
        # cells may now carry a width style (Card B); assert content + the
        # bold-inside-a-cell survive, robust to any cell geometry
        assert ">CellA</td>" in html
        assert "<strong>BoldCell</strong></td>" in html

    def test_page_break_chunk(self, imported):
        assert any(c.tag == "aim-page-break" for c in imported.chunks)

    def test_image_embeds_as_data_uri(self, html):
        assert 'src="data:image/png;base64,' in html

    def test_page_setup_carried(self, imported):
        text = imported.dumps()
        assert '"size":"Letter"' in text.replace(" ", "") or '"size": "Letter"' in text

    def test_document_title_falls_back_to_first_heading(self, imported):
        assert imported.title == "Heading One Alpha"


class TestConformance:
    def test_lints_clean(self, imported):
        assert [f for f in aim.lint(imported) if f.level == "error"] == []

    def test_roundtrip_is_byte_stable(self, imported):
        text = imported.dumps()
        assert aim.loads(text).dumps() == text

    def test_history_verifies(self, imported):
        assert imported.verify() == []


class TestTheme:
    def test_theme_derives_from_the_source(self, imported):
        text = imported.dumps()
        assert "--aim-font-heading:" in text and "--aim-font-body:" in text
        assert "--aim-brand-1:#" in text

    def test_caller_theme_slots_win(self):
        from aimformat.convert._docx_in import convert_docx

        doc = convert_docx(_styled_docx(), theme={"--aim-font-body": "Test Face"})
        assert "--aim-font-body:Test Face" in doc.dumps()


def test_a_list_starting_indented_keeps_its_outdented_items():
    """ilvl 1 then ilvl 0: nesting starts at the group's minimum level, so
    the outdented item must survive (not be dropped by the walk return)."""
    doc = Document()
    for text, ilvl in (("DeepFirst", 1), ("ShallowSecond", 0)):
        p = doc.add_paragraph(text, style="List Bullet")
        num_pr = OxmlElement("w:numPr")
        lvl = OxmlElement("w:ilvl")
        lvl.set(qn("w:val"), str(ilvl))
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(lvl)
        num_pr.append(num_id)
        p._p.get_or_add_pPr().append(num_pr)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    from aimformat.convert._docx_in import convert_docx

    imported = convert_docx(out)
    text = imported.dumps()
    assert "DeepFirst" in text and "ShallowSecond" in text
    assert [f for f in aim.lint(imported) if f.level == "error"] == []


# --------------------------------------------------------------------------
# Card A: edge-case ports (strict-OOXML, textboxes, OMML, checkboxes, symbols)
# --------------------------------------------------------------------------


from aimformat.convert._docx_in import convert_docx  # noqa: E402
from aimformat.convert._docx_seam import (  # noqa: E402
    _is_safe_zip_member,
    _is_strict_ooxml,
    _strict_ns_to_transitional,
    symbol_char,
)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
_V = "urn:schemas-microsoft-com:vml"
_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_WPG = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"


def _one_para_html(builder) -> str:
    doc = Document()
    builder(doc)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    imported = convert_docx(out)
    assert [f for f in aim.lint(imported) if f.level == "error"] == []
    return "\n".join(c.html for c in imported.chunks)


class TestSymbols:
    def test_wingdings_glyphs_map_to_unicode(self):
        def build(doc):
            p = doc.add_paragraph()
            for char in ("F0FC", "F0B7", "F0E0"):
                run = p.add_run()
                sym = OxmlElement("w:sym")
                sym.set(qn("w:font"), "Wingdings")
                sym.set(qn("w:char"), char)
                run._r.append(sym)

        assert "✓•→" in _one_para_html(build)

    def test_unmapped_wingdings_glyph_drops_never_leaks_the_hex(self):
        # F001 is not in the curated table: it must vanish, not print "F001"
        assert symbol_char("Wingdings", "F001") is None
        assert symbol_char("Wingdings", "F0FC") == "✓"

    def test_non_symbol_font_passes_real_characters_and_drops_pua(self):
        assert symbol_char("Calibri", "2022") == "•"  # real BMP char
        assert symbol_char("SomeFont", "F0FC") is None  # private-use, no table

    def test_a_bad_char_is_dropped(self):
        assert symbol_char("Wingdings", "nothex") is None
        assert symbol_char("Wingdings", None) is None


class TestEquations:
    def test_omml_survives_as_literal_text(self):
        def build(doc):
            p = doc.add_paragraph("Result ")
            p._p.append(parse_xml(f'<m:oMath xmlns:m="{_M}"><m:r><m:t>x=y+1</m:t></m:r></m:oMath>'))

        assert "Result x=y+1" in _one_para_html(build)


class TestCheckbox:
    def test_inline_content_control_checkbox_becomes_a_glyph(self):
        def build(doc):
            p = doc.add_paragraph()
            p._p.append(
                parse_xml(
                    f'<w:sdt xmlns:w="{_W}" xmlns:w14="{_W14}"><w:sdtPr>'
                    '<w14:checkbox><w14:checked w14:val="1"/></w14:checkbox>'
                    "</w:sdtPr><w:sdtContent><w:r><w:t>x</w:t></w:r></w:sdtContent></w:sdt>"
                )
            )

        assert "☑" in _one_para_html(build)

    def test_unchecked_checkbox_is_the_empty_box(self):
        def build(doc):
            p = doc.add_paragraph("Task ")
            p._p.append(
                parse_xml(
                    f'<w:sdt xmlns:w="{_W}" xmlns:w14="{_W14}"><w:sdtPr>'
                    '<w14:checkbox><w14:checked w14:val="0"/></w14:checkbox>'
                    "</w:sdtPr><w:sdtContent><w:r><w:t>x</w:t></w:r></w:sdtContent></w:sdt>"
                )
            )

        html = _one_para_html(build)
        assert "☐" in html and "Task" in html


class TestTextbox:
    def test_textbox_paragraph_follows_its_anchor(self):
        def build(doc):
            doc.add_paragraph("Before")
            anchor = doc.add_paragraph("Anchor")
            anchor._p.append(
                parse_xml(
                    f'<w:txbxContent xmlns:w="{_W}"><w:p><w:r>'
                    "<w:t>TextboxLine</w:t></w:r></w:p></w:txbxContent>"
                )
            )
            doc.add_paragraph("After")

        doc = Document()
        build(doc)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        imported = convert_docx(out)
        texts = [c.text for c in imported.chunks]
        # the textbox line sits between its anchor and the following paragraph
        assert texts == ["Before", "Anchor", "TextboxLine", "After"]

    def test_alternate_content_textbox_emits_once(self):
        # Word wraps every inserted shape in mc:AlternateContent with the SAME
        # w:txbxContent in both the DrawingML Choice and the VML Fallback —
        # MCE says read exactly one branch, so the text must appear once.
        mc = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        wps = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
        doc = Document()
        anchor = doc.add_paragraph("Anchor")
        run = parse_xml(f'<w:r xmlns:w="{_W}"/>')
        run.append(
            parse_xml(
                f'<mc:AlternateContent xmlns:mc="{mc}" xmlns:w="{_W}" '
                f'xmlns:wps="{wps}" xmlns:v="urn:schemas-microsoft-com:vml">'
                '<mc:Choice Requires="wps"><wps:wsp><wps:txbx>'
                "<w:txbxContent><w:p><w:r><w:t>BoxLine</w:t></w:r></w:p>"
                "</w:txbxContent></wps:txbx></wps:wsp></mc:Choice>"
                "<mc:Fallback><v:shape><v:textbox>"
                "<w:txbxContent><w:p><w:r><w:t>BoxLine</w:t></w:r></w:p>"
                "</w:txbxContent></v:textbox></v:shape></mc:Fallback>"
                "</mc:AlternateContent>"
            )
        )
        anchor._p.append(run)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        imported = convert_docx(out)
        texts = [c.text for c in imported.chunks]
        assert texts == ["Anchor", "BoxLine"]

    # An image inside a textbox has two failure modes that pull in OPPOSITE
    # directions, so each flavour must be pinned or a fix for one silently
    # breaks the other — which is exactly what happened twice here:
    #
    #   * dpc's typed model carries DrawingML pictures, so recovering blips
    #     from inside a textbox emits them TWICE;
    #   * dpc has no VML model at all, so NOT recovering there loses the
    #     image outright.
    #
    # These two tests are a pair. Do not "simplify" either away.

    @staticmethod
    def _textbox_with_image(flavour: str) -> str:
        doc = Document()
        anchor = doc.add_paragraph("Anchor")
        rid, _ = doc.part.get_or_add_image(io.BytesIO(_PNG))
        if flavour == "vml":
            inner = (
                f'<w:pict xmlns:w="{_W}" xmlns:v="{_V}" xmlns:r="{_R}">'
                '<v:shape style="width:60pt;height:60pt">'
                f'<v:imagedata r:id="{rid}"/></v:shape></w:pict>'
            )
        else:
            scratch = Document()
            run = scratch.add_paragraph().add_run()
            run.add_picture(io.BytesIO(_PNG), width=Inches(1))
            drawing = etree.tostring(run._r.find(f"{{{_W}}}drawing")).decode()
            inner = re.sub(r'r:embed="[^"]+"', f'r:embed="{rid}"', drawing)
            if flavour == "drawingml-mce":
                # how Word actually writes anything richer than a plain inline
                # picture: grouped art, picture fills, SmartArt
                inner = (
                    f'<mc:AlternateContent xmlns:mc="{_MC}" xmlns:w="{_W}" '
                    f'xmlns:wps="{_WPS}" xmlns:v="{_V}" xmlns:r="{_R}">'
                    f'<mc:Choice Requires="wps">{inner}</mc:Choice>'
                    f'<mc:Fallback><v:shape><v:imagedata r:id="{rid}"/>'
                    "</v:shape></mc:Fallback></mc:AlternateContent>"
                )
        run = parse_xml(f'<w:r xmlns:w="{_W}"/>')
        run.append(
            parse_xml(
                f'<w:pict xmlns:w="{_W}" xmlns:v="{_V}" xmlns:r="{_R}">'
                "<v:shape><v:textbox><w:txbxContent><w:p>"
                "<w:r><w:t>Caption</w:t></w:r>"
                f"<w:r>{inner}</w:r>"
                "</w:p></w:txbxContent></v:textbox></v:shape></w:pict>"
            )
        )
        anchor._p.append(run)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        imported = convert_docx(out)
        assert "Caption" in [c.text for c in imported.chunks]
        return "\n".join(c.html for c in imported.chunks)

    def test_vml_image_in_a_textbox_is_recovered_once(self):
        # dpc cannot see VML; without the seam's recovery this image is LOST
        assert self._textbox_with_image("vml").count("<img") == 1

    def test_drawingml_image_in_a_textbox_is_not_doubled(self):
        # dpc's own walk already emits this one, so the seam must not re-add it
        assert self._textbox_with_image("drawingml").count("<img") == 1

    def test_an_mce_wrapped_image_in_a_textbox_is_recovered(self):
        # the third case, and the one that makes "skip blips inside textboxes"
        # too blunt: dpc's run parser has no mc:AlternateContent branch, so it
        # never sees this picture and the seam is the only thing that can
        # recover it. Everything Word draws beyond a plain inline image —
        # grouped artwork, shapes with a picture fill, SmartArt — lands here.
        assert self._textbox_with_image("drawingml-mce").count("<img") == 1


class TestArchiveGuards:
    """Zip-slip / zip-bomb rejection on EVERY input, not only Strict OOXML."""

    @staticmethod
    def _plain_docx() -> io.BytesIO:
        doc = Document()
        doc.add_paragraph("ok")
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out

    def test_zip_slip_member_rejected(self):
        import zipfile

        src = zipfile.ZipFile(self._plain_docx())
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as z:
            for info in src.infolist():
                z.writestr(info, src.read(info.filename))
            z.writestr("../evil.txt", b"x")
        out.seek(0)
        with pytest.raises(ValueError, match="zip-slip"):
            convert_docx(out)

    def test_oversized_member_rejected(self, monkeypatch):
        from aimformat.convert import _docx_seam

        monkeypatch.setattr(_docx_seam, "_MAX_MEMBER_BYTES", 64)
        with pytest.raises(ValueError, match="oversized"):
            convert_docx(self._plain_docx())

    def test_total_size_cap_rejected(self, monkeypatch):
        from aimformat.convert import _docx_seam

        monkeypatch.setattr(_docx_seam, "_MAX_TOTAL_BYTES", 256)
        with pytest.raises(ValueError, match="size limit"):
            convert_docx(self._plain_docx())


class TestStrictOoxml:
    @staticmethod
    def _to_strict(transitional: io.BytesIO) -> io.BytesIO:
        import zipfile

        repl = [
            (f"{_W}", "http://purl.oclc.org/ooxml/wordprocessingml/main"),
            (
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "http://purl.oclc.org/ooxml/officeDocument/relationships",
            ),
        ]
        src = zipfile.ZipFile(transitional)
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as z:
            for info in src.infolist():
                data = src.read(info.filename)
                if info.filename.endswith((".xml", ".rels")):
                    text = data.decode()
                    for a, b in repl:
                        text = text.replace(a, b)
                    data = text.encode()
                z.writestr(info, data)
        out.seek(0)
        return out

    def test_strict_package_parses_after_normalization(self):
        doc = Document()
        doc.add_heading("StrictTitle", level=1)
        doc.add_paragraph("Strict body text.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        strict = self._to_strict(buf)

        import zipfile

        assert _is_strict_ooxml(zipfile.ZipFile(io.BytesIO(strict.getvalue())))
        imported = convert_docx(io.BytesIO(strict.getvalue()))
        text = imported.dumps()
        assert "StrictTitle" in text and "Strict body text." in text
        assert [f for f in aim.lint(imported) if f.level == "error"] == []

    def test_namespace_mapping_reverses_transitional(self):
        assert _strict_ns_to_transitional("http://purl.oclc.org/ooxml/wordprocessingml/main") == _W

    def test_zip_slip_members_are_rejected(self):
        assert not _is_safe_zip_member("../evil.xml")
        assert not _is_safe_zip_member("/etc/passwd")
        assert not _is_safe_zip_member("C:/windows")
        assert _is_safe_zip_member("word/document.xml")


# --------------------------------------------------------------------------
# Card B: table styling (cell shading + width; borders deliberately skipped)
# --------------------------------------------------------------------------


class TestTableStyling:
    @staticmethod
    def _shaded_table_html() -> str:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.cell(0, 0)
        cell.text = "Shaded"
        pr = cell._tc.get_or_add_tcPr()
        pr.append(parse_xml(f'<w:shd xmlns:w="{_W}" w:val="clear" w:fill="D9E2F3"/>'))
        pr.append(parse_xml(f'<w:tcW xmlns:w="{_W}" w:w="3000" w:type="dxa"/>'))
        table.cell(0, 1).text = "Plain"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        imported = convert_docx(out)
        assert [f for f in aim.lint(imported) if f.level == "error"] == []
        return "\n".join(c.html for c in imported.chunks)

    def test_cell_shading_becomes_background_paint(self):
        assert "background-color:#d9e2f3" in self._shaded_table_html()

    def test_cell_width_becomes_px_geometry(self):
        # python-docx redistributes the table width across cells, so the exact
        # px is its call; the conversion itself is pinned by the unit test.
        assert re.search(r"width:\d+px", self._shaded_table_html())

    def test_width_precedes_background_in_canonical_order(self):
        html = self._shaded_table_html()
        assert re.search(r'style="width:\d+px; background-color:#d9e2f3"', html)

    def test_dxa_conversion_and_non_dxa_skip(self):
        from aimformat.convert._docx_in import _cell_width_px

        assert _cell_width_px({"type": "dxa", "w": 1500}) == 100  # 1500 / 15
        assert _cell_width_px({"type": "dxa", "w": 3000}) == 200
        assert _cell_width_px({"type": "pct", "w": 5000}) is None
        assert _cell_width_px({"type": "auto"}) is None
        assert _cell_width_px(None) is None


class TestNumberingVocabularyIsValid:
    """The importer's own output must lint. Twice now a v0.5 construct was
    emitted before the registry admitted it — and since the editor rejects
    non-conforming uploads, that turns a real contract into a 422 rather
    than anything visible in a test of the markup alone."""

    def test_a_numbering_prefix_lints_on_the_blocks_that_carry_it(self):
        doc = aim.new_document(title="Prefix")
        for tag in ("p", "h1", "h2", "h6"):
            doc.add_chunk(
                f'<{tag} class="num-1" data-aim-num-prefix="Article ">Scope</{tag}>',
                author=aim.external("t"),
            )
        assert [f for f in aim.lint(doc) if f.level == "error"] == []

    def test_a_list_start_lints(self):
        doc = aim.new_document(title="Start")
        doc.add_chunk(
            '<ol class="list-multilevel" start="5"><li>five</li></ol>',
            author=aim.external("t"),
        )
        assert [f for f in aim.lint(doc) if f.level == "error"] == []

    def test_every_numbering_class_the_importer_emits_is_admitted(self):
        # the whole vocabulary at once, so adding a class without registering
        # its placement fails here rather than on a customer's upload
        doc = aim.new_document(title="All")
        for level in range(1, REGISTRY.num_levels + 1):
            doc.add_chunk(f'<p class="num-{level}">x</p>', author=aim.external("t"))
        doc.add_chunk('<p class="num-2 num-restart">x</p>', author=aim.external("t"))
        assert [f for f in aim.lint(doc) if f.level == "error"] == []

    @pytest.mark.parametrize("declared", ["0.4", "0.3", "0.1"])
    def test_a_05_class_is_gated_against_an_older_declaration(self, declared):
        # the gate named two eras explicitly, so every era added after them
        # went unchecked: a num-3 in a 0.4 document linted clean
        doc = aim.new_document(title="Gate")
        doc.add_chunk('<p class="num-3">x</p>', author=aim.external("t"))
        body = doc.dumps().replace('data-aim-version="0.5"', f'data-aim-version="{declared}"')
        codes = {f.code for f in aim.lint(aim.loads(body)) if f.level == "error"}
        assert "S034" in codes, f"a 0.5 class went unchecked under {declared}"

    def test_the_v05_attributes_are_gated_against_an_older_declaration(self):
        # The era gate reads classes and style props but never attributes, so
        # a 0.4 document carrying v0.5 numbering ATTRIBUTES linted clean —
        # and a writer recorded no floor for them either, so a document could
        # be written claiming an era that cannot render it.
        for markup in (
            '<p data-aim-num-prefix="Article ">x</p>',
            '<ol start="5"><li>five</li></ol>',
        ):
            doc = aim.new_document(title="Gate")
            doc.add_chunk(markup, author=aim.external("t"))
            body = doc.dumps().replace('data-aim-version="0.5"', 'data-aim-version="0.4"')
            codes = {f.code for f in aim.lint(aim.loads(body)) if f.level == "error"}
            assert "S034" in codes, f"{markup} went unchecked in a 0.4 document"


class TestNumberedSchemesImportAsOneShape:
    """A numbering scheme is one thing. Emitting part of it as a list and
    part as blocks leaves the blocks counting against a level nothing
    increments — they render 0.1, 0.2."""

    @staticmethod
    def _multilevel_docx(texts_and_levels) -> io.BytesIO:
        def lvl(i: int, text: str) -> str:
            return (
                f'<w:lvl w:ilvl="{i}"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
                f'<w:lvlText w:val="{text}"/></w:lvl>'
            )

        numbering = (
            f'<w:numbering xmlns:w="{_W}"><w:abstractNum w:abstractNumId="30">'
            + lvl(0, "%1.")
            + lvl(1, "%1.%2.")
            + lvl(2, "%1.%2.%3.")
            + '</w:abstractNum><w:num w:numId="30">'
            '<w:abstractNumId w:val="30"/></w:num></w:numbering>'
        )
        doc = Document()
        for text, ilvl in texts_and_levels:
            para = doc.add_paragraph(text)
            para._p.get_or_add_pPr().append(
                parse_xml(
                    f'<w:numPr xmlns:w="{_W}"><w:ilvl w:val="{ilvl}"/>'
                    '<w:numId w:val="30"/></w:numPr>'
                )
            )
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        source = zipfile.ZipFile(buf)
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as z:
            for name in source.namelist():
                z.writestr(
                    name,
                    numbering.encode() if name == "word/numbering.xml" else source.read(name),
                )
        out.seek(0)
        return out

    def test_word_stock_multilevel_list_emits_one_shape(self):
        imported = convert_docx(
            self._multilevel_docx(
                [("First", 0), ("Sub a", 1), ("Sub b", 1), ("Second", 0), ("Sub c", 1)]
            )
        )
        content = "\n".join(c.html for c in imported.chunks)
        assert "<ol" not in content, "the scheme was torn into list fragments"
        levels = re.findall(r'class="num-(\d)"', content)
        assert levels == ["1", "2", "2", "1", "2"], levels

    def test_the_rendered_numbers_match_what_word_draws(self):
        imported = convert_docx(
            self._multilevel_docx(
                [("First", 0), ("Sub a", 1), ("Sub b", 1), ("Second", 0), ("Sub c", 1)]
            )
        )
        counters = [0] * 10
        rendered = []
        for chunk in imported.chunks:
            match = re.search(r'class="num-(\d)"', chunk.html)
            if not match:
                continue
            level = int(match.group(1))
            counters[level] += 1
            for deeper in range(level + 1, 10):
                counters[deeper] = 0
            rendered.append(".".join(str(counters[i]) for i in range(1, level + 1)))
        assert rendered == ["1", "1.1", "1.2", "2", "2.1"], rendered

    def test_a_flat_numbered_list_is_still_a_list(self):
        imported = convert_docx(self._multilevel_docx([("One", 0), ("Two", 0), ("Three", 0)]))
        body = imported.dumps()
        assert "<ol" in body, "a plain numbered list lost its <ol>"
        # search the CONTENT, not the document: the embedded stylesheet
        # carries a .num-1 rule, so a whole-file search always matches
        assert not any("num-" in c.html for c in imported.chunks)


class TestTableStyleResolution:
    """A table's whole appearance usually lives in its STYLE, in conditional
    blocks (``w:tblStylePr``) that each apply to one band. Reading those
    wrongly is not a cosmetic miss — it repaints the entire table."""

    @staticmethod
    def _styled(*style_xml: str, rows: int = 4) -> str:
        """Import a table using custom styles injected into styles.xml. The
        table uses the LAST one, so a basedOn parent can be passed first."""
        doc = Document()
        table = doc.add_table(rows=rows, cols=2)
        for r in range(rows):
            for c in range(2):
                table.cell(r, c).text = f"r{r}c{c}"
        table._tbl.tblPr.append(parse_xml(f'<w:tblStyle xmlns:w="{_W}" w:val="Probe"/>'))
        # tblLook: firstRow + banded rows on, so the conditional blocks apply
        table._tbl.tblPr.append(
            parse_xml(
                f'<w:tblLook xmlns:w="{_W}" w:val="04A0" w:firstRow="1" '
                'w:lastRow="0" w:firstColumn="0" w:lastColumn="0" '
                'w:noHBand="0" w:noVBand="1"/>'
            )
        )
        for one in style_xml:
            doc.styles.element.append(parse_xml(one))
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        imported = convert_docx(out)
        assert [f for f in aim.lint(imported) if f.level == "error"] == []
        return "\n".join(c.html for c in imported.chunks if c.tag == "tr")

    def test_a_first_row_band_does_not_paint_every_row(self):
        # The wholeTable look is what the style declares DIRECTLY. Reading it
        # with a `.//` subtree search hoists the firstRow band onto every row
        # — a table of dark blue rows with white text on white paper.
        html = self._styled(
            f'<w:style xmlns:w="{_W}" w:type="table" w:styleId="Probe">'
            '<w:name w:val="Probe"/><w:tblStylePr w:type="firstRow">'
            '<w:tcPr><w:shd w:val="clear" w:fill="203864"/></w:tcPr>'
            '<w:rPr><w:color w:val="FFFFFF"/></w:rPr>'
            "</w:tblStylePr></w:style>"
        )
        rows = html.split("\n")
        assert "#203864" in rows[0], "the header band itself was lost"
        assert not any("#203864" in r for r in rows[1:]), (
            "the firstRow band leaked onto the body rows"
        )

    def test_a_band_that_sets_only_text_colour_is_dropped(self):
        # A conditional format with a colour and no fill is meaningless on
        # its own: the fill it was designed against is not there, so the
        # colour paints (often white) text onto white paper.
        html = self._styled(
            f'<w:style xmlns:w="{_W}" w:type="table" w:styleId="Probe">'
            '<w:name w:val="Probe"/><w:tblStylePr w:type="firstRow">'
            '<w:rPr><w:color w:val="FFFFFF"/></w:rPr>'
            "</w:tblStylePr></w:style>"
        )
        assert "color:#ffffff" not in html

    def test_an_inherited_fill_keeps_its_childs_text_colour(self):
        # The colour-without-fill rule must run AFTER basedOn inheritance. A
        # child style that only recolours the header text is completing its
        # parent's dark fill; dropping the colour per-style leaves dark text
        # on a dark band — the exact unreadable header the rule exists to
        # prevent, caused by the rule itself.
        html = self._styled(
            f'<w:style xmlns:w="{_W}" w:type="table" w:styleId="Base">'
            '<w:name w:val="Base"/><w:tblStylePr w:type="firstRow">'
            '<w:tcPr><w:shd w:val="clear" w:fill="1f4e79"/></w:tcPr>'
            "</w:tblStylePr></w:style>",
            f'<w:style xmlns:w="{_W}" w:type="table" w:styleId="Probe">'
            '<w:name w:val="Probe"/><w:basedOn w:val="Base"/>'
            '<w:tblStylePr w:type="firstRow">'
            '<w:rPr><w:color w:val="FFFFFF"/></w:rPr>'
            "</w:tblStylePr></w:style>",
        )
        header = html.split("\n")[0]
        assert "#1f4e79" in header, "the inherited fill was lost"
        assert "color:#ffffff" in header, "the child's text colour was dropped"

    def test_theme_named_fills_and_colours_resolve(self):
        # Word's own table styles name colours through the theme far more
        # often than they spell out hex. Literal-only reading left most real
        # tables unstyled. accent2 is C0504D in the default theme.
        html = self._styled(
            f'<w:style xmlns:w="{_W}" w:type="table" w:styleId="Probe">'
            '<w:name w:val="Probe"/><w:tblStylePr w:type="firstRow">'
            '<w:tcPr><w:shd w:val="clear" w:themeFill="accent2"/></w:tcPr>'
            '<w:rPr><w:color w:themeColor="background1"/></w:rPr>'
            "</w:tblStylePr></w:style>"
        )
        header = html.split("\n")[0]
        assert "#c0504d" in header, header
        assert "#ffffff" in header, header


# --------------------------------------------------------------------------
# Card C: to_docx export symmetry (DOCX → aim → DOCX round-trip idempotency)
# --------------------------------------------------------------------------


class TestGroupedPictureSizing:
    """A picture inside a group is authored in the GROUP's coordinate space.
    Read literally it renders at its full authored size — a 1.5-inch logo
    arriving 600px wide — so the scaling is what makes grouped artwork
    usable at all. Groups nest, and VML expresses the same idea in a
    completely different way; both are covered here because the real
    fixture only exercises one flat DrawingML group."""

    @staticmethod
    def _drawingml(depth: int) -> int | None:
        """A picture 100000 EMU wide, wrapped in *depth* nested groups. Each
        group halves the coordinate space, so the picture must come out at
        100000 / 2**depth EMU regardless of how deep it sits."""
        from aimformat.convert._docx_seam import _picture_width_px

        a, pic_ns, wpg = _A, _PIC, _WPG
        inner = (
            f'<pic:pic xmlns:pic="{pic_ns}" xmlns:a="{a}"><pic:spPr><a:xfrm>'
            '<a:ext cx="100000" cy="100000"/></a:xfrm></pic:spPr>'
            '<pic:blipFill><a:blip r:embed="rId1" '
            f'xmlns:r="{_R}"/></pic:blipFill></pic:pic>'
        )
        for level in range(depth):
            # ext is HALF of chExt at every level: each wrapper halves again
            child_space = 100000 * (2**level)
            own_ext = child_space // 2
            inner = (
                f'<wpg:wgp xmlns:wpg="{wpg}" xmlns:a="{a}"><wpg:grpSpPr><a:xfrm>'
                f'<a:ext cx="{own_ext}" cy="{own_ext}"/>'
                f'<a:chExt cx="{child_space}" cy="{child_space}"/>'
                f"</a:xfrm></wpg:grpSpPr>{inner}</wpg:wgp>"
            )
        root = etree.fromstring(f'<w:drawing xmlns:w="{_W}">{inner}</w:drawing>')
        blip = root.find(f".//{{{a}}}blip")
        return _picture_width_px(blip, is_vml=False)

    def test_a_single_group_scales_into_its_coordinate_space(self):
        # 100000 EMU authored, one group halving it → 50000 EMU ≈ 5px
        assert self._drawingml(depth=1) == round(50000 / 9525)

    def test_nested_groups_accumulate_every_ancestors_scale(self):
        # two groups, each halving → 25000 EMU. Stopping at the FIRST group
        # ancestor gives 50000 (5px) — twice the size Word draws.
        assert self._drawingml(depth=2) == round(25000 / 9525)
        assert self._drawingml(depth=3) == max(1, round(12500 / 9525))

    @staticmethod
    def _vml(child_style: str, group_style: str = "width:150pt", coordsize: str = "3000,3000"):
        from aimformat.convert._docx_seam import _picture_width_px

        root = etree.fromstring(
            f'<w:pict xmlns:w="{_W}" xmlns:v="{_V}" xmlns:r="{_R}">'
            f'<v:group style="{group_style}" coordsize="{coordsize}">'
            f'<v:shape style="{child_style}"><v:imagedata r:id="rId1"/></v:shape>'
            "</v:group></w:pict>"
        )
        return _picture_width_px(root.find(f".//{{{_V}}}imagedata"), is_vml=True)

    def test_a_vml_group_child_scales_by_the_groups_coordsize(self):
        # the child's "600" is in the group's 3000-unit space, and the group
        # is 150pt (200px) wide → 200 * 600/3000 = 40px. Reading the bare
        # number as a measurement, or falling through to the group's own
        # width, both render this child at the whole group's size.
        assert self._vml("position:absolute;width:600;height:600") == 40

    def test_nested_vml_groups_accumulate_every_coordinate_space(self):
        # Groups nest, and an inner group states its own width in its
        # PARENT's units — bare numbers either way. Reading an intermediate
        # width as points multiplies the error by that group's whole
        # coordinate space: this shape came out at 1000px instead of 50.
        from aimformat.convert._docx_seam import _picture_width_px

        root = etree.fromstring(
            f'<w:pict xmlns:w="{_W}" xmlns:v="{_V}" xmlns:r="{_R}">'
            '<v:group style="width:150pt;height:150pt" coordsize="3000,3000">'
            '<v:group style="position:absolute;width:1500;height:1500" '
            'coordsize="1500,1500">'
            '<v:shape style="position:absolute;width:750;height:750">'
            '<v:imagedata r:id="rId1"/></v:shape>'
            "</v:group></v:group></w:pict>"
        )
        # 150pt = 200px; the inner group is 1500/3000 of it = 100px; the
        # shape is 750/1500 of that = 50px
        node = root.find(f".//{{{_V}}}imagedata")
        assert _picture_width_px(node, is_vml=True) == 50

    def test_a_lone_vml_shape_uses_its_own_measurement(self):
        from aimformat.convert._docx_seam import _picture_width_px

        root = etree.fromstring(
            f'<w:pict xmlns:w="{_W}" xmlns:v="{_V}" xmlns:r="{_R}">'
            '<v:shape style="width:75pt;height:75pt"><v:imagedata r:id="rId1"/>'
            "</v:shape></w:pict>"
        )
        assert _picture_width_px(root.find(f".//{{{_V}}}imagedata"), is_vml=True) == 100


class TestImageParagraphs:
    def test_standalone_image_becomes_a_figure(self):
        # the system idiom is <figure> (from_docling, the editor's atomic
        # nodes, to_docx's figure exporter) — a paragraph that is only an
        # image must not stay a bare <p><img></p>
        from PIL import Image as PILImage

        doc = Document()
        img = io.BytesIO()
        PILImage.new("RGB", (12, 12), (10, 120, 40)).save(img, "PNG")
        img.seek(0)
        doc.add_picture(img)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        body = convert_docx(out).dumps()
        assert re.search(r"<figure[^>]*><img[^>]*data:image/png[^>]*></figure>", body), body[:400]

    def test_figure_roundtrips_through_export(self, tmp_path):
        from PIL import Image as PILImage

        doc = Document()
        img = io.BytesIO()
        PILImage.new("RGB", (12, 12), (10, 120, 40)).save(img, "PNG")
        img.seek(0)
        doc.add_picture(img)
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        a = convert_docx(out)
        path = tmp_path / "img.docx"
        aim.to_docx(a, str(path))
        assert "data:image/" in convert_docx(str(path)).dumps(), "image lost on export"

    def test_centered_image_keeps_alignment_both_ways(self, tmp_path):
        # a centered logo is the classic Word idiom: the figure carries the
        # class on import, and the exporter aligns the picture paragraph
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from PIL import Image as PILImage

        doc = Document()
        img = io.BytesIO()
        PILImage.new("RGB", (12, 12), (10, 120, 40)).save(img, "PNG")
        img.seek(0)
        doc.add_picture(img)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        a = convert_docx(out)
        assert re.search(r'<figure[^>]*class="text-center"[^>]*><img', a.dumps())
        path = tmp_path / "centered.docx"
        aim.to_docx(a, str(path))
        d = Document(str(path))
        pic_para = next(p for p in d.paragraphs if p._p.findall(".//" + qn("w:drawing")))
        assert pic_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        # and the re-import keeps the class
        assert re.search(r'<figure[^>]*class="text-center"', convert_docx(str(path)).dumps())


class TestMergedCells:
    def test_vertical_merge_becomes_rowspan_alongside_gridspan(self):
        # python-docx merge(): row 0 = [gridSpan-2 "wide", vMerge-restart
        # "tall"], row 1 = [x, y, vMerge-continue]. The restart cell must
        # survive with rowspan=2 (dpc models w:vMerge as a plain string —
        # reading .val off it silently dropped the whole merged column).
        doc = Document()
        t = doc.add_table(rows=2, cols=3)
        t.cell(0, 0).merge(t.cell(0, 1))
        t.cell(0, 2).merge(t.cell(1, 2))
        t.cell(0, 0).text = "wide"
        t.cell(0, 2).text = "tall"
        t.cell(1, 0).text = "x"
        t.cell(1, 1).text = "y"
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        dumped = convert_docx(out).dumps()
        table = re.search(r"<table.*?</table>", dumped, re.S)
        assert table, dumped[:600]
        html = table.group(0)
        assert 'colspan="2"' in html and ">wide<" in html
        assert 'rowspan="2"' in html and ">tall<" in html
        # the continuation slot collapses into the restart cell
        assert html.count("<td") + html.count("<th") == 4


class TestExportSymmetry:
    def _roundtrip(self, tmp_path):
        aim_doc = convert_docx(_styled_docx())
        out = tmp_path / "roundtrip.docx"
        aim.to_docx(aim_doc, str(out))
        return Document(str(out))

    def test_font_size_and_family_survive(self, tmp_path):
        d = self._roundtrip(tmp_path)
        faces = {
            (r.font.name, r.font.size.pt if r.font.size else None)
            for p in d.paragraphs
            for r in p.runs
            if r.text.strip()
        }
        assert ("Georgia", 18.0) in faces
        assert ("Courier New", 9.0) in faces

    def test_alignment_survives(self, tmp_path):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        got = {p.alignment for p in self._roundtrip(tmp_path).paragraphs if p.alignment is not None}
        assert WD_ALIGN_PARAGRAPH.CENTER in got
        assert WD_ALIGN_PARAGRAPH.JUSTIFY in got

    def test_theme_fonts_reach_the_styles(self, tmp_path):
        d = self._roundtrip(tmp_path)
        # the source theme (Cambria body / Calibri headings) rides the styles
        assert d.styles["Normal"].font.name == "Cambria"
        assert d.styles["Heading 1"].font.name == "Calibri"

    def test_type_scale_class_exports_to_points(self, tmp_path):
        # text-2xl resolves through the normative pt table to 18pt
        doc = aim.new_document(title="Scale")
        doc.add_chunk('<p class="text-2xl">Big</p>', author=aim.external("t"))
        out = tmp_path / "scale.docx"
        aim.to_docx(doc, str(out))
        d = Document(str(out))
        sizes = [
            r.font.size.pt for p in d.paragraphs for r in p.runs if r.text == "Big" and r.font.size
        ]
        assert sizes == [18.0]

    def test_list_item_alignment_and_size_survive_export(self, tmp_path):
        # the synthetic li the exporter rebuilds must keep class/style — a
        # centered, sized list item exports with alignment and run size
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = aim.new_document(title="Li")
        doc.add_chunk(
            '<ul><li class="text-center" style="font-size:14pt">Item</li></ul>',
            author=aim.external("t"),
        )
        out = tmp_path / "li.docx"
        aim.to_docx(doc, str(out))
        d = Document(str(out))
        para = next(p for p in d.paragraphs if "Item" in p.text)
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert [r.font.size.pt for r in para.runs if r.text == "Item"] == [14.0]

    def test_uppercase_class_exports_as_all_caps(self, tmp_path):
        doc = aim.new_document(title="Caps")
        doc.add_chunk('<p><span class="uppercase">shout</span></p>', author=aim.external("t"))
        out = tmp_path / "caps.docx"
        aim.to_docx(doc, str(out))
        d = Document(str(out))
        flags = [r.font.all_caps for p in d.paragraphs for r in p.runs if r.text == "shout"]
        assert flags == [True]

    def test_font_stack_exports_its_first_family(self, tmp_path):
        # the inline grammar allows a stack; Word run props name one face
        doc = aim.new_document(title="Stack")
        doc.add_chunk(
            "<p><span style=\"font-family:'Segoe UI', Arial, sans-serif\">S</span></p>",
            author=aim.external("t"),
        )
        out = tmp_path / "stack.docx"
        aim.to_docx(doc, str(out))
        d = Document(str(out))
        names = [r.font.name for p in d.paragraphs for r in p.runs if r.text == "S"]
        assert names == ["Segoe UI"]

    def test_inline_typography_beats_the_class_on_the_same_run(self, tmp_path):
        # inline font-size wins over a type-scale class (CSS specificity)
        doc = aim.new_document(title="Override")
        doc.add_chunk(
            '<p><span class="text-2xl" style="font-size:30pt">X</span></p>',
            author=aim.external("t"),
        )
        out = tmp_path / "override.docx"
        aim.to_docx(doc, str(out))
        d = Document(str(out))
        sizes = [
            r.font.size.pt for p in d.paragraphs for r in p.runs if r.text == "X" and r.font.size
        ]
        assert sizes == [30.0]


class TestNumberingStaysInSyncAcrossTheDocument:
    """One numbering scheme is one continuous sequence, wherever its
    paragraphs sit and whatever shape the vocabulary can draw.

    Every failure these pin is silent: the document renders, nothing raises,
    and the numbers are simply wrong from some point onward.
    """

    @staticmethod
    def _with_numbering(doc, numbering: str) -> io.BytesIO:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        source = zipfile.ZipFile(buf)
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as z:
            for name in source.namelist():
                z.writestr(
                    name,
                    numbering.encode() if name == "word/numbering.xml" else source.read(name),
                )
        out.seek(0)
        return out

    @staticmethod
    def _lvl(i: int, text: str, fmt: str = "decimal") -> str:
        return (
            f'<w:lvl w:ilvl="{i}"><w:start w:val="1"/><w:numFmt w:val="{fmt}"/>'
            f'<w:lvlText w:val="{text}"/></w:lvl>'
        )

    @classmethod
    def _chain(cls, num_ids=(30,), abstract: int = 30) -> str:
        instances = "".join(
            f'<w:num w:numId="{n}"><w:abstractNumId w:val="{abstract}"/></w:num>' for n in num_ids
        )
        return (
            f'<w:numbering xmlns:w="{_W}"><w:abstractNum w:abstractNumId="{abstract}">'
            + cls._lvl(0, "%1.")
            + cls._lvl(1, "%1.%2")
            + cls._lvl(2, "%1.%2.%3")
            + f"</w:abstractNum>{instances}</w:numbering>"
        )

    @staticmethod
    def _number(para, num_id: int, ilvl: int) -> None:
        para._p.get_or_add_pPr().append(
            parse_xml(
                f'<w:numPr xmlns:w="{_W}"><w:ilvl w:val="{ilvl}"/>'
                f'<w:numId w:val="{num_id}"/></w:numPr>'
            )
        )

    def test_a_numbered_paragraph_in_a_table_cell_keeps_its_number(self):
        # Word numbers paragraphs inside table cells like any other. Skipping
        # them loses the number AND leaves the counter unadvanced, so every
        # clause after the table is off by one — the damage outlives the table.
        doc = Document()
        self._number(doc.add_paragraph("Top"), 30, 0)
        self._number(doc.add_paragraph("First clause"), 30, 1)
        table = doc.add_table(rows=1, cols=1)
        cell_para = table.rows[0].cells[0].paragraphs[0]
        cell_para.add_run("Clause in a cell")
        self._number(cell_para, 30, 1)
        self._number(doc.add_paragraph("Clause after the table"), 30, 1)

        imported = aim.from_docx(self._with_numbering(doc, self._chain()))
        html = "\n".join(c.html for c in imported.chunks)
        assert "1.2" in html, "the cell paragraph lost its number"
        assert "1.3" in html, "the counter did not advance for the cell paragraph"

    def test_a_continuation_instance_does_not_tear_the_scheme(self):
        # Word mints a fresh w:num for the same definition whenever a list is
        # interrupted. Judged per instance the continuation uses only deep
        # levels, fails the contiguity rule alone, and emits as <li> — one
        # visible sequence rendered as blocks and then as a fresh list at 1.
        doc = Document()
        self._number(doc.add_paragraph("Top"), 30, 0)
        self._number(doc.add_paragraph("One"), 30, 1)
        self._number(doc.add_paragraph("Deep one"), 30, 2)
        self._number(doc.add_paragraph("Deep two"), 2, 2)  # same abstract, new instance

        imported = aim.from_docx(self._with_numbering(doc, self._chain(num_ids=(30, 2))))
        html = "\n".join(c.html for c in imported.chunks)
        assert "<ol" not in html, "the continuation instance tore the scheme into a list"
        assert re.findall(r'class="num-(\d)"', html) == ["1", "2", "3", "3"]

    def test_a_chained_scheme_the_vocabulary_cannot_draw_bakes_its_label(self):
        # upperRoman over decimal children on PLAIN paragraphs: not a heading,
        # not outline-drawable. Spec §3.8 says such a writer MUST write the
        # computed number as text; dropping it loses the numbering outright.
        numbering = (
            f'<w:numbering xmlns:w="{_W}"><w:abstractNum w:abstractNumId="95">'
            + self._lvl(0, "Article %1", fmt="upperRoman")
            + self._lvl(1, "%1.%2")
            + '</w:abstractNum><w:num w:numId="95">'
            '<w:abstractNumId w:val="95"/></w:num></w:numbering>'
        )
        doc = Document()
        self._number(doc.add_paragraph("Definitions"), 95, 0)
        self._number(doc.add_paragraph("means x"), 95, 1)
        self._number(doc.add_paragraph("means y"), 95, 1)

        imported = aim.from_docx(self._with_numbering(doc, numbering))
        html = "\n".join(c.html for c in imported.chunks)
        assert "Article I" in html, "the top-level label vanished"
        assert "I.1" in html and "I.2" in html, "the chained labels vanished"

    def test_a_second_scheme_starts_at_one_rather_than_continuing_the_first(self):
        # Two independent outline schemes share the same CSS counters, so the
        # second must say "restart" or it carries on from the first: its
        # opening clause renders 2. where the document says 1.
        doc = Document()
        for text, ilvl in (("A top", 0), ("A sub", 1)):
            self._number(doc.add_paragraph(text), 30, ilvl)
        doc.add_paragraph("Interlude")
        for text, ilvl in (("B top", 0), ("B sub", 1)):
            self._number(doc.add_paragraph(text), 31, ilvl)

        numbering = (
            f'<w:numbering xmlns:w="{_W}"><w:abstractNum w:abstractNumId="30">'
            + self._lvl(0, "%1.")
            + self._lvl(1, "%1.%2")
            + '</w:abstractNum><w:abstractNum w:abstractNumId="31">'
            + self._lvl(0, "%1.")
            + self._lvl(1, "%1.%2")
            + '</w:abstractNum><w:num w:numId="30"><w:abstractNumId w:val="30"/></w:num>'
            '<w:num w:numId="31"><w:abstractNumId w:val="31"/></w:num></w:numbering>'
        )
        imported = aim.from_docx(self._with_numbering(doc, numbering))
        html = "\n".join(c.html for c in imported.chunks)
        blocks = re.findall(r'class="([^"]*num-1[^"]*)"[^>]*>([^<]*)', html)
        second = [cls for cls, text in blocks if "B top" in text]
        assert second and "num-restart" in second[0], (
            f"the second scheme continues the first's counters: {blocks}"
        )

    def test_an_empty_numbered_paragraph_still_advances_the_counter(self):
        # Word draws a number for an empty numbered paragraph and the next one
        # carries on from it. Dropping the paragraph without advancing shifts
        # every label after it down by one.
        doc = Document()
        self._number(doc.add_paragraph("Top"), 30, 0)
        self._number(doc.add_paragraph("One"), 30, 1)
        self._number(doc.add_paragraph(""), 30, 1)  # empty, but numbered
        self._number(doc.add_paragraph("Three"), 30, 1)

        imported = aim.from_docx(self._with_numbering(doc, self._chain()))
        html = "\n".join(c.html for c in imported.chunks)
        texts = re.findall(r">([^<]*Three[^<]*)<", html)
        assert texts, "the trailing clause vanished"
        # three sub-clauses were numbered, so the last one is 1.3 — whether it
        # is drawn dynamically or baked, the counter must have moved three times
        assert html.count('class="num-2"') == 3, (
            "the empty numbered paragraph did not consume a number"
        )

    def test_a_list_reopened_after_prose_continues_its_numbering(self):
        # Word keeps counting across the interruption — the fifth item is 5,
        # not 1. A fresh <ol> with no start renders 1. again, so the document
        # silently contains two item 1s and no item 3.
        numbering = (
            f'<w:numbering xmlns:w="{_W}"><w:abstractNum w:abstractNumId="40">'
            + self._lvl(0, "%1.")
            + '</w:abstractNum><w:num w:numId="40">'
            '<w:abstractNumId w:val="40"/></w:num></w:numbering>'
        )
        doc = Document()
        self._number(doc.add_paragraph("One"), 40, 0)
        self._number(doc.add_paragraph("Two"), 40, 0)
        doc.add_paragraph("An interrupting paragraph.")
        self._number(doc.add_paragraph("Three"), 40, 0)
        self._number(doc.add_paragraph("Four"), 40, 0)

        imported = aim.from_docx(self._with_numbering(doc, numbering))
        # the rendered document only — the history log repeats every markup
        content = imported.dumps().split("<body", 1)[1].split("<script", 1)[0]
        assert content.count("<ol") == 2, "the interruption should split the list"
        assert 'start="3"' in content, f"the reopened list restarts at 1: {content}"


class TestExportedNumberingIsOneSequence:
    """What Word draws from the exported file must be what the .aim draws.

    Both defects here were invisible to a test that reads the XML: the parts
    are all present and well-formed, they just belong to different counter
    streams or to no numbering at all.
    """

    @staticmethod
    def _contract() -> aim.AimDocument:
        doc = aim.new_document(title="Contract")
        who = aim.external("t")
        for markup in (
            '<h1 class="num-1" data-aim-num-prefix="Article ">Scope</h1>',
            '<p class="num-2">First sub-clause.</p>',
            '<p class="num-2">Second sub-clause.</p>',
            '<h1 class="num-1" data-aim-num-prefix="Article ">Term</h1>',
            '<p class="num-2">Third sub-clause.</p>',
        ):
            doc.add_chunk(markup, author=who)
        return doc

    @staticmethod
    def _numbering_of(path) -> tuple[dict[int, int], list[int]]:
        """``{numId: abstractNumId}`` and the numIds document.xml actually uses."""
        with zipfile.ZipFile(str(path)) as z:
            numbering = etree.fromstring(z.read("word/numbering.xml"))
            document = etree.fromstring(z.read("word/document.xml"))
        w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        mapping = {
            int(n.get(f"{w}numId")): int(n.find(f"{w}abstractNumId").get(f"{w}val"))
            for n in numbering.findall(f"{w}num")
        }
        used = [int(e.get(f"{w}val")) for e in document.iter(f"{w}numId") if e.get(f"{w}val")]
        return mapping, used

    def test_a_prefixed_scheme_stays_one_counter_stream(self, tmp_path):
        # A literal lives in the level's lvlText, so a per-block prefix set
        # minted an abstract per shape: the "Article" headings counted in one
        # definition and their sub-clauses in another, whose level-0 counter
        # nothing ever advanced. Word renders Article 2 / 1.3 where the
        # document says Article 2 / 2.1 — and OOXML cannot share a counter
        # across two definitions, so this has to be one.
        out = aim.to_docx(self._contract(), tmp_path / "contract.docx")
        mapping, used = self._numbering_of(out)
        abstracts = {mapping[n] for n in used if n in mapping}
        assert len(abstracts) == 1, f"the scheme was split across {abstracts}"

    def test_the_exported_prefix_rides_the_level_it_belongs_to(self, tmp_path):
        out = aim.to_docx(self._contract(), tmp_path / "contract.docx")
        with zipfile.ZipFile(str(out)) as z:
            numbering = z.read("word/numbering.xml").decode()
        assert 'w:val="Article %1"' in numbering, numbering[:400]

    def test_the_export_reimports_as_the_same_shape(self, tmp_path):
        # the round trip is where a split shows up structurally: the
        # sub-clauses come back as a plain <ol> because their scheme no
        # longer uses level 0
        out = aim.to_docx(self._contract(), tmp_path / "contract.docx")
        back = aim.from_docx(str(out))
        content = back.dumps().split("<body", 1)[1].split("<script", 1)[0]
        assert content.count("<ol") == 0, "the sub-clauses degraded to a list"
        assert re.findall(r'class="num-(\d)', content) == ["1", "2", "2", "1", "2"]

    def test_a_pending_change_keeps_the_clause_numbered(self, tmp_path):
        # tracked export is the DEFAULT, and it built its revision paragraphs
        # without ever numbering them: every clause touched by a pending
        # proposal lost its w:numPr, so the delivered file showed the rest of
        # the contract shifted up by one.
        doc = self._contract()
        doc.propose_modify(
            doc.chunks[1].id,
            '<p class="num-2">First sub-clause, amended.</p>',
            author=aim.external("bot"),
            explanation="tighter",
        )
        out = aim.to_docx(doc, tmp_path / "tracked.docx", pending="tracked")
        _, used = self._numbering_of(out)
        # four untouched clauses, plus the struck original and its
        # replacement — Word numbers a pending deletion until it is accepted
        assert len(used) == 6, f"a tracked clause lost its numbering: {len(used)} numbered of 6"


class TestThemeFontsComeFromTheStylesWordRenders:
    """The document's face is what the STYLE resolves to, not what the theme
    table declares — and a style may name its font through the theme, or
    carry a localized styleId, or not exist at all under that name.

    Every miss here is the same silent failure: the whole document renders in
    the wrong family, and a test that only asks whether the slot is non-empty
    never notices.
    """

    @staticmethod
    def _docx(patch_styles=None, *, major="Georgia", minor="Verdana", heading_style=True):
        doc = Document()
        doc.add_paragraph("Body text")
        if heading_style:
            doc.add_paragraph("A heading").style = "Heading 1"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        src = zipfile.ZipFile(buf)
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as z:
            for name in src.namelist():
                data = src.read(name)
                if name == "word/theme/theme1.xml":
                    s = data.decode()
                    s = re.sub(r'(<a:majorFont>\s*<a:latin typeface=")[^"]*', r"\g<1>" + major, s)
                    s = re.sub(r'(<a:minorFont>\s*<a:latin typeface=")[^"]*', r"\g<1>" + minor, s)
                    data = s.encode()
                elif name == "word/styles.xml" and patch_styles is not None:
                    data = patch_styles(data.decode()).encode()
                z.writestr(name, data)
        out.seek(0)
        return out

    @staticmethod
    def _heading_slot(source) -> str | None:
        from aimformat.convert._docx_in import _derived_theme_slots
        from aimformat.convert._docx_seam import parse_docx

        return _derived_theme_slots(parse_docx(source)).get("--aim-font-heading")

    def test_a_style_naming_its_font_through_the_theme_is_resolved(self):
        # Word's own default heading style carries asciiTheme="majorHAnsi"
        # rather than a literal face. The parse layer drops rFonts entirely
        # when it is a theme reference, so the resolved props came back empty
        # and the slot silently fell back to the theme table. Pointed at the
        # MINOR font, the two answers differ and the bug is visible.
        def minor_ref(s):
            return re.sub(
                r'(<w:style [^>]*w:styleId="Heading1">)',
                r'\1<w:rPr><w:rFonts w:asciiTheme="minorHAnsi" w:hAnsiTheme="minorHAnsi"/></w:rPr>',
                s,
            )

        assert self._heading_slot(self._docx(minor_ref)) == "Verdana"

    def test_a_localized_style_id_is_found_by_its_english_name(self):
        # A German Word writes w:styleId="berschrift1" and keeps
        # <w:name w:val="heading 1"/>. Keyed on the styleId alone, every
        # non-English document loses its heading face.
        def localized(s):
            s = s.replace('w:styleId="Heading1"', 'w:styleId="berschrift1"')
            return re.sub(
                r'(<w:style [^>]*w:styleId="berschrift1">)',
                r'\1<w:rPr><w:rFonts w:ascii="Palatino" w:hAnsi="Palatino"/></w:rPr>',
                s,
            )

        assert self._heading_slot(self._docx(localized)) == "Palatino"

    def test_a_document_without_heading_1_falls_through_to_the_next_level(self):
        # Plenty of documents start at Heading 2. Looking only for Heading 1
        # leaves the slot on the theme table's word.
        def only_h2(s):
            return re.sub(
                r'(<w:style [^>]*w:styleId="Heading2">)',
                r'\1<w:rPr><w:rFonts w:ascii="Baskerville" w:hAnsi="Baskerville"/></w:rPr>',
                s,
            )

        doc = Document()
        doc.add_paragraph("Body")
        doc.add_paragraph("Sub heading").style = "Heading 2"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        src = zipfile.ZipFile(buf)
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as z:
            for name in src.namelist():
                data = src.read(name)
                if name == "word/theme/theme1.xml":
                    s = data.decode()
                    s = re.sub(r'(<a:majorFont>\s*<a:latin typeface=")[^"]*', r"\g<1>Georgia", s)
                    data = s.encode()
                elif name == "word/styles.xml":
                    data = only_h2(data.decode()).encode()
                z.writestr(name, data)
        out.seek(0)
        assert self._heading_slot(out) == "Baskerville"


class TestMalformedInputFailsTyped:
    """An unreadable file is a user error, not a crash. Without a typed
    error the CLI prints a traceback ending in someone else's exception
    class, and every caller has to guess what to catch."""

    @staticmethod
    def _import(data: bytes):
        return aim.from_docx(io.BytesIO(data))

    def test_an_empty_zip_raises_a_typed_error(self):
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w"):
            pass
        with pytest.raises(aim.ParseError):
            self._import(buf.getvalue())

    def test_a_truncated_archive_raises_a_typed_error(self):
        doc = Document()
        doc.add_paragraph("ok")
        buf = io.BytesIO()
        doc.save(buf)
        with pytest.raises(aim.ParseError):
            self._import(buf.getvalue()[: len(buf.getvalue()) // 2])

    def test_a_file_that_is_not_a_zip_raises_a_typed_error(self):
        # a legacy .doc renamed .docx is the everyday version of this
        with pytest.raises(aim.ParseError):
            self._import(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1 legacy OLE compound file")


class TestListsStayAddressable:
    """A list is a CONTAINER whose items carry their own ids — that is what
    lets a proposal target one item. The rule that decides this matched the
    literal string "<ol>", so the moment a list carried an attribute it
    silently became one atomic chunk and every item lost its id."""

    @staticmethod
    def _containerize(markup: str) -> str:
        from aimformat.ingest import _containerize

        return _containerize(markup)

    def test_a_list_with_an_attribute_is_still_a_container(self):
        out = self._containerize('<ol start="6"><li>six</li><li>seven</li></ol>')
        assert "data-aim-container" in out, out
        assert out.count('<li data-aim=""') == 2, out

    def test_the_plain_list_is_unchanged(self):
        out = self._containerize("<ol><li>one</li></ol>")
        assert "data-aim-container" in out
        assert out.count('<li data-aim=""') == 1

    def test_the_real_document_keeps_its_items_addressable(self):
        # sample3's reopened list is the one that grew a start attribute
        imported = aim.from_docx("tests/fixtures/docxs/sample3.docx")
        content = imported.dumps().split("<body", 1)[1].split("<script", 1)[0]
        for ol in re.findall(r"<ol[^>]*>.*?</ol>", content, re.S):
            assert "data-aim-container" in ol, ol[:160]
            assert "<li data-aim=" in ol, ol[:160]
