"""Regression tests for findings from the v0.1 self-review.

Each test pins a bug that shipped in the initial 0.1.0 commit and was found
during the post-ship review (docs/log review entry). Keep them even when
they look redundant with broader tests — they encode the exact failure.
"""

import re

import pytest

import aimformat as aim
from aimformat.errors import InvalidOperation
from conftest import BOT, ME, ts


@pytest.fixture
def table_doc():
    doc = aim.new_document(title="T")
    doc.add_chunk(
        '<table data-aim-container="tbl">'
        '<thead><tr data-aim="h"><th>K</th></tr></thead>'
        '<tbody><tr data-aim="r1"><td>1</td></tr>'
        '<tr data-aim="r2"><td>2</td></tr></tbody></table>',
        author=ME,
        at=ts(0),
    )
    return doc


class TestShellAnchors:
    """Deleting the first tbody row must not un-delete into the thead."""

    def test_delete_first_body_row_records_shell(self, table_doc):
        table_doc.delete_chunk("r1", author=ME, at=ts(1))
        anchor = table_doc.history[-1].get("anchor")
        assert anchor == {"after": None, "container": "tbl", "shell": "tbody"}

    def test_undo_restores_row_into_its_shell(self, table_doc):
        before = table_doc._state.serial("tbl")
        table_doc.delete_chunk("r1", author=ME, at=ts(1))
        table_doc.undo(author=ME, at=ts(2))
        assert table_doc._state.serial("tbl") == before
        assert table_doc.verify() == []

    def test_delete_header_row_round_trips(self, table_doc):
        before = table_doc._state.serial("tbl")
        table_doc.delete_chunk("h", author=ME, at=ts(1))
        assert table_doc.history[-1].get("anchor")["shell"] == "thead"
        table_doc.undo(author=ME, at=ts(2))
        assert table_doc._state.serial("tbl") == before
        assert table_doc.verify() == []

    def test_state_at_walks_shelled_deletes(self, table_doc):
        h0 = table_doc.doc_hash
        table_doc.delete_chunk("r1", author=ME, at=ts(1))
        past = table_doc.state_at(1)
        assert past.doc_hash == h0

    def test_list_anchors_carry_no_shell(self, rich_doc):
        rich_doc.delete_chunk("li1", author=ME, at=ts(30))
        assert "shell" not in rich_doc.history[-1].get("anchor")


class TestUndoRedoZone:
    """Stacked and interleaved undo/redo walk the zone as a proper stack."""

    def test_undo_undo_redo_redo_full_cycle(self, basic_doc):
        ids_full = basic_doc.body_ids
        basic_doc.undo(author=ME, at=ts(10))
        basic_doc.undo(author=ME, at=ts(11))
        assert basic_doc.body_ids == []
        basic_doc.redo(author=ME, at=ts(12))
        basic_doc.redo(author=ME, at=ts(13))
        assert basic_doc.body_ids == ids_full
        assert basic_doc.verify() == []

    def test_undo_after_redo_targets_the_redone_edit(self, basic_doc):
        basic_doc.undo(author=ME, at=ts(10))  # undoes intro-add
        basic_doc.redo(author=ME, at=ts(11))  # restores intro
        basic_doc.undo(author=ME, at=ts(12))  # must undo intro again
        assert basic_doc.body_ids == ["h1"]
        assert basic_doc.verify() == []

    def test_third_redo_raises_cleanly(self, basic_doc):
        basic_doc.undo(author=ME, at=ts(10))
        basic_doc.redo(author=ME, at=ts(11))
        with pytest.raises(InvalidOperation):
            basic_doc.redo(author=ME, at=ts(12))


class TestThemePayloadValidation:
    def test_accept_with_hostile_theme_tweak_rejected(self, basic_doc):
        p = basic_doc.propose_theme({"--aim-brand-1": "#333333"}, author=BOT, at=ts(10))
        with pytest.raises(InvalidOperation):
            basic_doc.accept(
                p.id,
                decided_by=ME,
                at=ts(11),
                applied="<style data-aim-theme>:root{--aim-brand-1:"
                "#333333} body{background:red}</style>",
            )

    def test_accept_with_valid_theme_tweak(self, basic_doc):
        p = basic_doc.propose_theme({"--aim-brand-1": "#333333"}, author=BOT, at=ts(10))
        basic_doc.accept(
            p.id,
            decided_by=ME,
            at=ts(11),
            applied="<style data-aim-theme>:root{--aim-brand-1:#444444}</style>",
        )
        assert basic_doc.theme["--aim-brand-1"] == "#444444"
        assert basic_doc.verify() == []

    def test_propose_raw_theme_markup_validated(self, basic_doc):
        with pytest.raises(InvalidOperation):
            basic_doc.propose_modify(
                "aim:theme",
                "<style data-aim-theme>:root{--aim-evil:#000}</style>",
                author=BOT,
                at=ts(10),
            )


class TestSupersededByIntegrity:
    def test_superseded_by_is_never_a_placeholder(self, basic_doc):
        basic_doc.propose_modify("intro", '<p data-aim="intro">v1</p>', author=BOT, at=ts(10))
        p2 = basic_doc.propose_modify("intro", '<p data-aim="intro">v2</p>', author=BOT, at=ts(11))
        ev = next(e for e in basic_doc.history if e.get("decision") == "superseded")
        assert ev.get("superseded_by") == p2.id
        assert "(new)" not in basic_doc.dumps()

    def test_supersede_chain_of_three(self, basic_doc):
        basic_doc.propose_modify("intro", '<p data-aim="intro">v1</p>', author=BOT, at=ts(10))
        basic_doc.propose_modify("intro", '<p data-aim="intro">v2</p>', author=BOT, at=ts(11))
        p3 = basic_doc.propose_modify("intro", '<p data-aim="intro">v3</p>', author=BOT, at=ts(12))
        assert [p.id for p in basic_doc.proposals] == [p3.id]
        chain = [
            e.get("superseded_by") for e in basic_doc.history if e.get("decision") == "superseded"
        ]
        assert len(chain) == 2 and all(chain)
        assert basic_doc.verify() == []


# ===========================================================================
# Wave 2: findings from the three independent review agents
# ===========================================================================


class TestAnchorResolutionFamily:
    """The systemic family: resolution must honor the context it claims."""

    def test_move_to_last_when_already_last_is_a_noop_error(self, basic_doc):
        with pytest.raises(InvalidOperation):
            basic_doc.move_chunk("intro", author=ME, at=ts(10))  # already last
        basic_doc.chunk("intro")  # still present, tree unharmed
        assert basic_doc.verify() == []

    def test_failed_move_never_mutates(self, rich_doc):
        before = rich_doc.doc_hash
        with pytest.raises(aim.TargetNotFound):
            rich_doc.move_chunk("li1", container="list", after="ghost", author=ME, at=ts(30))
        assert rich_doc.doc_hash == before and rich_doc.verify() == []

    def test_delete_after_nested_container_records_container_anchor(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<aim-slide data-aim-container="s1" '
            'style="width:1920px; height:1080px">'
            '<ul data-aim-container="lst" '
            'style="left:10px; top:10px; width:100px">'
            '<li data-aim="i1">a</li></ul>'
            '<p data-aim="x" style="left:10px; top:200px">after</p>'
            "</aim-slide>",
            author=ME,
            at=ts(0),
        )
        doc.checkpoint("cp", at=ts(1))
        doc.delete_chunk("x", author=ME, at=ts(2))
        ev = doc.history[-1]
        assert ev.get("anchor")["after"] == "lst"
        doc.undo(author=ME, at=ts(3))
        assert doc.verify() == []
        past = doc.state_at(1)
        slide_kids = [
            e.chunk_id or e.container_id for e in past._state.container_node("s1").elements()
        ]
        assert slide_kids == ["lst", "x"]

    def test_delete_of_nested_container_itself_round_trips(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<aim-slide data-aim-container="s1" '
            'style="width:1920px; height:1080px">'
            '<h2 data-aim="t" style="left:10px; top:10px">T</h2>'
            '<ul data-aim-container="lst" '
            'style="left:10px; top:100px; width:100px">'
            '<li data-aim="i1">a</li></ul></aim-slide>',
            author=ME,
            at=ts(0),
        )
        h0 = doc.doc_hash
        doc.delete_chunk("lst", author=ME, at=ts(1))
        doc.undo(author=ME, at=ts(2))
        assert doc.doc_hash == h0 and doc.verify() == []

    def test_insert_anchor_must_live_in_stated_container(self, rich_doc):
        rich_doc.add_chunk(
            '<ul data-aim-container="lb"><li data-aim="b1">x</li></ul>', author=ME, at=ts(30)
        )
        with pytest.raises(aim.TargetNotFound):
            rich_doc.add_chunk("<li>stray</li>", author=ME, container="list", after="b1", at=ts(31))

    def test_last_in_slide_ignores_items_of_nested_containers(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<aim-slide data-aim-container="s1" '
            'style="width:1920px; height:1080px">'
            '<ul data-aim-container="lst" '
            'style="left:10px; top:10px; width:100px">'
            '<li data-aim="i1">a</li></ul></aim-slide>',
            author=ME,
            at=ts(0),
        )
        doc.add_chunk(
            '<p data-aim="cap" style="left:10px; top:400px">caption</p>',
            author=ME,
            container="s1",
            at=ts(1),
        )
        slide = doc._state.container_node("s1")
        assert [e.chunk_id or e.container_id for e in slide.elements()] == ["lst", "cap"]
        assert doc.history[-1].get("anchor") == {"after": "lst", "container": "s1"}

    def test_public_add_first_into_table_defaults_to_tbody(self, table_doc):
        doc = table_doc
        doc.add_chunk(
            '<tr data-aim="r0"><td>zero</td></tr>', author=ME, container="tbl", after=None, at=ts(1)
        )
        html = doc._state.serial("tbl")
        thead = html[html.index("<thead>") : html.index("</thead>")]
        assert "r0" not in thead and doc.verify() == []
        assert doc.history[-1].get("anchor")["shell"] == "tbody"

    def test_chunk_ids_with_proposal_prefix_are_reassigned(self, basic_doc):
        c = basic_doc.add_chunk('<p data-aim="p-note1">n</p>', author=ME, at=ts(10))
        assert not c.id.startswith("p-")


class TestPayloadAndIdIntegrity:
    def test_modify_container_keeps_marker_and_covers_new_items(self, rich_doc):
        got = rich_doc.modify_chunk(
            "list",
            '<ul data-aim-container="list"><li data-aim="li1">First</li><li>NEW ITEM</li></ul>',
            author=ME,
            at=ts(30),
        )
        assert got.id == "list"
        assert "list" in rich_doc.containers
        html = rich_doc._state.serial("list")
        assert 'data-aim-container="list"' in html
        assert html.count("data-aim=") == 2  # li1 kept + NEW ITEM covered
        assert not [f for f in aim.lint(rich_doc) if f.level == "error"]

    def test_modify_container_with_chunk_marker_rejected(self, rich_doc):
        with pytest.raises(InvalidOperation):
            rich_doc.modify_chunk(
                "list", '<ul data-aim="list"><li>x</li></ul>', author=ME, at=ts(30)
            )

    def test_payload_only_ids_stay_burned(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<ul data-aim-container="lst"><li data-aim="item1">a</li></ul>', author=ME, at=ts(0)
        )
        doc.delete_chunk("lst", author=ME, at=ts(1))
        c = doc.add_chunk('<p data-aim="item1">new life</p>', author=ME, at=ts(2))
        assert c.id != "item1"

    def test_accepted_move_resolution_passes_event_schema(self, basic_doc):
        pr = basic_doc.propose_move("intro", container="body", after=None, author=BOT, at=ts(10))
        basic_doc.accept(pr.id, decided_by=ME, at=ts(11))
        ev = basic_doc.history[-1]
        assert ev.get("from") and ev.get("to")
        assert ev.validate() == []
        assert not [f for f in aim.lint(basic_doc) if f.level == "error"]

    def test_theme_value_grammar_enforced_on_write(self, basic_doc):
        with pytest.raises(InvalidOperation):
            basic_doc.set_theme(
                {"--aim-brand-1": 'url("https://evil.example/x")'}, author=ME, at=ts(10)
            )
        with pytest.raises(InvalidOperation):
            basic_doc.propose_theme({"--aim-font-body": "x;}body{color:red"}, author=BOT, at=ts(11))

    def test_pack_assets_atomic_on_undecodable_image(self):
        from test_css_assets_meta import DATA_URI

        doc = aim.new_document(title="T")
        doc.add_chunk(
            f'<figure data-aim="f"><img alt="ok" src="{DATA_URI}">'
            '<img alt="bad" src="data:image/svg+xml,<svg/>">'
            "</figure>",
            author=ME,
            at=ts(0),
        )
        h0 = doc.doc_hash
        with pytest.raises(InvalidOperation):
            doc.pack_assets(author=aim.external("packer"), at=ts(1))
        assert doc.doc_hash == h0 and doc.verify() == []

    def test_pack_assets_events_share_one_batch(self):
        from test_css_assets_meta import DATA_URI

        doc = aim.new_document(title="T")
        for i, cid in enumerate(("f1", "f2")):
            doc.add_chunk(
                f'<figure data-aim="{cid}"><img alt="a" src="{DATA_URI}"></figure>',
                author=ME,
                at=ts(i),
            )
        doc.pack_assets(author=aim.external("packer"), at=ts(5))
        packs = [e for e in doc.history if e.action == "modify"]
        assert len({e.batch for e in packs}) == 1

    def test_prune_refuses_to_drop_everything(self, rich_doc):
        with pytest.raises(InvalidOperation):
            rich_doc.prune(before=10_000)


class TestVerifierHardening:
    """lint_text must convert hostile input into findings, never raise."""

    HOSTILE_META = '<script type="application/aim-meta+json">\n{not json]\n</script>\n'

    def hostile(self, basic_doc, mutate):
        return mutate(basic_doc.dumps())

    @pytest.mark.parametrize(
        "mutate",
        [
            lambda t: t.replace("<title>", TestVerifierHardening.HOSTILE_META + "<title>"),
            lambda t: t.replace(
                "</body>",
                '<script type="application/aim-embeddings+jsonl">\n[1,2,3]\n</script>\n</body>',
            ),
            lambda t: t.replace('{"action":"add"', "42 ", 1),
            lambda t: t.replace(
                "</body>",
                '<script type="application/'
                'aim-meta+json">\n{"summary":"a string"}\n'
                "</script>\n</body>",
            ),
        ],
        ids=["malformed-meta", "embeddings-array", "history-non-object", "summary-not-dict"],
    )
    def test_never_raises_on_hostile_caches(self, basic_doc, mutate):
        findings = aim.lint_text(self.hostile(basic_doc, mutate))
        assert any(f.level == "error" for f in findings)

    def test_meta_missing_summary_is_M004(self, basic_doc):
        text = basic_doc.dumps().replace(
            "<title>", '<script type="application/aim-meta+json">\n{"toc":[]}\n</script>\n<title>'
        )
        assert "M004" in {f.code for f in aim.lint_text(text)}

    def test_asset_registry_not_exempt_from_security(self, basic_doc):
        text = basic_doc.dumps().replace(
            "</body>",
            "<aim-assets>\n"
            '<svg aria-hidden="true" height="0" width="0">\n'
            '<symbol id="asset-aaaaaaaaaaaa" viewBox="0 0 10 10">'
            '<image height="10" width="10" '
            'href="javascript:alert(1)"/></symbol>\n'
            "</svg>\n</aim-assets>\n</body>",
        )
        codes = {f.code for f in aim.lint_text(text) if f.level == "error"}
        assert codes & {"X003", "V009"}

    def test_container_modify_proposal_lints_clean(self, rich_doc):
        rich_doc.propose_modify(
            "list",
            '<ul data-aim-container="list"><li data-aim="li1">First</li></ul>',
            author=BOT,
            at=ts(30),
        )
        assert not [f for f in aim.lint_text(rich_doc.dumps()) if f.level == "error"]

    def test_nested_chunk_is_S024(self, basic_doc):
        text = basic_doc.dumps().replace(
            '<p data-aim="intro">Intro paragraph.</p>',
            '<section data-aim="intro"><p data-aim="inner">x</p></section>',
        )
        assert "S024" in {f.code for f in aim.lint_text(text)}

    def test_add_anchor_cycle_is_P015(self, basic_doc):
        basic_doc.propose_add("<p>one</p>", author=BOT, at=ts(10))
        text = basic_doc.dumps()
        card = text[
            text.index("<aim-proposal ") : text.index("</aim-proposal>") + len("</aim-proposal>")
        ]
        pid = card.split('id="')[1].split('"')[0]
        looped = card.replace('data-anchor-after="intro"', 'data-anchor-after="p-zzzz"')
        twin = card.replace(pid, "p-zzzz").replace(
            'data-anchor-after="intro"', f'data-anchor-after="{pid}"'
        )
        text = text.replace(card, looped + "\n" + twin)
        assert "P015" in {f.code for f in aim.lint_text(text)}


class TestNormalFormHardening:
    def test_doc_hash_single_valued_for_void_and_style_spellings(self):
        a = aim.loads(_mini('<p data-aim="x" style="left:3px; top:5px">t<br></p>'))
        b = aim.loads(_mini('<p data-aim="x" style="top:5px;left:3px">t<br/></p>'))
        assert a.doc_hash == b.doc_hash

    def test_c001_rejects_non_normal_style_order(self, basic_doc):
        basic_doc.add_chunk(
            '<aim-slide data-aim-container="s" '
            'style="width:1920px; height:1080px">'
            '<p data-aim="q" style="left:1px; top:2px">x</p>'
            "</aim-slide>",
            author=ME,
            at=ts(10),
        )
        text = basic_doc.dumps().replace("left:1px; top:2px", "top:2px;left:1px")
        assert "C001" in {f.code for f in aim.lint_text(text)}

    def test_class_tokens_deduped(self):
        doc = aim.new_document(title="T")
        c = doc.add_chunk('<p class="font-bold font-bold text-lg">x</p>', author=ME, at=ts(0))
        assert 'class="font-bold text-lg"' in c.html

    def test_duplicate_style_props_last_wins(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<aim-slide data-aim-container="s" '
            'style="width:1920px; height:1080px">'
            '<p data-aim="q" style="left:1px; left:9px; top:2px">x'
            "</p></aim-slide>",
            author=ME,
            at=ts(0),
        )
        assert 'style="left:9px; top:2px"' in doc._state.serial("q")


# ===========================================================================
# Wave 3: findings from the Codex deep code review (2026-07-08)
# ===========================================================================


class TestFragmentIsTheTrustBoundary:
    """AIM-01: lint validates the whole parsed file, not just the first
    <html>. Forbidden markup outside the modeled body must not lint clean."""

    def test_top_level_script_after_html_is_S028(self, basic_doc):
        text = basic_doc.dumps() + "<script>alert(1)</script>\n"
        assert "S028" in {f.code for f in aim.lint_text(text)}

    def test_top_level_element_after_html_is_S028(self, basic_doc):
        text = basic_doc.dumps() + "<p>loose</p>\n"
        assert "S028" in {f.code for f in aim.lint_text(text)}

    def test_second_html_element_is_S028(self, basic_doc):
        text = basic_doc.dumps() + '<html data-aim-version="0.1"></html>\n'
        assert "S028" in {f.code for f in aim.lint_text(text)}

    def test_forbidden_head_child_is_flagged(self, basic_doc):
        text = basic_doc.dumps().replace(
            "</title>", '</title>\n<iframe src="https://evil.example"></iframe>'
        )
        codes = {f.code for f in aim.lint_text(text)}
        assert codes & {"S029", "X001"}

    def test_event_handler_on_proposal_card_is_X002(self, basic_doc):
        basic_doc.propose_delete("intro", author=BOT, at=ts(10))
        canon = aim.loads(
            basic_doc.dumps().replace("<aim-proposal ", '<aim-proposal onclick="alert(1)" ', 1)
        ).dumps()
        assert "X002" in {f.code for f in aim.lint_text(canon)}


class TestEmbeddedCssIsVerified:
    """AIM-02: the machine-managed aim.css block is trusted at the raw tier,
    so lint must pin it to the generated stylesheet."""

    def test_tampered_aim_css_is_X006(self, basic_doc):
        v = re.escape(aim.SPEC_VERSION)
        bad = re.sub(
            rf'<style data-aim-css="{v}">[\s\S]*?</style>',
            f'<style data-aim-css="{aim.SPEC_VERSION}">'
            "\n@import url(https://evil.example/x"
            ".css);\nbody{background:red}\n</style>",
            basic_doc.dumps(),
        )
        assert "X006" in {f.code for f in aim.lint_text(bad)}

    def test_generated_css_lints_clean(self, basic_doc):
        assert "X006" not in {f.code for f in aim.lint_text(basic_doc.dumps())}


class TestProposalAnchorsAreContainerScoped:
    """AIM-03: proposal anchors resolve container-scoped, like direct ops —
    a proposal that can never be accepted must never be created or lint clean."""

    def test_propose_add_cross_container_anchor_raises(self, rich_doc):
        with pytest.raises(aim.TargetNotFound):
            rich_doc.propose_add(
                "<li>x</li>", author=BOT, container="list", after="intro", at=ts(30)
            )

    def test_propose_move_to_foreign_anchor_raises(self, rich_doc):
        with pytest.raises(aim.TargetNotFound):
            rich_doc.propose_move("li1", author=BOT, container="body", after="row1", at=ts(30))

    def test_cross_container_add_anchor_is_P016(self, rich_doc):
        rich_doc.propose_add(
            '<li data-aim="x">new</li>', author=BOT, container="list", after="li1", at=ts(30)
        )
        text = rich_doc.dumps().replace('data-anchor-after="li1"', 'data-anchor-after="intro"')
        assert "P016" in {f.code for f in aim.lint_text(text)}


class TestDuplicateProposalIds:
    """AIM-04: duplicate pending ids are rejected, not silently shadowed."""

    def test_duplicate_ids_are_P017(self, basic_doc):
        basic_doc.propose_add("<p>one</p>", author=BOT, at=ts(10))
        p2 = basic_doc.propose_add("<p>two</p>", author=BOT, at=ts(11))
        p1_id = basic_doc.proposals[0].id
        text = basic_doc.dumps().replace(p2.id, p1_id)
        assert "P017" in {f.code for f in aim.lint_text(text)}

    def test_accept_on_duplicate_id_raises(self, basic_doc):
        basic_doc.propose_add("<p>one</p>", author=BOT, at=ts(10))
        p2 = basic_doc.propose_add("<p>two</p>", author=BOT, at=ts(11))
        p1_id = basic_doc.proposals[0].id
        dup = aim.loads(basic_doc.dumps().replace(p2.id, p1_id))
        with pytest.raises(InvalidOperation):
            dup.accept(p1_id, decided_by=ME, at=ts(12))


class TestHistoryValidationIsActionAware:
    """AIM-05: malformed history yields targeted H-findings, never a generic
    S000 verifier crash."""

    def _move_doc(self, basic_doc):
        basic_doc.move_chunk("h1", author=ME, container="body", after="intro", at=ts(10))
        return basic_doc.dumps()

    def test_move_missing_to_is_H003_not_S000(self, basic_doc):
        broken = re.sub(r',"to":\{[^}]*?\}', "", self._move_doc(basic_doc), count=1)
        codes = {f.code for f in aim.lint_text(broken)}
        assert "H003" in codes and "S000" not in codes

    def test_move_missing_from_is_not_S000(self, basic_doc):
        broken = re.sub(r'"from":\{[^}]*?\},', "", self._move_doc(basic_doc), count=1)
        codes = {f.code for f in aim.lint_text(broken)}
        assert "S000" not in codes and "H003" in codes

    def test_add_missing_anchor_is_H003(self, basic_doc):
        broken = re.sub(r',"anchor":\{[^}]*?\}', "", basic_doc.dumps(), count=1)
        codes = {f.code for f in aim.lint_text(broken)}
        assert "H003" in codes and "S000" not in codes

    def test_event_missing_seq_is_not_S000(self, basic_doc):
        broken = re.sub(r'"seq":\d+,', "", basic_doc.dumps(), count=1)
        codes = {f.code for f in aim.lint_text(broken)}
        assert "S000" not in codes


class TestUrlSchemeValidation:
    """AIM-06: URL checks match by scheme, not raw prefix."""

    @pytest.mark.parametrize(
        "href", ["httpjavascript:alert(1)", "httpsx://example.com", "mailtox:test"]
    )
    def test_fake_schemes_rejected(self, href):
        doc = aim.new_document(title="T")
        doc.add_chunk(f'<p data-aim="p1"><a href="{href}">l</a></p>', author=ME, at=ts(0))
        codes = {f.code for f in aim.lint_text(doc.dumps())}
        assert codes & {"V009", "X003"}

    @pytest.mark.parametrize(
        "href", ["https://example.com", "http://example.com", "mailto:a@b.co", "#sec"]
    )
    def test_real_schemes_pass(self, href):
        doc = aim.new_document(title="T")
        doc.add_chunk(f'<p data-aim="p1"><a href="{href}">l</a></p>', author=ME, at=ts(0))
        assert not [f for f in aim.lint_text(doc.dumps()) if f.level == "error"]


class TestConstructorThemeValues:
    """AIM-07: new_document validates theme values, not just slot names."""

    def test_bad_theme_value_raises(self):
        with pytest.raises(InvalidOperation):
            aim.new_document(title="x", theme={"--aim-brand-1": "not-a-color"})

    def test_valid_theme_value_lints_clean(self):
        doc = aim.new_document(title="x", theme={"--aim-brand-1": "#123456"})
        assert not [f for f in aim.lint_text(doc.dumps()) if f.level == "error"]


class TestProposeMoveNoop:
    """AIM-08: a no-op move proposal is rejected like the direct move is."""

    def test_propose_move_noop_raises(self, basic_doc):
        with pytest.raises(InvalidOperation):
            basic_doc.propose_move("intro", author=BOT, container="body", after="h1", at=ts(10))

    def test_propose_move_real_move_lints_clean(self, basic_doc):
        p = basic_doc.propose_move("h1", author=BOT, container="body", after="intro", at=ts(10))
        assert p.action == "move"
        assert not [f for f in aim.lint_text(basic_doc.dumps()) if f.level == "error"]


# ===========================================================================
# Wave 4: findings from the agent-core round-2 review (2026-07-13)
# ===========================================================================


class TestTableGridOverflow:
    """A-13: spans exceeding the grid must clamp/widen, never crash the
    DOCX export of a lint-clean table."""

    def _export(self, html, tmp_path):
        docx = pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk(html, author=ME, at=ts(0))
        assert aim.lint(doc) == []
        out = aim.to_docx(doc, tmp_path / "t.docx")
        return docx.Document(str(out))

    def test_rowspan_exceeding_rows_exports(self, tmp_path):
        d = self._export(
            '<table data-aim-container="t1">'
            '<tr data-aim="r1"><td rowspan="3">A</td><td>B</td></tr>'
            '<tr data-aim="r2"><td>C</td></tr></table>',
            tmp_path,
        )
        assert len(d.tables) == 1 and len(d.tables[0].rows) == 2
        texts = {c.text for row in d.tables[0].rows for c in row.cells}
        assert {"A", "B", "C"} <= texts

    def test_colspan_shifted_past_grid_exports(self, tmp_path):
        d = self._export(
            '<table data-aim-container="t1">'
            '<tr data-aim="r1"><td rowspan="2">A</td><td>B</td></tr>'
            '<tr data-aim="r2"><td colspan="2">C</td></tr></table>',
            tmp_path,
        )
        texts = {c.text for row in d.tables[0].rows for c in row.cells}
        assert {"A", "B", "C"} <= texts

    def test_base_cell_pushed_past_grid_widens(self, tmp_path):
        # the surplus cell is a legitimate grid column — it must survive,
        # not be dropped to dodge the crash
        d = self._export(
            '<table data-aim-container="t1">'
            '<tr data-aim="r1"><td rowspan="2">A</td><td>B</td></tr>'
            '<tr data-aim="r2"><td>C</td><td>D</td></tr></table>',
            tmp_path,
        )
        texts = {c.text for row in d.tables[0].rows for c in row.cells}
        assert "D" in texts

    def test_valid_rowspan_still_merges(self, tmp_path):
        d = self._export(
            '<table data-aim-container="t1">'
            '<tr data-aim="r1"><td rowspan="2">A</td><td>B</td></tr>'
            '<tr data-aim="r2"><td>C</td></tr></table>',
            tmp_path,
        )
        table = d.tables[0]
        assert "vMerge" in table._tbl.xml
        assert table.cell(0, 0).text == table.cell(1, 0).text == "A"
        assert table.cell(0, 1).text == "B" and table.cell(1, 1).text == "C"


# ===========================================================================
# Wave 5: findings from the final Codex round on the fixed-layout-pages PR
# (docs/log/2026-07-16_1334_review_pptx-codex-final-round.md)
# ===========================================================================


class TestModifyAcceptValidatesEveryRoot:
    """A1: accept must validate EVERY root of a modify payload before the
    write — a hand-edited card whose payload smuggles a second root behind
    a valid first one lints clean, then accept corrupts the document
    (S031/S024 in the body plus a permanent H006 history mismatch)."""

    def _smuggled(self, basic_doc):
        p = basic_doc.propose_modify(
            "intro", '<p data-aim="intro">changed</p>', author=BOT, at=ts(10)
        )
        needle = '<p data-aim="intro">changed</p>'
        evil = needle + '<aim-slide data-aim="intro"><p data-aim="x9">smuggled</p></aim-slide>'
        return aim.loads(basic_doc.dumps().replace(needle, evil, 1)), p.id

    def test_accept_rejects_smuggled_second_root(self, basic_doc):
        doc, pid = self._smuggled(basic_doc)
        body_before = doc._state.serial("intro")
        with pytest.raises(InvalidOperation):
            doc.accept(pid, decided_by=ME, at=ts(11))
        # the failed accept wrote nothing: body intact, history replays,
        # and the document still round-trips
        assert doc._state.serial("intro") == body_before
        assert doc.verify() == []
        assert aim.loads(doc.dumps()).dumps() == doc.dumps()
        codes = {f.code for f in aim.lint(doc) if f.level == "error"}
        assert not codes & {"S031", "S024", "H006"}

    def test_smuggled_card_fails_lint_before_accept(self, basic_doc):
        doc, _ = self._smuggled(basic_doc)
        assert "P010" in {f.code for f in aim.lint(doc) if f.level == "error"}

    def test_multi_root_li_run_modify_still_accepts(self, rich_doc):
        p = rich_doc.propose_modify(
            "li2",
            '<li data-aim="li2">rewritten, part one…</li>'
            '<li data-aim="li2">…rewritten, part two</li>',
            author=BOT,
            at=ts(30),
        )
        rich_doc.accept(p.id, decided_by=ME, at=ts(31))
        assert "rewritten, part one" in rich_doc._state.serial("li2")
        assert rich_doc.verify() == []
        assert not [f for f in aim.lint(rich_doc) if f.level == "error"]

    def test_plain_accept_keeps_card_reserved_item_ids(self, rich_doc):
        # ids the card itself minted for new items are its own to spend:
        # the accept-time re-validation must not remap them (nor record a
        # spurious "applied" divergence for an untouched SDK card)
        p = rich_doc.propose_modify(
            "list",
            '<ul data-aim-container="list"><li data-aim="li1">First</li><li>NEW</li></ul>',
            author=BOT,
            at=ts(30),
        )
        new_id = re.findall(r'data-aim="([a-z0-9_-]+)"', p.payload_html)[-1]
        rich_doc.accept(p.id, decided_by=ME, at=ts(31))
        assert f'data-aim="{new_id}"' in rich_doc._state.serial("list")
        assert rich_doc.history[-1].get("applied") is None
        assert rich_doc.verify() == []

    def test_kind_mismatched_card_fails_lint_before_accept(self, rich_doc):
        # follow-up (Codex on the fix PR): a hand-authored card can keep the
        # target id but put it on the wrong KIND of root — accept rejects it,
        # so lint must flag the card while it is still pending
        p = rich_doc.propose_modify(
            "list",
            '<ul data-aim-container="list"><li data-aim="li1">First</li></ul>',
            author=BOT,
            at=ts(30),
        )
        needle = p.payload_html
        doc = aim.loads(rich_doc.dumps().replace(needle, '<p data-aim="list">oops</p>', 1))
        assert "P010" in {f.code for f in aim.lint(doc) if f.level == "error"}
        with pytest.raises(InvalidOperation):
            doc.accept(p.id, decided_by=ME, at=ts(31))
        assert doc.verify() == []

    def test_double_marked_card_fails_lint_before_accept(self, rich_doc):
        # follow-up (Codex, second pass): the RIGHT marker plus the WRONG one
        # on the same root read as kind-consistent, yet accept rejects the
        # wrong marker — the pending card must not lint green
        p = rich_doc.propose_modify(
            "list",
            '<ul data-aim-container="list"><li data-aim="li1">Rewritten</li></ul>',
            author=BOT,
            at=ts(30),
        )
        needle = '<ul data-aim-container="list"><li data-aim="li1">Rewritten</li></ul>'
        assert p.payload_html == needle
        evil = needle.replace(
            '<ul data-aim-container="list">', '<ul data-aim="list" data-aim-container="list">', 1
        )
        doc = aim.loads(rich_doc.dumps().replace(needle, evil, 1))
        assert "P010" in {f.code for f in aim.lint(doc) if f.level == "error"}
        with pytest.raises(InvalidOperation):
            doc.accept(p.id, decided_by=ME, at=ts(31))
        assert doc.verify() == []

    def test_wrong_marker_on_a_later_run_root_is_rejected(self, rich_doc):
        # the wrong-marker guard must scan every run root: a container marker
        # smuggled onto the SECOND li of a run would previously be written
        # into the body verbatim
        p = rich_doc.propose_modify(
            "li2",
            '<li data-aim="li2">part one</li><li data-aim="li2">part two</li>',
            author=BOT,
            at=ts(30),
        )
        needle = '<li data-aim="li2">part two</li>'
        evil = '<li data-aim="li2" data-aim-container="li2">part two</li>'
        doc = aim.loads(rich_doc.dumps().replace(needle, evil, 1))
        assert "P010" in {f.code for f in aim.lint(doc) if f.level == "error"}
        with pytest.raises(InvalidOperation):
            doc.accept(p.id, decided_by=ME, at=ts(31))
        assert doc.verify() == []
        assert 'data-aim-container="li2"' not in doc.dumps().split("<aim-proposals")[0]


class TestEmptySlidesKeepTheirDocxPage:
    """A2: a childless slide is a valid blank page (PDF prints it); the
    DOCX linearization dropped it whenever the neighbor was another slide —
    nothing emitted, and the next slide reset ``_break_before_next``."""

    @staticmethod
    def _slide(sid, *blocks):
        return (
            f'<aim-slide data-aim-container="{sid}" style="width:420px; height:595px">'
            + "".join(blocks)
            + "</aim-slide>"
        )

    BODY = '<p data-aim="{cid}" style="left:10px; top:10px; width:300px">{text}</p>'

    def _breaks(self, d):
        from docx.oxml.ns import qn

        return d.element.body.findall(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']")

    def test_leading_empty_slide_keeps_its_page(self, tmp_path):
        docx = pytest.importorskip("docx")
        doc = aim.new_document(title="Deck")
        doc.add_chunk(self._slide("s1"), author=ME, at=ts(0))
        doc.add_chunk(
            self._slide("s2", self.BODY.format(cid="b2", text="Second body.")), author=ME, at=ts(1)
        )
        assert not [f for f in aim.lint(doc) if f.level == "error"]  # blank slides are valid
        d = docx.Document(str(aim.to_docx(doc, tmp_path / "deck.docx")))
        assert len(self._breaks(d)) == 1  # blank page → s2's page
        # placeholder paragraph, page break, s2 body
        assert [p.text for p in d.paragraphs] == ["", "", "Second body."]

    def test_mid_deck_empty_slide_keeps_its_page(self, tmp_path):
        docx = pytest.importorskip("docx")
        doc = aim.new_document(title="Deck")
        doc.add_chunk(
            self._slide("s1", self.BODY.format(cid="b1", text="First body.")), author=ME, at=ts(0)
        )
        doc.add_chunk(self._slide("s2"), author=ME, at=ts(1))
        doc.add_chunk(
            self._slide("s3", self.BODY.format(cid="b3", text="Third body.")), author=ME, at=ts(2)
        )
        d = docx.Document(str(aim.to_docx(doc, tmp_path / "deck.docx")))
        assert len(self._breaks(d)) == 2  # s1 → blank page → s3
        assert [p.text for p in d.paragraphs] == ["First body.", "", "", "", "Third body."]

    def test_only_empty_slide_exports_one_blank_page(self, tmp_path):
        docx = pytest.importorskip("docx")
        doc = aim.new_document(title="Deck")
        doc.add_chunk(self._slide("s1"), author=ME, at=ts(0))
        d = docx.Document(str(aim.to_docx(doc, tmp_path / "deck.docx")))
        assert len(d.paragraphs) == 1 and not self._breaks(d)


class TestSlideRootsForceTheContainerMarker:
    """A4: a caller-supplied valid unused id was honored on whatever marker
    it arrived on — so ``add_chunk('<aim-slide data-aim="sx">…')`` minted an
    S031-failing document, and the propose_add card lint-passed while
    pending only to fail after acceptance."""

    SLIDE = (
        '<aim-slide data-aim="sx" style="width:420px; height:595px">'
        '<h2 style="left:10px; top:10px; width:300px">X</h2></aim-slide>'
    )

    def test_add_chunk_moves_the_id_to_the_container_marker(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(self.SLIDE, author=ME, at=ts(0))
        assert "sx" in doc.containers
        text = doc.dumps()
        assert 'data-aim-container="sx"' in text
        assert 'data-aim="sx"' not in text
        # the slide's children are addressable chunks, not swallowed content
        slide = doc._state.container_node("sx")
        assert all(e.chunk_id for e in slide.elements())
        assert not [f for f in aim.lint(doc) if f.level == "error"]

    def test_propose_add_then_accept_ends_lint_clean(self):
        doc = aim.new_document(title="T")
        p = doc.propose_add(self.SLIDE, author=BOT, at=ts(0))
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        doc.accept(p.id, decided_by=ME, at=ts(1))
        assert "sx" in doc.containers
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        assert doc.verify() == []

    def test_reminted_slide_id_drops_the_stale_chunk_marker(self):
        # follow-up (Codex, third pass): when the supplied id is already
        # taken, the fresh-id branch re-mints — but must also REMOVE the
        # stale data-aim, or the card carries both markers and accepting
        # writes an S012/S031 document
        doc = aim.new_document(title="T")
        doc.add_chunk(self.SLIDE, author=ME, at=ts(0))
        assert "sx" in doc.containers  # the id is now taken
        p = doc.propose_add(self.SLIDE, author=BOT, at=ts(1))
        assert 'data-aim="' not in p.payload_html.split(">")[0]  # root carries no chunk marker
        assert "data-aim-container=" in p.payload_html.split(">")[0]
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        doc.accept(p.id, decided_by=ME, at=ts(2))
        assert len([c for c in doc.containers if c != "sx"]) >= 1
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        assert doc.verify() == []


class TestDefaultCanvasSlidesGetPrintSize:
    """A5: a slide that omits its canvas size got a 960×540 named page but
    an element rule with only page+zoom — the slide box collapsed (children
    are absolutely positioned, overflow hidden) and printed blank. The
    element rule must carry the same resolved geometry as its @page rule."""

    def _css(self, style):
        from aimformat.convert._pdf_out import _slide_page_css

        doc = aim.new_document(title="T")
        s = f' style="{style}"' if style else ""
        doc.add_chunk(
            f'<aim-slide data-aim-container="s1"{s}>'
            '<p data-aim="b1" style="left:10px; top:10px; width:300px">x</p>'
            "</aim-slide>",
            author=ME,
            at=ts(0),
        )
        css = _slide_page_css(doc)
        rule = re.search(r'aim-slide\[data-aim-container="s1"\]\{([^}]*)\}', css)
        return css, rule.group(1)

    def test_unsized_slide_rule_carries_the_default_box(self):
        css, rule = self._css(None)
        assert "@page pg-s1{size:960pt 540pt;margin:0}" in css
        assert "width:960px" in rule and "height:540px" in rule

    def test_partial_size_defaults_only_the_missing_axis(self):
        css, rule = self._css("height:595px")
        assert "@page pg-s1{size:960pt 595pt;margin:0}" in css
        assert "width:960px" in rule and "height:595px" in rule


def _mini(construct: str) -> str:
    doc = aim.new_document(title="mini")
    text = doc.dumps()
    return text.replace("<body>", "<body>\n" + construct)


# ===========================================================================
# Wave 6: findings from the 2026-07 deep review (consolidated ids AF-xx)
# ===========================================================================


class TestMoveStaysOutOfItsOwnSubtree:
    """AF-01: a move whose destination is the moved construct itself or any
    of its descendants used to remove the subtree, then fail the re-insert —
    mutation with no event, silent data loss, verify() broken."""

    @pytest.fixture
    def nested_doc(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">before</p>', author=BOT, at=ts(0))
        doc.add_chunk(
            '<aim-slide data-aim-container="s1">'
            '<ul data-aim-container="l2"><li data-aim="i1">x</li></ul>'
            "</aim-slide>",
            author=BOT,
            at=ts(1),
        )
        return doc

    def test_move_container_into_itself_fails_closed(self, nested_doc):
        before = nested_doc.dumps()
        with pytest.raises(InvalidOperation):
            nested_doc.move_chunk("s1", author=ME, container="s1", at=ts(2))
        assert nested_doc.dumps() == before
        assert nested_doc.verify() == []

    def test_move_container_into_own_descendant_fails_closed(self, nested_doc):
        before = nested_doc.dumps()
        with pytest.raises(InvalidOperation):
            nested_doc.move_chunk("s1", author=ME, container="l2", at=ts(2))
        assert nested_doc.dumps() == before
        assert nested_doc.verify() == []

    def test_accepting_a_self_move_proposal_fails_closed(self, nested_doc):
        # an aim-slide destination passes the item-carrier guard (slides
        # take any member), so this reaches the accept-time self-move check
        p = nested_doc.propose_move("s1", author=BOT, container="s1", at=ts(2))
        before = nested_doc.dumps()
        with pytest.raises(InvalidOperation):
            nested_doc.accept(p.id, decided_by=ME, at=ts(3))
        # fail closed: nothing mutated, the card is still pending
        assert nested_doc.dumps() == before
        assert [q.id for q in nested_doc.proposals] == [p.id]
        assert nested_doc.verify() == []

    def test_low_level_move_rejects_descendant_after_anchor(self, nested_doc):
        from aimformat.document import Anchor

        with pytest.raises(InvalidOperation):
            nested_doc._state.move("s1", Anchor("l2", "i1"))
        assert nested_doc.verify() == []

    def test_legit_move_still_works(self, nested_doc):
        nested_doc.move_chunk("s1", author=ME, container="body", after=None, at=ts(2))
        assert nested_doc.body_ids == ["s1", "p1"]
        assert nested_doc.verify() == []


class TestAcceptedAddValidatesItsPayload:
    """AF-02: the accepted-add branch of _resolve inserted the card payload
    raw — no root marker, id validity/uniqueness, run-shape, or nested-id
    checks — while the modify branch fully re-validates. A hand-authored
    card could smuggle an unmarked root (unreplayable target "") or a live
    id (duplicate body ids) past accept."""

    def _with_hand_card(self, payload, pid="p-handmade"):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="live">real</p>', author=BOT, at=ts(0))
        card = (
            f'<aim-proposal id="{pid}" data-action="add" '
            'data-anchor-container="body" data-anchor-after="live" '
            f'data-at="{ts(5)}" data-author="human" data-author-id="eve">'
            f"<template>{payload}</template></aim-proposal>"
        )
        text = doc.dumps().replace("</body>", f"<aim-proposals>\n{card}\n</aim-proposals>\n</body>")
        return aim.loads(text)

    def test_unmarked_root_fails_closed(self):
        doc = self._with_hand_card("<p>smuggled</p>")
        with pytest.raises(InvalidOperation):
            doc.accept("p-handmade", decided_by=ME, at=ts(6))
        assert doc.body_ids == ["live"]
        assert doc.verify() == []

    def test_live_id_collision_fails_closed(self):
        doc = self._with_hand_card('<p data-aim="live">impostor</p>')
        with pytest.raises(InvalidOperation):
            doc.accept("p-handmade", decided_by=ME, at=ts(6))
        assert doc.body_ids == ["live"]
        assert doc.verify() == []

    def test_nested_live_id_is_reminted_not_duplicated(self):
        doc = self._with_hand_card('<ul data-aim-container="c9"><li data-aim="live">x</li></ul>')
        doc.accept("p-handmade", decided_by=ME, at=ts(6))
        all_ids = doc._state.all_ids()
        assert doc.body_ids == ["live", "c9"]
        assert len([i for i in all_ids if i == "live"]) == 1
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        assert doc.verify() == []

    def test_valid_hand_card_still_accepts(self):
        doc = self._with_hand_card('<p data-aim="fresh1">fine</p>')
        doc.accept("p-handmade", decided_by=ME, at=ts(6))
        assert doc.body_ids == ["live", "fresh1"]
        assert doc.verify() == []

    def test_normal_add_accept_records_no_spurious_applied(self, basic_doc):
        p = basic_doc.propose_add('<p data-aim="new1">hello</p>', author=BOT, at=ts(10))
        ev = basic_doc.accept(p.id, decided_by=ME, at=ts(11))
        assert ev.get("applied") is None
        assert basic_doc.verify() == []


class TestContainersHoldOnlyItemCarriers:
    """AF-03: a non-item chunk (e.g. <p data-aim>) inside a list/table
    container linted clean — S022 fired only for item carriers in the WRONG
    container kind — and every write path (add/move/modify/accept) let one
    in. Item-aware consumers can't see such a member: an editor hides it,
    then the next container-level write destroys it."""

    def test_p_chunk_inside_ul_is_S022(self, rich_doc):
        text = rich_doc.dumps().replace(
            '<li data-aim="li1">First</li>',
            '<li data-aim="li1">First</li><p data-aim="px">hidden</p>',
        )
        assert "S022" in {f.code for f in aim.lint_text(text)}

    def test_wrong_kind_item_carrier_still_fires_S022(self, rich_doc):
        text = rich_doc.dumps().replace(
            '<li data-aim="li1">First</li>',
            '<li data-aim="li1">First</li><tr data-aim="trx"><td>x</td></tr>',
        )
        assert "S022" in {f.code for f in aim.lint_text(text)}

    def test_add_chunk_rejects_non_item_member(self, rich_doc):
        with pytest.raises(InvalidOperation):
            rich_doc.add_chunk('<p data-aim="px">bad</p>', author=ME, container="list", at=ts(20))
        assert rich_doc.verify() == []

    def test_move_chunk_rejects_non_item_member(self, rich_doc):
        before = rich_doc.dumps()
        with pytest.raises(InvalidOperation):
            rich_doc.move_chunk("intro", author=ME, container="list", at=ts(20))
        assert rich_doc.dumps() == before
        assert rich_doc.verify() == []

    def test_modify_cannot_swap_item_for_non_item(self, rich_doc):
        with pytest.raises(InvalidOperation):
            rich_doc.modify_chunk("li1", '<p data-aim="li1">now a para</p>', author=ME, at=ts(20))
        assert rich_doc.verify() == []

    def test_propose_paths_reject_at_creation(self, rich_doc):
        with pytest.raises(InvalidOperation):
            rich_doc.propose_add(
                '<p data-aim="px">bad</p>', author=BOT, container="list", at=ts(20)
            )
        with pytest.raises(InvalidOperation):
            rich_doc.propose_move("intro", author=BOT, container="list", at=ts(21))

    def test_hand_authored_add_card_fails_closed_at_accept(self, rich_doc):
        card = (
            '<aim-proposal id="p-handmade" data-action="add" '
            'data-anchor-container="list" data-anchor-after="li1" '
            f'data-at="{ts(20)}" data-author="human" data-author-id="eve">'
            '<template><p data-aim="px">bad</p></template></aim-proposal>'
        )
        text = rich_doc.dumps().replace("<aim-proposals>", f"<aim-proposals>\n{card}", 1)
        if "<aim-proposals>" not in text:
            text = rich_doc.dumps().replace(
                "</body>", f"<aim-proposals>\n{card}\n</aim-proposals>\n</body>"
            )
        doc = aim.loads(text)
        with pytest.raises(InvalidOperation):
            doc.accept("p-handmade", decided_by=ME, at=ts(21))
        assert doc.verify() == []

    def test_legal_item_writes_still_work(self, rich_doc):
        rich_doc.add_chunk(
            '<li data-aim="li9">new item</li>', author=ME, container="list", at=ts(20)
        )
        rich_doc.add_chunk(
            '<tr data-aim="row9"><td>gamma</td><td>3</td></tr>',
            author=ME,
            container="tbl",
            at=ts(21),
        )
        rich_doc.move_chunk("li9", author=ME, container="list", after=None, at=ts(22))
        assert rich_doc.verify() == []
        assert not [f for f in aim.lint(rich_doc) if f.level == "error"]


class TestContainerPayloadsValidateTheirMembers:
    """Codex-1: body-level container payloads bypassed the destination
    container guard, so their own direct members could violate S022."""

    def test_add_rejects_invalid_list_container_payload(self):
        doc = aim.new_document(title="T")
        before = doc.dumps()

        with pytest.raises(InvalidOperation):
            doc.add_chunk(
                '<ul data-aim-container="listx"><p data-aim="p1">bad</p></ul>',
                author=ME,
                at=ts(0),
            )

        assert doc.dumps() == before
        assert doc.verify() == []

    def test_modify_rejects_invalid_table_container_payload(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tablex"><tbody>'
            '<tr data-aim="row1"><td>good</td></tr>'
            "</tbody></table>",
            author=ME,
            at=ts(0),
        )
        before = doc.dumps()

        with pytest.raises(InvalidOperation):
            doc.modify_chunk(
                "tablex",
                '<table data-aim-container="tablex"><tbody>'
                '<p data-aim="p1">bad</p>'
                "</tbody></table>",
                author=ME,
                at=ts(1),
            )

        assert doc.dumps() == before
        assert doc.verify() == []


class TestShellChildrenAreItemCarriers:
    """codex-r2-3 (refines codex-1/AF-03): the lint carrier check skipped
    children nested under thead/tbody/tfoot — the shell branch only verified
    each row HAS a chunk id, so ``<tbody><p data-aim="r">x</p></tbody>``
    linted clean and item-aware consumers could accept a table whose row is
    invisible to them."""

    ROW = '<tr data-aim="r1"><td>ok</td></tr>'

    def _table_text(self, row_html):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            f'<table data-aim-container="tbl"><tbody>{self.ROW}</tbody></table>',
            author=ME,
            at=ts(0),
        )
        return doc.dumps().replace(self.ROW, row_html)

    def test_non_carrier_chunk_inside_tbody_is_S022(self):
        text = self._table_text('<p data-aim="r1">x</p>')
        assert "S022" in {f.code for f in aim.lint_text(text)}

    def test_wrong_kind_carrier_inside_tbody_is_S022(self):
        text = self._table_text('<li data-aim="r1">x</li>')
        assert "S022" in {f.code for f in aim.lint_text(text)}

    def test_uncovered_shell_child_still_fires_S021_not_S022(self):
        codes = {f.code for f in aim.lint_text(self._table_text("<tr><td>bare</td></tr>"))}
        assert "S021" in codes and "S022" not in codes

    def test_legal_rows_stay_clean(self):
        assert not [f for f in aim.lint_text(self._table_text(self.ROW)) if f.level == "error"]


class TestResolutionOrderProtectsAnchors:
    """AF-04: resolution_order sorted ready cards only delete-last, so a
    pending MOVE of an existing chunk that a sibling add anchors on could
    resolve first — the add then failed TargetNotFound, aborting
    `accept --all` and `to_docx(pending="accept-all")` on card order alone."""

    @pytest.fixture
    def doc_with_lane(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="x1">anchor chunk</p>', author=BOT, at=ts(0))
        doc.add_chunk(
            '<aim-slide data-aim-container="s1"><p data-aim="sp">slide p</p></aim-slide>',
            author=BOT,
            at=ts(1),
        )
        # move card BEFORE the add that anchors on its target
        doc.propose_move("x1", author=BOT, container="s1", at=ts(2))
        doc.propose_add(
            '<p data-aim="n1">new</p>', author=BOT, container="body", after="x1", at=ts(3)
        )
        return doc

    def test_add_orders_before_the_move_of_its_anchor(self, doc_with_lane):
        from aimformat.document import resolution_order

        actions = [p.action for p in resolution_order(doc_with_lane.proposals)]
        assert actions == ["add", "move"]

    def test_accept_all_resolves_the_whole_lane(self, doc_with_lane):
        from aimformat.document import resolution_order

        for p in resolution_order(doc_with_lane.proposals):
            doc_with_lane.accept(p.id, decided_by=ME, at=ts(4))
        assert doc_with_lane.proposals == []
        assert doc_with_lane.body_ids == ["n1", "s1"]
        assert doc_with_lane.verify() == []

    def test_docx_accept_all_export_succeeds(self, doc_with_lane, tmp_path):
        pytest.importorskip("docx")
        out = aim.to_docx(doc_with_lane, tmp_path / "t.docx", pending="accept-all")
        assert out.exists()

    def test_delete_still_goes_after_the_anchored_add(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="x1">anchor</p>', author=BOT, at=ts(0))
        doc.propose_delete("x1", author=BOT, at=ts(1))
        doc.propose_add(
            '<p data-aim="n1">new</p>', author=BOT, container="body", after="x1", at=ts(2)
        )
        from aimformat.document import resolution_order

        for p in resolution_order(doc.proposals):
            doc.accept(p.id, decided_by=ME, at=ts(3))
        assert doc.body_ids == ["n1"]
        assert doc.verify() == []


class TestMovesStayAheadOfAncestorReplacements:
    """codex-r2-5 (refines AF-04): ranking every modify below every move
    broke lanes where a move rescues a descendant before its container is
    replaced — ``accept --all`` resolved the container modify first,
    deleting the member, and the move then raised TargetNotFound after
    partially resolving the document. A container modify whose payload
    drops a pending move's target now waits for that move; a payload that
    keeps the member stays ahead of the move (relocating first would make
    the payload re-introduce a duplicate id)."""

    @pytest.fixture
    def rescue_lane(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<ul data-aim-container="lst"><li data-aim="x">keep me</li>'
            '<li data-aim="y">other</li></ul>',
            author=ME,
            at=ts(0),
        )
        # move card FIRST rescues x; the container replacement drops it
        doc.propose_move("x", author=BOT, container="body", at=ts(1))
        doc.propose_modify(
            "lst",
            '<table data-aim-container="lst"><tbody>'
            '<tr data-aim="r1"><td>other</td></tr></tbody></table>',
            author=BOT,
            at=ts(2),
        )
        return doc

    def test_move_orders_before_the_container_replacement(self, rescue_lane):
        from aimformat.document import resolution_order

        actions = [p.action for p in resolution_order(rescue_lane.proposals, rescue_lane)]
        assert actions == ["move", "modify"]

    def test_accept_all_resolves_the_whole_lane(self, rescue_lane):
        from aimformat.document import resolution_order

        for p in resolution_order(rescue_lane.proposals, rescue_lane):
            rescue_lane.accept(p.id, decided_by=ME, at=ts(3))
        assert rescue_lane.proposals == []
        assert "x" in rescue_lane.body_ids
        assert rescue_lane.verify() == []

    def test_docx_accept_all_export_succeeds(self, rescue_lane, tmp_path):
        pytest.importorskip("docx")
        out = aim.to_docx(rescue_lane, tmp_path / "t.docx", pending="accept-all")
        assert out.exists()

    def test_payload_keeping_the_member_stays_ahead_of_the_move(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<ul data-aim-container="lst"><li data-aim="x">keep</li></ul>', author=ME, at=ts(0)
        )
        doc.propose_move("x", author=BOT, container="body", at=ts(1))
        doc.propose_modify(
            "lst",
            '<ul data-aim-container="lst"><li data-aim="x">keep!</li></ul>',
            author=BOT,
            at=ts(2),
        )
        from aimformat.document import resolution_order

        assert [p.action for p in resolution_order(doc.proposals, doc)] == ["modify", "move"]
        for p in resolution_order(doc.proposals, doc):
            doc.accept(p.id, decided_by=ME, at=ts(3))
        assert doc.body_ids == ["lst", "x"]
        assert doc.verify() == []

    def test_without_the_doc_the_static_order_is_unchanged(self, rescue_lane):
        from aimformat.document import resolution_order

        actions = [p.action for p in resolution_order(rescue_lane.proposals)]
        assert actions == ["modify", "move"]


class TestPruneFlattenKeepBurnedIds:
    """AF-05: prune()/flatten() dropped the only record of burned ids, so a
    later write re-honored a previously seen id — an external reference to
    the old id then aliases unrelated content while verify() stays clean
    (spec §4.4: an id is never reused within a document lifetime)."""

    def _doc(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="victim">to be deleted</p>', author=BOT, at=ts(0))
        doc.add_chunk('<p data-aim="keeper">stays</p>', author=BOT, at=ts(1))
        doc.delete_chunk("victim", author=ME, at=ts(2))
        doc.checkpoint("cut", at=ts(3))
        doc.modify_chunk("keeper", '<p data-aim="keeper">stays!</p>', author=ME, at=ts(4))
        return doc

    def test_pruned_burned_id_is_not_rehonored(self):
        doc = self._doc()
        doc.prune(before="cut")
        # the burn record left the retained log…
        assert all(e.get("target") != "victim" for e in doc.history)
        # …but the id must still mint fresh, not alias the old reference
        c = doc.add_chunk('<p data-aim="victim">impostor</p>', author=BOT, at=ts(5))
        assert c.id != "victim"
        assert doc.verify() == []

    def test_flattened_burned_id_is_not_rehonored(self):
        doc = self._doc()
        doc.flatten()
        c = doc.add_chunk('<p data-aim="victim">impostor</p>', author=BOT, at=ts(5))
        assert c.id != "victim"
        assert doc.verify() == []

    def test_pending_card_still_accepts_after_prune(self):
        doc = self._doc()
        p = doc.propose_add('<p data-aim="pnew">proposed</p>', author=BOT, at=ts(5))
        doc.prune(before="cut")
        doc.accept(p.id, decided_by=ME, at=ts(6))
        assert "pnew" in doc.body_ids


class TestPruneCliResolvesNumericCheckpointLabels:
    """codex-r2-4: ``aim prune`` documents BEFORE as a seq number or
    checkpoint label, but an unconditional digit conversion made numeric
    labels unselectable — with a checkpoint labeled "1" at a later seq,
    ``aim prune f.aim 1`` kept from seq 1 and silently retained the history
    the user asked to drop. An exact label match now wins."""

    def _saved(self, tmp_path):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="a">one</p>', author=ME, at=ts(0))
        doc.add_chunk('<p data-aim="b">two</p>', author=ME, at=ts(1))
        doc.checkpoint("1", at=ts(2))  # numeric label, seq 3
        doc.modify_chunk("b", '<p data-aim="b">two!</p>', author=ME, at=ts(3))
        path = tmp_path / "doc.aim"
        doc.save(path)
        return path

    def test_numeric_label_selects_the_checkpoint(self, tmp_path):
        from aimformat.cli import main

        path = self._saved(tmp_path)
        assert main(["prune", str(path), "1"]) == 0
        kept = aim.load(path).history
        assert kept[0].kind == "checkpoint" and kept[0].get("label") == "1"
        assert kept[0].seq == 3  # cut at the checkpoint, not at seq 1

    def test_plain_seq_still_works_when_no_label_matches(self, tmp_path):
        from aimformat.cli import main

        path = self._saved(tmp_path)
        assert main(["prune", str(path), "2"]) == 0  # no checkpoint labeled "2"
        assert [e.seq for e in aim.load(path).history][0] == 2


class TestReconcileHandlesWrappingContainers:
    """AF-07: a hand edit that wraps existing units into a NEW container is
    one of the most natural out-of-band edits, but _drive counted a
    container's interior as covered only when the container existed on both
    sides — the A-only wrapper was added whole while the units' E-side
    copies survived as duplicates, and reconcile raised 'did not converge'
    (failing closed on a legitimate repair of a headline feature, §6.8)."""

    def _base_text(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<h1 data-aim="h1">Title</h1>', author=BOT, at=ts(0))
        doc.add_chunk('<p data-aim="c1">alpha</p>', author=BOT, at=ts(1))
        doc.add_chunk('<p data-aim="c2">beta</p>', author=BOT, at=ts(2))
        return doc.dumps()

    def _reconciled(self, old, new):
        text = self._base_text()
        assert old in text
        doc = aim.loads(text.replace(old, new))
        report = doc.reconcile(at=ts(3))
        assert report.residual == []
        assert doc.verify() == []
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        return doc, report

    def test_slide_wrapping_existing_chunks_reconciles(self):
        doc, report = self._reconciled(
            '<p data-aim="c1">alpha</p>\n<p data-aim="c2">beta</p>',
            '<aim-slide><p data-aim="c1">alpha</p><p data-aim="c2">beta</p></aim-slide>',
        )
        assert [e.get("action") for e in report.events] == ["delete", "delete", "add"]
        assert doc._state.find_chunk("c1")[1]  # the unit lives inside the wrapper

    def test_list_wrapping_converted_items_reconciles(self):
        doc, _ = self._reconciled(
            '<p data-aim="c1">alpha</p>\n<p data-aim="c2">beta</p>',
            '<ul data-aim-container="wrap">'
            '<li data-aim="c1">alpha</li><li data-aim="c2">beta</li></ul>',
        )
        assert doc.body_ids == ["h1", "wrap"]

    def test_wrapper_swallowing_a_whole_container_reconciles(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<ul data-aim-container="l1"><li data-aim="i1">x</li></ul>', author=BOT, at=ts(0)
        )
        text = doc.dumps()
        wrapped = text.replace(
            '<ul data-aim-container="l1"><li data-aim="i1">x</li></ul>',
            '<aim-slide data-aim-container="ws">'
            '<ul data-aim-container="l1"><li data-aim="i1">x</li></ul></aim-slide>',
        )
        d = aim.loads(wrapped)
        report = d.reconcile(at=ts(3))
        assert report.residual == []
        assert d.verify() == []
        assert d.body_ids == ["ws"]


class TestEveryAimCssBlockIsVerified:
    """AF-21: the X006 byte-match ran only when data-aim-css equalled the
    CURRENT spec version, and only on the first head block — a
    version-mismatched <style data-aim-css="0.1"> (or a second attr-bearing
    block) smuggled arbitrary CSS (@import beacons, phishing overlays) past
    aim lint with zero errors."""

    def _text(self, basic_doc):
        return basic_doc.dumps()

    def test_stale_version_css_is_X006(self, basic_doc):
        text = re.sub(
            r'<style data-aim-css="[^"]*">.*?</style>',
            '<style data-aim-css="0.1">@import url(https://evil.example/x.css);</style>',
            self._text(basic_doc),
            flags=re.S,
        )
        assert "X006" in {f.code for f in aim.lint_text(text)}

    def test_second_attr_bearing_head_block_is_X006(self, basic_doc):
        text = self._text(basic_doc).replace(
            "</title>",
            '</title>\n<style data-aim-css="0.1">body{display:none}</style>',
        )
        assert "X006" in {f.code for f in aim.lint_text(text)}

    def test_inert_stale_placeholder_stays_warning_only(self, basic_doc):
        text = re.sub(
            r'<style data-aim-css="[^"]*">.*?</style>',
            '<style data-aim-css="0.1">/* aim.css placeholder */</style>',
            self._text(basic_doc),
            flags=re.S,
        )
        codes = {f.code: f.level for f in aim.lint_text(text)}
        assert "X006" not in codes and "X005" not in codes
        assert codes.get("S006") == "warning"

    def test_pristine_document_still_lints_clean(self, basic_doc):
        assert aim.lint_text(self._text(basic_doc)) == []


class TestTrackedTableContainerRevisions:
    """AF-35: ``emit_construct`` called ``emit_table(el, force="del")``
    without ``prop=``, so a pending-deleted (or container-modified) table
    exported as plain untracked content — Word showed the doomed table as
    ordinary accepted text, and the modify case kept BOTH grids on
    accept-all."""

    @staticmethod
    def _del_texts(path):
        import docx
        from docx.oxml.ns import qn

        d = docx.Document(str(path))
        return [t.text for t in d.element.body.iter(qn("w:delText"))]

    @staticmethod
    def _ins_texts(path):
        import docx
        from docx.oxml.ns import qn

        d = docx.Document(str(path))
        return [t.text for w in d.element.body.iter(qn("w:ins")) for t in w.iter(qn("w:t"))]

    def test_pending_delete_of_table_container_is_tracked(self, table_doc, tmp_path):
        pytest.importorskip("docx")
        table_doc.propose_delete("tbl", author=BOT, at=ts(1))
        out = aim.to_docx(table_doc, tmp_path / "d.docx")
        deleted = self._del_texts(out)
        for text in ("K", "1", "2"):
            assert text in deleted

    def test_container_modify_tracks_old_grid_as_deleted(self, table_doc, tmp_path):
        pytest.importorskip("docx")
        table_doc.propose_modify(
            "tbl",
            '<table data-aim-container="tbl"><tbody>'
            '<tr data-aim="n1"><td>only</td></tr></tbody></table>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(table_doc, tmp_path / "m.docx")
        deleted, inserted = self._del_texts(out), self._ins_texts(out)
        for text in ("K", "1", "2"):
            assert text in deleted
        assert "only" in inserted


class TestMarkdownKeepsGroupingBlockText:
    """AF-36: the div/section and blockquote branches of the Markdown
    exporter iterated only ``elements()``, so lint-clean direct text of a
    grouping block (``<blockquote>Direct quote</blockquote>``) exported as
    nothing — permanent loss on reimport."""

    def test_blockquote_direct_text_survives(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<blockquote data-aim="q">Direct quote</blockquote>', author=ME, at=ts(0))
        assert "> Direct quote" in aim.to_markdown(doc)

    def test_div_mixed_content_keeps_order(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<div data-aim="d">Lead-in text<p>First para.</p>Trailing text</div>',
            author=ME,
            at=ts(0),
        )
        md = aim.to_markdown(doc)
        lead = md.index("Lead-in text")
        para = md.index("First para.")
        trail = md.index("Trailing text")
        assert lead < para < trail

    def test_section_direct_text_survives(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<section data-aim="s"><h2>Head</h2>Loose section text</section>',
            author=ME,
            at=ts(0),
        )
        md = aim.to_markdown(doc)
        assert "## Head" in md and "Loose section text" in md

    def test_whitespace_between_blocks_adds_no_paragraph(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<div data-aim="d">\n  <p>One.</p>\n  <p>Two.</p>\n</div>', author=ME, at=ts(0)
        )
        md = aim.to_markdown(doc)
        assert "One.\n\nTwo." in md


class TestDocxKeepsInlineImages:
    """AF-37: ``_runs_of`` recursed into an ``<img>``'s (nonexistent)
    children, so ``Before <img> after`` exported as ``Before  after`` —
    no alt, no URL, anywhere; the supported md→AIM→DOCX path silently
    dropped every inline image."""

    def _texts(self, path):
        import docx

        return [p.text for p in docx.Document(str(path)).paragraphs]

    def test_url_image_leaves_placeholder_and_url(self, tmp_path):
        pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<p data-aim="p1">Before <img src="https://x.example/i.png" alt="chart"> after</p>',
            author=ME,
            at=ts(0),
        )
        texts = self._texts(aim.to_docx(doc, tmp_path / "i.docx"))
        assert "Before [image: chart] (https://x.example/i.png) after" in texts

    def test_data_image_placeholder_elides_the_blob(self, tmp_path):
        pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<p data-aim="p1">See <img src="data:image/png;base64,iVBORw0KGgo=" alt="logo">.</p>',
            author=ME,
            at=ts(0),
        )
        texts = self._texts(aim.to_docx(doc, tmp_path / "d.docx"))
        assert "See [image: logo]." in texts
        assert not any("data:" in t for t in texts)

    def test_tracked_delete_keeps_the_placeholder(self, tmp_path):
        pytest.importorskip("docx")
        import docx
        from docx.oxml.ns import qn

        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<p data-aim="p1">Before <img src="https://x.example/i.png" alt="chart"> after</p>',
            author=ME,
            at=ts(0),
        )
        doc.propose_delete("p1", author=BOT, at=ts(1))
        out = aim.to_docx(doc, tmp_path / "t.docx")
        deleted = "".join(
            t.text or "" for t in docx.Document(str(out)).element.body.iter(qn("w:delText"))
        )
        assert "[image: chart]" in deleted


class TestCriticMarkupKeepsHeaderAnchoredRowAdds:
    """AF-38: the header loop of the Markdown table renderer never drained
    ``adds_by_anchor``, so a pending ``<tr>`` anchored after a thead row —
    and every add chained on it — vanished from criticmarkup output."""

    def test_row_add_after_header_is_rendered(self, table_doc):
        table_doc.propose_add(
            '<tr data-aim="nr"><td>9</td></tr>',
            author=BOT,
            container="tbl",
            after="h",
            at=ts(1),
        )
        md = aim.to_markdown(table_doc, pending="criticmarkup")
        assert "{++" in md and "9" in md

    def test_chained_add_on_the_header_anchor_is_rendered(self, table_doc):
        first = table_doc.propose_add(
            '<tr data-aim="nr"><td>9</td></tr>',
            author=BOT,
            container="tbl",
            after="h",
            at=ts(1),
        )
        table_doc.propose_add(
            '<tr data-aim="nr2"><td>10</td></tr>',
            author=BOT,
            container="tbl",
            after=first.id,
            at=ts(2),
        )
        md = aim.to_markdown(table_doc, pending="criticmarkup")
        assert "9" in md and "10" in md

    def test_add_lands_after_the_separator(self, table_doc):
        table_doc.propose_add(
            '<tr data-aim="nr"><td>9</td></tr>',
            author=BOT,
            container="tbl",
            after="h",
            at=ts(1),
        )
        lines = aim.to_markdown(table_doc, pending="criticmarkup").splitlines()
        sep = next(i for i, ln in enumerate(lines) if ln.startswith("| ---"))
        added = next(i for i, ln in enumerate(lines) if "{++" in ln)
        assert added == sep + 1


class TestTrackedStructuralRowModify:
    """AF-39: a row modify that changes the grid shape was forced into the
    old row's cells — surplus replacement cells fused into the last cell
    (``A|B`` → ``X|Y|Z`` exported as ``X | YZ``), a grid no Word
    accept/reject sequence could turn into the AIM state. Shape-changing
    modifies now emit a tracked row-delete plus inserted replacement rows."""

    @pytest.fixture
    def two_col_doc(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl">'
            '<thead><tr data-aim="h"><th>K</th><th>V</th></tr></thead>'
            '<tbody><tr data-aim="r1"><td>A</td><td>B</td></tr></tbody></table>',
            author=ME,
            at=ts(0),
        )
        return doc

    @staticmethod
    def _ins_cells_per_row(path):
        import docx
        from docx.oxml.ns import qn

        d = docx.Document(str(path))
        out = []
        for tbl in d.tables:
            for tr in tbl._tbl.findall(qn("w:tr")):
                out.append(
                    [
                        "".join(
                            t.text or "" for w in tc.iter(qn("w:ins")) for t in w.iter(qn("w:t"))
                        )
                        for tc in tr.findall(qn("w:tc"))
                    ]
                )
        return out

    def test_wider_replacement_row_keeps_cell_boundaries(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        import docx
        from docx.oxml.ns import qn

        two_col_doc.propose_modify(
            "r1",
            '<tr data-aim="r1"><td>X</td><td>Y</td><td>Z</td></tr>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(two_col_doc, tmp_path / "s.docx")
        ins_rows = self._ins_cells_per_row(out)
        assert ["X", "Y", "Z"] in ins_rows  # one inserted row, one cell each
        assert not any("YZ" in c for row in ins_rows for c in row)  # no fusion
        deleted = [t.text for t in docx.Document(str(out)).element.body.iter(qn("w:delText"))]
        assert "A" in deleted and "B" in deleted

    def test_wider_replacement_does_not_widen_original_rows(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        import docx
        from docx.oxml.ns import qn

        two_col_doc.propose_modify(
            "r1",
            '<tr data-aim="r1"><td>X</td><td>Y</td><td>Z</td></tr>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(two_col_doc, tmp_path / "width.docx")
        rows = docx.Document(str(out)).tables[0]._tbl.findall(qn("w:tr"))
        header = next(row for row in rows if "K" in [t.text for t in row.iter(qn("w:t"))])
        old_row = next(row for row in rows if "A" in [t.text for t in row.iter(qn("w:delText"))])
        new_row = next(row for row in rows if "X" in [t.text for t in row.iter(qn("w:t"))])

        assert len(header.findall(qn("w:tc"))) == 2
        assert len(old_row.findall(qn("w:tc"))) == 2
        assert len(new_row.findall(qn("w:tc"))) == 3

    def test_shape_changing_rows_carry_structural_revisions(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        import docx
        from docx.oxml.ns import qn

        two_col_doc.propose_modify(
            "r1",
            '<tr data-aim="r1"><td>X</td><td>Y</td><td>Z</td></tr>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(two_col_doc, tmp_path / "structural.docx")
        rows = docx.Document(str(out)).tables[0]._tbl.findall(qn("w:tr"))
        old_row = next(row for row in rows if "A" in [t.text for t in row.iter(qn("w:delText"))])
        new_row = next(row for row in rows if "X" in [t.text for t in row.iter(qn("w:t"))])

        old_props = old_row.find(qn("w:trPr"))
        new_props = new_row.find(qn("w:trPr"))
        assert old_props is not None and old_props.find(qn("w:del")) is not None
        assert new_props is not None and new_props.find(qn("w:ins")) is not None

    def test_same_shape_modify_stays_cellwise(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        two_col_doc.propose_modify(
            "r1", '<tr data-aim="r1"><td>X</td><td>Y</td></tr>', author=BOT, at=ts(1)
        )
        out = aim.to_docx(two_col_doc, tmp_path / "c.docx")
        ins_rows = self._ins_cells_per_row(out)
        # replacement lands inside the original row's cells, not a new row
        assert len(ins_rows) == 2 and ["X", "Y"] in ins_rows

    def test_multi_row_replacement_inserts_every_row(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        two_col_doc.propose_modify(
            "r1",
            '<tr data-aim="r1"><td>X</td><td>Y</td></tr>'
            '<tr data-aim="r1"><td>P</td><td>Q</td></tr>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(two_col_doc, tmp_path / "m.docx")
        ins_rows = self._ins_cells_per_row(out)
        first = ins_rows.index(["X", "Y"])
        assert ins_rows[first + 1] == ["P", "Q"]  # both rows, in payload order


class TestStructuralReplacementRecreatesSpans:
    """codex-r2-2 (refines AF-39): a shape-changing row modify created one
    unspanned Word cell per payload cell, so ``<td colspan="2">X</td>``
    became a single one-column cell and no accept sequence in Word could
    reach the proposed AIM grid. Inserted replacement rows now carry the
    payload's gridSpan/vMerge structure."""

    @pytest.fixture
    def two_col_doc(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl">'
            '<tbody><tr data-aim="r1"><td>A</td><td>B</td></tr></tbody></table>',
            author=ME,
            at=ts(0),
        )
        return doc

    @staticmethod
    def _inserted_rows(path):
        import docx
        from docx.oxml.ns import qn

        rows = docx.Document(str(path)).tables[0]._tbl.findall(qn("w:tr"))
        return [
            r
            for r in rows
            if (pr := r.find(qn("w:trPr"))) is not None and pr.find(qn("w:ins")) is not None
        ]

    def test_colspan_payload_cell_gets_a_gridspan(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        from docx.oxml.ns import qn

        two_col_doc.propose_modify(
            "r1", '<tr data-aim="r1"><td colspan="2">X</td></tr>', author=BOT, at=ts(1)
        )
        out = aim.to_docx(two_col_doc, tmp_path / "colspan.docx")
        (new_row,) = self._inserted_rows(out)
        (tc,) = new_row.findall(qn("w:tc"))  # one cell spanning the grid,
        span = tc.find(qn("w:tcPr")).find(qn("w:gridSpan"))  # not one column
        assert span is not None and span.get(qn("w:val")) == "2"

    def test_rowspan_payload_recreates_the_vertical_merge(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        from docx.oxml.ns import qn

        two_col_doc.propose_modify(
            "r1",
            '<tr data-aim="r1"><td rowspan="2">X</td><td>Y</td></tr>'
            '<tr data-aim="r1"><td>Z</td></tr>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(two_col_doc, tmp_path / "rowspan.docx")
        first, second = self._inserted_rows(out)

        top = first.findall(qn("w:tc"))[0].find(qn("w:tcPr")).find(qn("w:vMerge"))
        assert top is not None and top.get(qn("w:val")) == "restart"
        cont_tc = second.findall(qn("w:tc"))[0]
        cont = cont_tc.find(qn("w:tcPr")).find(qn("w:vMerge"))
        assert cont is not None and cont.get(qn("w:val")) in (None, "continue")
        assert "".join(t.text or "" for t in cont_tc.iter(qn("w:t"))) == ""
        # Z sits beside the continuation, in the second grid column
        tcs = second.findall(qn("w:tc"))
        texts = ["".join(t.text or "" for t in tc.iter(qn("w:t"))) for tc in tcs]
        assert len(tcs) == 2 and texts == ["", "Z"]

    def test_unspanned_payload_rows_stay_unspanned(self, two_col_doc, tmp_path):
        pytest.importorskip("docx")
        from docx.oxml.ns import qn

        two_col_doc.propose_modify(
            "r1", '<tr data-aim="r1"><td>X</td><td>Y</td><td>Z</td></tr>', author=BOT, at=ts(1)
        )
        out = aim.to_docx(two_col_doc, tmp_path / "plain.docx")
        (new_row,) = self._inserted_rows(out)
        tcs = new_row.findall(qn("w:tc"))
        assert len(tcs) == 3
        assert all(
            tc.find(qn("w:tcPr")) is None or tc.find(qn("w:tcPr")).find(qn("w:gridSpan")) is None
            for tc in tcs
        )


class TestTrackedTableContainerTracksItsStructure:
    """codex-r2-6 (refines codex-3/AF-39): a pending delete or modify of a
    TABLE container only wrapped each cell's text in run-level w:del — the
    rows carried no structural deletion, so accepting the revisions in Word
    left an empty grid behind (and a table modify left that grid alongside
    the inserted replacement). The container-level force passes now mark
    every row: del on the old grid, ins on an inserted one."""

    @pytest.fixture
    def table_doc(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl">'
            '<thead><tr data-aim="h"><th>K</th><th>V</th></tr></thead>'
            '<tbody><tr data-aim="r1"><td>A</td><td>B</td></tr></tbody></table>',
            author=ME,
            at=ts(0),
        )
        return doc

    @staticmethod
    def _row_marks(path, tag):
        import docx
        from docx.oxml.ns import qn

        out = []
        for tbl in docx.Document(str(path)).tables:
            rows = tbl._tbl.findall(qn("w:tr"))
            out.append(
                [
                    (pr := r.find(qn("w:trPr"))) is not None and pr.find(qn(tag)) is not None
                    for r in rows
                ]
            )
        return out

    def test_container_delete_marks_every_row_deleted(self, table_doc, tmp_path):
        pytest.importorskip("docx")
        table_doc.propose_delete("tbl", author=BOT, at=ts(1))
        out = aim.to_docx(table_doc, tmp_path / "del.docx")
        assert self._row_marks(out, "w:del") == [[True, True]]

    def test_container_modify_deletes_old_grid_and_inserts_the_new(self, table_doc, tmp_path):
        pytest.importorskip("docx")
        table_doc.propose_modify(
            "tbl",
            '<table data-aim-container="tbl"><tbody>'
            '<tr data-aim="n1"><td>X</td></tr></tbody></table>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(table_doc, tmp_path / "mod.docx")
        assert self._row_marks(out, "w:del") == [[True, True], [False]]
        assert self._row_marks(out, "w:ins") == [[False, False], [True]]

    def test_pending_table_add_marks_its_rows_inserted(self, tmp_path):
        pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="a">before</p>', author=ME, at=ts(0))
        doc.propose_add(
            '<table data-aim-container="t2"><tbody>'
            '<tr data-aim="n1"><td>X</td></tr></tbody></table>',
            author=BOT,
            container="body",
            after="a",
            at=ts(1),
        )
        out = aim.to_docx(doc, tmp_path / "add.docx")
        assert self._row_marks(out, "w:ins") == [[True]]

    def test_resolve_modes_still_export(self, table_doc, tmp_path):
        pytest.importorskip("docx")
        table_doc.propose_delete("tbl", author=BOT, at=ts(1))
        assert aim.to_docx(table_doc, tmp_path / "acc.docx", pending="accept-all").exists()
        assert aim.to_docx(table_doc, tmp_path / "rej.docx", pending="reject-all").exists()


class TestDocxSplitsGroupingBlocks:
    """AF-40: ``_block_children`` unpacked only section/slides, so a
    div/blockquote fell through whole to one ``add_paragraph`` whose
    ``_runs_of`` treated the nested ``<p>``s as inline marks —
    '<blockquote><p>Quote one.</p><p>Quote two.</p></blockquote>' exported
    as the single paragraph 'Quote one.Quote two.'."""

    def _paras(self, path):
        import docx

        return [(p.style.name, p.text) for p in docx.Document(str(path)).paragraphs]

    def test_blockquote_paragraphs_stay_separate(self, tmp_path):
        pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<blockquote data-aim="q"><p>Quote one.</p><p>Quote two.</p></blockquote>',
            author=ME,
            at=ts(0),
        )
        paras = self._paras(aim.to_docx(doc, tmp_path / "q.docx"))
        assert ("Quote", "Quote one.") in paras
        assert ("Quote", "Quote two.") in paras
        assert not any("Quote one.Quote two." in text for _, text in paras)

    def test_div_paragraphs_stay_separate(self, tmp_path):
        pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk('<div data-aim="d"><p>First.</p><p>Second.</p></div>', author=ME, at=ts(0))
        texts = [text for _, text in self._paras(aim.to_docx(doc, tmp_path / "d.docx"))]
        assert "First." in texts and "Second." in texts

    def test_div_direct_text_is_kept(self, tmp_path):
        pytest.importorskip("docx")
        doc = aim.new_document(title="T")
        doc.add_chunk('<div data-aim="d">Lead-in text<p>Body para.</p></div>', author=ME, at=ts(0))
        texts = [text for _, text in self._paras(aim.to_docx(doc, tmp_path / "m.docx"))]
        assert "Lead-in text" in texts and "Body para." in texts


class TestTrackedTableToListModify:
    """AF-41: the table container-modify branch re-emitted the payload via
    ``emit_table(new_el, force="ins")``, which early-returns when the root
    has no ``<tr>`` — so a legal table→list replacement produced no w:ins
    and no text at all."""

    def test_list_payload_is_inserted(self, table_doc, tmp_path):
        pytest.importorskip("docx")
        import docx
        from docx.oxml.ns import qn

        table_doc.propose_modify(
            "tbl",
            '<ul data-aim-container="tbl"><li data-aim="a">one</li><li data-aim="b">two</li></ul>',
            author=BOT,
            at=ts(1),
        )
        out = aim.to_docx(table_doc, tmp_path / "l.docx")
        d = docx.Document(str(out))
        inserted = "".join(
            t.text or "" for w in d.element.body.iter(qn("w:ins")) for t in w.iter(qn("w:t"))
        )
        assert "one" in inserted and "two" in inserted
        deleted = [t.text for t in d.element.body.iter(qn("w:delText"))]
        for text in ("K", "1", "2"):
            assert text in deleted


class TestCriticMarkupSurfacesPageSetupProposals:
    """AF-42: the criticmarkup triage excluded only ``aim:theme``, so a
    pending ``aim:doc`` (page-setup) modify landed in ``mods["aim:doc"]``
    — a key no body element ever matches — and appeared nowhere in the
    output; theme changes got the note branch."""

    def test_pending_page_setup_gets_a_note(self, basic_doc):
        basic_doc.propose_page_setup({"size": "A5"}, author=BOT, at=ts(5), explanation="Booklet.")
        md = aim.to_markdown(basic_doc, pending="criticmarkup")
        assert "aim:doc" in md and "Booklet." in md

    def test_theme_note_still_renders(self, basic_doc):
        basic_doc.propose_theme({"--aim-brand-1": "#000000"}, author=BOT, at=ts(5))
        md = aim.to_markdown(basic_doc, pending="criticmarkup")
        assert "aim:theme" in md


class TestCriticMarkupPayloadStructure:
    """AF-43: ``_payload_md`` rendered tr/li payload roots via ``_inline``,
    so the new side of ``{~~old~>new~~}`` fused a two-cell row into
    'gammadelta' and dropped list markers — accepting the suggestion
    corrupted the table/list."""

    def test_row_replacement_keeps_cell_boundaries(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl"><tbody>'
            '<tr data-aim="r1"><td>alpha</td><td>beta</td></tr></tbody></table>',
            author=ME,
            at=ts(0),
        )
        doc.propose_modify(
            "r1",
            '<tr data-aim="r1"><td>gamma</td><td>delta</td></tr>',
            author=BOT,
            at=ts(1),
        )
        md = aim.to_markdown(doc, pending="criticmarkup")
        assert "| gamma | delta |" in md
        assert "gammadelta" not in md

    def test_list_item_add_keeps_its_marker(self, rich_doc):
        rich_doc.propose_add(
            '<li data-aim="nli">Brand new item</li>',
            author=BOT,
            container="list",
            after="li1",
            at=ts(9),
        )
        md = aim.to_markdown(rich_doc, pending="criticmarkup")
        assert "{++- Brand new item++}" in md

    def test_ordered_list_item_add_keeps_its_numbered_marker(self):
        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<ol data-aim-container="steps">'
            '<li data-aim="one">First</li><li data-aim="two">Second</li>'
            "</ol>",
            author=ME,
            at=ts(0),
        )
        doc.propose_add(
            '<li data-aim="new">Inserted</li>',
            author=BOT,
            container="steps",
            after="one",
            at=ts(1),
        )

        critic = aim.to_markdown(doc, pending="criticmarkup")
        accepted = aim.to_markdown(doc, pending="accept-all")
        assert "{++2. Inserted++}" in critic
        assert "2. Inserted" in accepted


class TestMidParagraphBreakAnchorsOnItsOwnParagraph:
    """AF-44: a mid-paragraph page break was anchored on ``(prefix,
    seen+1)`` — assuming docling splits the paragraph at the break. It
    doesn't: the paragraph ingests whole, so the anchor matched whatever
    OTHER paragraph read exactly like the prefix and the break landed
    there. The break now anchors on the containing paragraph's completed
    text."""

    @pytest.fixture
    def broken_docx(self, tmp_path):
        docx = pytest.importorskip("docx")
        from docx.enum.text import WD_BREAK

        d = docx.Document()
        d.add_paragraph("Shared prefix")
        mid = d.add_paragraph()
        mid.add_run("Shared prefix")
        mid.add_run().add_break(WD_BREAK.PAGE)
        mid.add_run(" tail")
        d.add_paragraph("Shared prefix")
        path = tmp_path / "mid.docx"
        d.save(str(path))
        return path

    def _ingested(self):
        # the chunk shapes docling produces: paragraphs ingest whole
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Shared prefix</p>', author=ME, at=ts(0))
        doc.add_chunk('<p data-aim="p2">Shared prefix tail</p>', author=ME, at=ts(1))
        doc.add_chunk('<p data-aim="p3">Shared prefix</p>', author=ME, at=ts(2))
        return doc

    def test_anchor_is_the_full_paragraph_text(self, broken_docx):
        import docx

        from aimformat.convert._docx_pages import _read_break_anchors

        anchors = _read_break_anchors(docx.Document(str(broken_docx)))
        assert anchors == [("Shared prefix tail", 1)]

    def test_break_lands_after_the_containing_paragraph(self, broken_docx):
        from aimformat.convert._docx_pages import apply_docx_pagination

        doc = self._ingested()
        apply_docx_pagination(doc, broken_docx, author=ME)
        texts = [c.text for c in doc.chunks]
        assert texts == ["Shared prefix", "Shared prefix tail", "", "Shared prefix"]


class TestIngestKeepsRowspanCoveredRows:
    """AF-45: ``_table_markup`` dropped any physical grid row consisting
    entirely of rowspan continuations while the covering cells kept their
    span counts — the ``rowspan=2`` then swallowed the NEXT real row:
    lint-clean geometry corruption nothing downstream detects."""

    @staticmethod
    def _grid_cell(text, r, c, rs=1, cs=1):
        return {
            "text": text,
            "start_row_offset_idx": r,
            "start_col_offset_idx": c,
            "row_span": rs,
            "col_span": cs,
        }

    def test_fully_covered_row_emits_an_empty_tr(self):
        from aimformat.ingest import _table_markup

        a = self._grid_cell("A", 0, 0, rs=2)
        b = self._grid_cell("B", 0, 1, rs=2)
        grid = [[a, b], [a, b], [self._grid_cell("C", 2, 0), self._grid_cell("D", 2, 1)]]
        html = _table_markup({"data": {"grid": grid}})
        assert html.count("<tr") == 3
        assert "<tr></tr>" in html
        assert html.index("<tr></tr>") < html.index("<tr><td>C")

    def test_from_docling_geometry_survives(self):
        docling_core = pytest.importorskip("docling_core")  # noqa: F841
        from docling_core.types.doc import DoclingDocument, TableCell, TableData

        d = DoclingDocument(name="spans")
        td = TableData(
            num_rows=3,
            num_cols=2,
            table_cells=[
                TableCell(
                    text="A",
                    start_row_offset_idx=0,
                    end_row_offset_idx=2,
                    start_col_offset_idx=0,
                    end_col_offset_idx=1,
                    row_span=2,
                ),
                TableCell(
                    text="B",
                    start_row_offset_idx=0,
                    end_row_offset_idx=2,
                    start_col_offset_idx=1,
                    end_col_offset_idx=2,
                    row_span=2,
                ),
                TableCell(
                    text="C",
                    start_row_offset_idx=2,
                    end_row_offset_idx=3,
                    start_col_offset_idx=0,
                    end_col_offset_idx=1,
                ),
                TableCell(
                    text="D",
                    start_row_offset_idx=2,
                    end_row_offset_idx=3,
                    start_col_offset_idx=1,
                    end_col_offset_idx=2,
                ),
            ],
        )
        d.add_table(data=td)
        doc = aim.from_docling(d)
        assert not [f for f in aim.lint(doc) if f.level == "error"]
        rows = [c for c in doc.chunks if c.tag == "tr"]
        assert len(rows) == 3  # A/B row, covered row, C/D row
        assert rows[1].text == ""  # the covered row survives, empty
        assert "C" in rows[2].text and "D" in rows[2].text


class TestConsecutiveDocxBreaksAllSurvive:
    """AF-46: after the first insertion ``start = i+1``, so the second of
    two consecutive page breaks resolved to the same anchor and failed the
    ``hit[0] < start`` guard — two Ctrl+Enter breaks imported as one,
    deleting the intentional blank page."""

    def test_double_break_imports_as_two_chunks(self, tmp_path):
        docx = pytest.importorskip("docx")
        from docx.enum.text import WD_BREAK

        from aimformat.convert._docx_pages import apply_docx_pagination

        d = docx.Document()
        d.add_paragraph("Alpha")
        d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        d.add_paragraph("Beta")
        path = tmp_path / "double.docx"
        d.save(str(path))

        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Alpha</p>', author=ME, at=ts(0))
        doc.add_chunk('<p data-aim="p2">Beta</p>', author=ME, at=ts(1))
        apply_docx_pagination(doc, path, author=ME)
        texts = [c.text for c in doc.chunks]
        assert texts == ["Alpha", "", "", "Beta"]  # both breaks, in order


class TestSiblingAddOrderMatchesAcceptAll:
    """AF-47: resolution inserts every same-anchor add at index(anchor)+1,
    so accept-all leaves the LAST-proposed sibling closest to the anchor —
    but the tracked-DOCX paragraph lane and criticmarkup rendered proposal
    order. The two views of the same pending lane disagreed."""

    @pytest.fixture
    def two_adds(self):
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="p1">Anchor</p>', author=ME, at=ts(0))
        doc.propose_add('<p data-aim="a1">First card</p>', author=BOT, after="p1", at=ts(1))
        doc.propose_add('<p data-aim="a2">Second card</p>', author=BOT, after="p1", at=ts(2))
        return doc

    def test_accept_all_ground_truth(self, two_adds):
        md = aim.to_markdown(two_adds, pending="accept-all")
        assert md.index("Second card") < md.index("First card")

    def test_criticmarkup_matches_accept_all(self, two_adds):
        md = aim.to_markdown(two_adds, pending="criticmarkup")
        assert md.index("Second card") < md.index("First card")

    def test_tracked_docx_matches_accept_all(self, two_adds, tmp_path):
        docx = pytest.importorskip("docx")
        from docx.oxml.ns import qn

        out = aim.to_docx(two_adds, tmp_path / "o.docx")
        texts = [  # p.text does not see runs inside w:ins revisions
            "".join(t.text or "" for t in p._p.iter(qn("w:t")))
            for p in docx.Document(str(out)).paragraphs
        ]
        assert texts.index("Second card") < texts.index("First card")

    def test_chained_add_still_follows_its_parent(self, two_adds):
        first = next(p for p in two_adds.proposals if "First card" in (p.payload_html or ""))
        two_adds.propose_add(
            '<p data-aim="c1">Chained on first</p>', author=BOT, after=first.id, at=ts(3)
        )
        accept = aim.to_markdown(two_adds, pending="accept-all")
        critic = aim.to_markdown(two_adds, pending="criticmarkup")
        for md in (accept, critic):
            assert md.index("First card") < md.index("Chained on first")
            assert md.index("Second card") < md.index("First card")


class TestRowAddsAnchorAfterTheWholeRunChunk:
    """AF-48: ``emit_table`` drained row-adds per physical ``tr``, so the
    FIRST member of a run chunk popped the proposal and the added row
    landed mid-run — splitting the chunk in tracked DOCX."""

    def test_added_row_lands_after_the_last_run_member(self, tmp_path):
        docx = pytest.importorskip("docx")
        from docx.oxml.ns import qn

        doc = aim.new_document(title="T")
        doc.add_chunk(
            '<table data-aim-container="tbl"><tbody>'
            '<tr data-aim="rr"><td>one</td></tr>'
            '<tr data-aim="rr"><td>two</td></tr></tbody></table>',
            author=ME,
            at=ts(0),
        )
        doc.propose_add(
            '<tr data-aim="nr"><td>new</td></tr>',
            author=BOT,
            container="tbl",
            after="rr",
            at=ts(1),
        )
        out = aim.to_docx(doc, tmp_path / "r.docx")
        tbl = docx.Document(str(out)).tables[0]
        row_texts = [
            "".join(t.text or "" for t in tr.iter(qn("w:t"))) for tr in tbl._tbl.findall(qn("w:tr"))
        ]
        assert row_texts == ["one", "two", "new"]


class TestDirectModifyOfReservedTargets:
    """AF-08: ``modify_chunk`` on ``aim:theme``/``aim:doc`` fell through
    ``_normalize_payload``'s generic funnel, which stamped
    ``data-aim="aim:theme"`` onto the head <style>/<script> — persisted,
    hashed, lint-clean; the propose/set_theme/accept paths all routed
    around it."""

    def test_theme_modify_keeps_the_reserved_grammar(self, basic_doc):
        basic_doc.modify_chunk(
            "aim:theme",
            "<style data-aim-theme>:root{--aim-brand-1: #123456}</style>",
            author=ME,
            at=ts(5),
        )
        text = basic_doc.dumps()
        assert 'data-aim="aim:theme"' not in text
        assert basic_doc.theme.get("--aim-brand-1") == "#123456"
        assert basic_doc.verify() == []
        assert not [f for f in aim.lint_text(text) if f.level == "error"]

    def test_doc_modify_keeps_the_reserved_grammar(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        basic_doc.modify_chunk(
            "aim:doc",
            '<script type="application/aim-doc+json">{"page": {"size": "Letter"}}</script>',
            author=ME,
            at=ts(6),
        )
        assert 'data-aim="aim:doc"' not in basic_doc.dumps()
        assert basic_doc.page_setup.size == "Letter"
        assert basic_doc.verify() == []

    def test_non_theme_payload_is_rejected(self, basic_doc):
        with pytest.raises(InvalidOperation):
            basic_doc.modify_chunk("aim:theme", "<p>not a theme block</p>", author=ME, at=ts(5))


class TestPageSettingsPreserveNestedExtensions:
    """Codex-2: whole aim:doc validation resolved ``page`` through
    PageSetup.to_obj(), silently dropping forward-compatible nested data."""

    MARKUP = (
        '<script type="application/aim-doc+json">'
        '{"page":{"margins":{"top":"20mm","x_gutter":"5mm"},'
        '"size":"Letter","x_duplex":"long-edge"}}'
        "</script>"
    )

    @staticmethod
    def assert_extensions_survive(doc):
        page = doc.doc_settings["page"]
        assert page["x_duplex"] == "long-edge"
        assert page["margins"]["x_gutter"] == "5mm"
        assert doc.page_setup.size == "Letter"
        assert doc.page_setup.margins_mm["top"] == 20.0
        assert doc.verify() == []

    def test_direct_modify_preserves_extensions(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        basic_doc.modify_chunk("aim:doc", self.MARKUP, author=ME, at=ts(6))
        self.assert_extensions_survive(basic_doc)

    def test_amend_preserves_extensions(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        proposal = basic_doc.propose_page_setup({"size": "A4"}, author=BOT, at=ts(6))
        basic_doc.amend_proposal(proposal.id, self.MARKUP, at=ts(7))
        basic_doc.accept(proposal.id, decided_by=ME, at=ts(8))
        self.assert_extensions_survive(basic_doc)

    def test_accept_with_tweaks_preserves_extensions(self, basic_doc):
        basic_doc.set_page_setup({"size": "A5"}, author=ME, at=ts(5))
        proposal = basic_doc.propose_page_setup({"size": "A4"}, author=BOT, at=ts(6))
        basic_doc.accept(proposal.id, decided_by=ME, applied=self.MARKUP, at=ts(7))
        self.assert_extensions_survive(basic_doc)


def _tiny_png(width, height):
    """A minimal opaque RGBA PNG of the given size, stdlib-only."""
    import struct
    import zlib

    def chunk(tag, data):
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body))

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x00\x00\x00\xff" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


class TestPackAssetsKeepsImageFidelity:
    """AF-09: pack_assets froze a hardcoded 100×100 viewBox/image size into
    every symbol (letterboxing any non-square image) and rebuilt the inline
    ``<svg>`` from scratch, dropping the img's class/title and the other
    registry-global attributes — both losses recorded forever by the modify
    event."""

    def _doc_with_png(self, width, height, extra_attrs=""):
        import base64 as b64

        uri = "data:image/png;base64," + b64.b64encode(_tiny_png(width, height)).decode()
        doc = aim.new_document(title="T")
        doc.add_chunk(
            f'<figure data-aim="fig"><img alt="dot"{extra_attrs} src="{uri}"></figure>',
            author=ME,
            at=ts(0),
        )
        return doc

    def test_intrinsic_dimensions_reach_the_symbol(self):
        doc = self._doc_with_png(3, 2)
        doc.pack_assets(author=aim.external("packer"), at=ts(1))
        text = doc.dumps()
        assert 'viewBox="0 0 3 2"' in text
        assert 'viewBox="0 0 100 100"' not in text

    def test_global_attributes_survive_packing(self):
        doc = self._doc_with_png(3, 2, extra_attrs=' class="border" title="A dot"')
        doc.pack_assets(author=aim.external("packer"), at=ts(1))
        html = doc.chunk("fig").html
        assert 'class="border"' in html and 'title="A dot"' in html
        assert not [f for f in aim.lint_text(doc.dumps()) if f.level == "error"]
        assert doc.verify() == []

    def test_unparseable_blob_falls_back_to_square(self):
        assert aim.AimDocument._image_dimensions(b"not an image") is None


class TestPackKeepsOneAxisStyledAspect:
    """codex-r2-1 (refines AF-09): the packed wrapper svg copied a one-axis
    style (only ``width:`` or only ``height:``) while still pinning BOTH
    presentation attributes to the intrinsic size — the styled axis followed
    the style but the free axis stayed fixed, so a 4×2 image styled
    ``width:100px`` rendered as a distorted 100×2. The wrapper now carries
    the intrinsic ratio as a viewBox and leaves the unstyled axis auto."""

    def _packed_svg(self, extra_attrs=""):
        import base64 as b64

        uri = "data:image/png;base64," + b64.b64encode(_tiny_png(4, 2)).decode()
        doc = aim.new_document(title="T")
        doc.add_chunk(
            f'<figure data-aim="fig"><img alt="dot"{extra_attrs} src="{uri}"></figure>',
            author=ME,
            at=ts(0),
        )
        doc.pack_assets(author=aim.external("packer"), at=ts(1))
        svg = re.search(r"<svg[^>]*>", doc.chunk("fig").html)
        assert svg is not None
        return doc, svg.group(0)

    def test_width_only_style_leaves_height_auto(self):
        doc, svg = self._packed_svg(' style="width:100px"')
        assert 'viewBox="0 0 4 2"' in svg
        assert "height=" not in svg
        assert 'width="4"' in svg  # intrinsic fallback; the style overrides it
        assert not [f for f in aim.lint_text(doc.dumps()) if f.level == "error"]
        assert doc.verify() == []

    def test_height_only_style_leaves_width_auto(self):
        doc, svg = self._packed_svg(' style="height:50px"')
        assert 'viewBox="0 0 4 2"' in svg
        assert "width=" not in svg
        assert 'height="2"' in svg

    def test_unstyled_image_still_pins_intrinsic_size(self):
        _, svg = self._packed_svg()
        assert 'width="4"' in svg and 'height="2"' in svg
        assert 'viewBox="0 0 4 2"' in svg

    def test_both_axes_styled_keep_the_intrinsic_attrs(self):
        _, svg = self._packed_svg(' style="width:60px; height:30px"')
        assert 'width="4"' in svg and 'height="2"' in svg  # fully author-sized anyway


class TestTweaksKeepTheCardsOwnNestedIds:
    """AF-10: ``_taken_ids()`` includes the card's own payload ids, so a
    text-only amend (or accept-with-tweaks) of a pending add reminted an
    unchanged nested ``child`` id to a random one — recorded as if a human
    renamed it; only the root escaped via ``payload_id == expect_id``."""

    PAYLOAD = '<ul data-aim-container="nlist"><li data-aim="child">{text}</li></ul>'

    def test_amend_keeps_nested_ids(self, basic_doc):
        p = basic_doc.propose_add(self.PAYLOAD.format(text="old"), author=BOT, at=ts(5))
        basic_doc.amend_proposal(p.id, self.PAYLOAD.format(text="new"), at=ts(6))
        amended = basic_doc.proposal(p.id)
        assert 'data-aim="child"' in amended.payload_html
        assert "new" in amended.payload_html

    def test_accept_with_tweaks_keeps_nested_ids(self, basic_doc):
        p = basic_doc.propose_add(self.PAYLOAD.format(text="old"), author=BOT, at=ts(5))
        basic_doc.accept(p.id, decided_by=ME, applied=self.PAYLOAD.format(text="tweaked"), at=ts(6))
        assert basic_doc.chunk("child").text == "tweaked"
        assert basic_doc.verify() == []


class TestShellIdentityOnEveryTableAnchor:
    """AF-12: destination anchors carried a shell only in the ``after is
    None`` branch, no-op checks compared just (container, after), and
    ``move_chunk`` exposed no shell — distinct first positions in
    thead/tbody collapsed into one 'already at that position' and a
    first-of-thead move was inexpressible."""

    def test_concrete_row_anchor_records_its_shell(self, table_doc):
        table_doc.add_chunk(
            '<tr data-aim="r3"><td>3</td></tr>', author=ME, container="tbl", after="r1", at=ts(1)
        )
        assert table_doc.history[-1].get("anchor") == {
            "after": "r1",
            "container": "tbl",
            "shell": "tbody",
        }

    def test_first_of_thead_move_is_expressible(self, table_doc):
        table_doc.move_chunk("r1", author=ME, container="tbl", after=None, shell="thead", at=ts(1))
        serial = table_doc._state.serial("tbl")
        head = serial.index("<thead>"), serial.index("</thead>")
        assert 'data-aim="r1"' in serial[head[0] : head[1]]
        assert table_doc.verify() == []

    def test_same_shell_first_position_is_still_a_noop(self, table_doc):
        with pytest.raises(InvalidOperation):
            table_doc.move_chunk("r1", author=ME, container="tbl", after=None, at=ts(1))

    def test_wrong_shell_for_a_concrete_anchor_is_rejected(self, table_doc):
        with pytest.raises(InvalidOperation):
            table_doc.move_chunk(
                "r2", author=ME, container="tbl", after="r1", shell="thead", at=ts(1)
            )

    def test_propose_move_distinguishes_shells(self, table_doc):
        p = table_doc.propose_move(
            "r1", author=BOT, container="tbl", after=None, shell="thead", at=ts(1)
        )
        table_doc.accept(p.id, decided_by=ME, at=ts(2))
        serial = table_doc._state.serial("tbl")
        assert serial.index('data-aim="r1"') < serial.index("</thead>")
        assert table_doc.verify() == []


class TestReconcileRejectsSchemaInvalidHistory:
    """AF-13: ``_check_log`` validated seq shape and known kinds but never
    ``Event.validate()`` — a history the linter itself flags (H003, e.g. a
    deleted ``author``) counted as an intact baseline and reconcile happily
    'repaired' on top of corrupt provenance."""

    def test_missing_author_fails_closed(self, basic_doc):
        text = basic_doc.dumps()
        corrupted = re.sub(r'"author":\{[^}]*\},', "", text, count=1)
        assert "H003" in {f.code for f in aim.lint_text(corrupted)}  # linter agrees
        corrupted = corrupted.replace("Intro paragraph.</p>", "Intro paragraph, edited.</p>")
        doc = aim.AimDocument.loads(corrupted)
        with pytest.raises(aim.HistoryError):
            doc.reconcile(at=ts(60))

    def test_intact_history_still_reconciles(self, basic_doc):
        text = basic_doc.dumps().replace("Intro paragraph.</p>", "Intro paragraph, edited.</p>")
        doc = aim.AimDocument.loads(text)
        report = doc.reconcile(at=ts(60))
        assert report.changed and report.residual == []


class TestEventValidationIsTypeStrict:
    """AF-14: ``isinstance(seq, int)`` let ``"seq": true`` through (bool ⊂
    int in Python) and the digit-shape-only timestamp regex accepted
    impossible instants like 2026-99-99T99:99:99Z — all the way through
    validate(), verify(), and lint_text. An independent implementation
    doing real validation rejects the same bytes."""

    @staticmethod
    def _event(**overrides):
        from aimformat.events import Event

        data = {
            "seq": 1,
            "kind": "direct_edit",
            "t": ts(0),
            "target": "x",
            "action": "modify",
            "before": '<p data-aim="x">a</p>',
            "after": '<p data-aim="x">b</p>',
            "author": {"type": "human", "id": "luca"},
            "batch": "b-1",
        }
        data.update(overrides)
        return Event(data)

    def test_boolean_seq_is_rejected(self):
        assert any("seq" in p for p in self._event(seq=True).validate())

    def test_zero_and_negative_seq_are_rejected(self):
        assert self._event(seq=0).validate()
        assert self._event(seq=-3).validate()

    def test_impossible_timestamp_is_rejected(self):
        assert any("t is not" in p for p in self._event(t="2026-99-99T99:99:99Z").validate())

    def test_real_utc_instant_still_passes(self):
        assert self._event().validate() == []
        assert self._event(t="2024-02-29T23:59:59Z").validate() == []  # leap day


class TestChunkModifyRequiresBefore:
    """AF-15: a chunk ``modify`` event needed only ``after``, so with no
    ``before`` the inverse walk fell through silently and ``state_at``
    reconstructed the post-modify content as the past — fabricated history
    with no diagnostic once the original add was pruned."""

    def _doc_with_strippable_modify(self, basic_doc):
        basic_doc.modify_chunk("intro", '<p data-aim="intro">Rewritten.</p>', author=ME, at=ts(5))
        text = basic_doc.dumps()
        line_re = re.compile(r'\{[^\n]*"action":"modify"[^\n]*"target":"intro"[^\n]*\}')
        line = line_re.search(text).group(0)
        stripped = re.sub(r'"before":"(?:[^"\\]|\\.)*",', "", line, count=1)
        return text.replace(line, stripped)

    def test_validate_flags_the_stripped_event(self, basic_doc):
        from aimformat.events import Event

        text = self._doc_with_strippable_modify(basic_doc)
        doc = aim.AimDocument.loads(text)
        ev = next(e for e in doc.history if e.action == "modify" and e.target == "intro")
        assert Event(ev.data).validate()  # non-empty: 'before' is required

    def test_lint_reports_the_gap(self, basic_doc):
        text = self._doc_with_strippable_modify(basic_doc)
        assert "H003" in {f.code for f in aim.lint_text(text)}


class TestAssetGcIsWiredIn:
    """AF-17: ``gc_assets`` was specced (§9.3) as the final pass of
    pack/flatten/prune but called only from a test, and the CLI had no
    pack/prune/gc verbs at all — a flattened 'clean file' shipped every
    dead multi-hundred-KB data-URI forever."""

    def _packed_then_deleted(self):
        import base64 as b64

        uri = "data:image/png;base64," + b64.b64encode(_tiny_png(2, 2)).decode()
        doc = aim.new_document(title="T")
        doc.add_chunk('<p data-aim="keep">Kept text.</p>', author=ME, at=ts(0))
        doc.add_chunk(
            f'<figure data-aim="fig"><img alt="dot" src="{uri}"></figure>',
            author=ME,
            at=ts(1),
        )
        doc.pack_assets(author=aim.external("packer"), at=ts(2))
        doc.delete_chunk("fig", author=ME, at=ts(3))
        return doc

    def test_flatten_collects_dead_assets(self):
        doc = self._packed_then_deleted()
        doc.flatten()
        assert "<aim-assets>" not in doc.dumps()

    def test_prune_collects_dead_assets(self):
        doc = self._packed_then_deleted()
        doc.modify_chunk("keep", '<p data-aim="keep">Kept, edited.</p>', author=ME, at=ts(4))
        last = doc.history[-1].seq
        doc.prune(before=last)  # cut above the delete that referenced it
        assert "<aim-assets>" not in doc.dumps()

    def test_cli_pack_and_gc_verbs_exist(self, tmp_path, capsys):
        import base64 as b64

        from aimformat.cli import main

        uri = "data:image/png;base64," + b64.b64encode(_tiny_png(2, 2)).decode()
        doc = aim.new_document(title="T")
        doc.add_chunk(
            f'<figure data-aim="fig"><img alt="dot" src="{uri}"></figure>',
            author=ME,
            at=ts(0),
        )
        path = tmp_path / "doc.aim"
        doc.save(path)
        assert main(["pack", str(path)]) == 0
        packed = aim.AimDocument.load(path)
        assert "<aim-assets>" in packed.dumps()
        assert "data:image" not in packed.chunk("fig").html
        assert main(["gc", str(path)]) == 0  # nothing dead: a no-op
        assert "<aim-assets>" in aim.AimDocument.load(path).dumps()


class TestTemplatelessAddFailsAsAimError:
    """AF-18: ``_payload_like`` indexed ``orig_nodes[0]`` unguarded, so
    accept-with-tweaks (or amend) on a hand-authored template-less add
    card (P006) raised a bare IndexError — not an AimError, so it escaped
    the MCP boundary's ``except AimError``."""

    @pytest.fixture
    def templateless(self, basic_doc):
        basic_doc.propose_add('<p data-aim="a1">new</p>', author=BOT, at=ts(5))
        text = re.sub(r"<template>.*?</template>", "", basic_doc.dumps(), flags=re.S)
        doc = aim.AimDocument.loads(text)
        assert "P006" in {f.code for f in aim.lint_text(text)}
        return doc

    def test_accept_with_tweaks_raises_invalid_operation(self, templateless):
        with pytest.raises(InvalidOperation):
            templateless.accept(
                templateless.proposals[0].id,
                decided_by=ME,
                applied='<p data-aim="a1">tweak</p>',
                at=ts(6),
            )

    def test_amend_raises_invalid_operation(self, templateless):
        with pytest.raises(InvalidOperation):
            templateless.amend_proposal(
                templateless.proposals[0].id, '<p data-aim="a1">amended</p>', at=ts(6)
            )


class TestDuplicateAttributesResolveFirstWins:
    """AF-19: the reader resolves duplicate attributes first-wins (HTML
    semantics) but the canonical serializer built its map last-wins — on
    '<p data-aim="a" data-aim="b">' the document answered to 'a' while
    dumps()/normalize emitted 'b', renaming the chunk and breaking every
    event that targeted 'a'."""

    def test_serializer_matches_the_reader(self):
        from aimformat.canonical import serialize
        from aimformat.dom import parse_fragment

        el = parse_fragment('<p data-aim="a" data-aim="b">x</p>')[0]
        assert el.chunk_id == "a"  # the reader's resolution
        out = serialize(el)
        assert 'data-aim="a"' in out and 'data-aim="b"' not in out

    def test_normalize_keeps_the_resolved_id(self, basic_doc):
        text = basic_doc.dumps().replace(
            '<p data-aim="intro">', '<p data-aim="intro" data-aim="other">'
        )
        doc = aim.AimDocument.loads(text)
        assert 'data-aim="intro"' in doc.dumps()
        assert 'data-aim="other"' not in doc.dumps()


class TestWrongMarkerStrippedFromEveryRunMember:
    """AF-20: the honored-id branch of ``_normalize_payload`` gated the
    wrong-marker cleanup on the FIRST run member only — a dual-marked
    second member survived, and a single public ``add_chunk`` wrote a
    document its own linter rejects (S025/S019/V003/H006)."""

    def test_dual_marked_second_member_is_cleaned(self, rich_doc):
        rich_doc.add_chunk(
            '<li data-aim="x9">part one</li>'
            '<li data-aim="x9" data-aim-container="x9">part two</li>',
            author=ME,
            container="list",
            at=ts(9),
        )
        text = rich_doc.dumps()
        assert 'data-aim-container="x9"' not in text
        assert not [f for f in aim.lint_text(text) if f.level == "error"]
        assert rich_doc.verify() == []


class TestMcpWarnsWhenUnscoped:
    """AF-22: with AIMFORMAT_MCP_ROOT unset (the default) the MCP server
    reaches any host path silently. The unscoped default stays — the
    documented local-trusted-stdio contract — but startup now warns loudly
    so wiring it to a semi-trusted client is a visible decision."""

    def test_unscoped_startup_warns(self, monkeypatch, capsys):
        mcp_mod = pytest.importorskip("aimformat.mcp")
        monkeypatch.delenv("AIMFORMAT_MCP_ROOT", raising=False)
        assert mcp_mod._warn_if_unscoped() is True
        assert "AIMFORMAT_MCP_ROOT" in capsys.readouterr().err

    def test_scoped_startup_is_quiet(self, monkeypatch, capsys, tmp_path):
        mcp_mod = pytest.importorskip("aimformat.mcp")
        monkeypatch.setenv("AIMFORMAT_MCP_ROOT", str(tmp_path))
        assert mcp_mod._warn_if_unscoped() is False
        assert capsys.readouterr().err == ""


_HISTORY_OPEN = '<script type="application/aim-history+jsonl">'


def _before_history(text, insert):
    """Insert a body section right before the history script (keeps the
    content → proposals → assets → history order valid)."""
    return text.replace(_HISTORY_OPEN, insert + "\n" + _HISTORY_OPEN, 1)


def _card(attrs):
    return (
        f"<aim-proposals><aim-proposal {attrs}>"
        "<template><p>x</p></template></aim-proposal></aim-proposals>"
    )


_ADD_ATTRS = (
    'id="p-aaaaaaaa" data-action="add" data-anchor-container="body" '
    'data-at="2026-07-07T10:00:05Z" data-author="agent" data-author-model="m"'
)
_MOD_ATTRS = (
    'id="p-aaaaaaaa" data-action="modify" '
    'data-at="2026-07-07T10:00:05Z" data-author="agent" data-author-model="m"'
)

_RULE_PROBES = {
    "H002": lambda t: t.replace(_HISTORY_OPEN, _HISTORY_OPEN + "\nnot json", 1),
    "P001": lambda t: _before_history(t, "<aim-proposals><div>bad</div></aim-proposals>"),
    "P002": lambda t: _before_history(t, _card(_ADD_ATTRS.replace("p-aaaaaaaa", "bad id"))),
    "P003": lambda t: _before_history(t, _card(_ADD_ATTRS.replace("add", "frobnicate"))),
    "P004": lambda t: _before_history(t, _card(_MOD_ATTRS)),  # modify without data-for
    "P005": lambda t: _before_history(
        t, _card(_ADD_ATTRS.replace(' data-anchor-container="body"', ""))
    ),
    "P012": lambda t: _before_history(t, _card(_ADD_ATTRS + ' data-depends-on="p-00000000"')),
    "P013": lambda t: _before_history(
        t, _card(_ADD_ATTRS.replace('data-at="2026-07-07T10:00:05Z"', 'data-at="yesterday"'))
    ),
    "S006": lambda t: re.sub(
        r'<style data-aim-css="[^"]*">.*?</style>',
        '<style data-aim-css="0.1">/* aim.css placeholder */</style>',
        t,
        flags=re.S,
    ),
    "S010": lambda t: _before_history(t, '<script type="application/weird+json">{}</script>'),
    "S013": lambda t: t.replace(
        "</body>", _HISTORY_OPEN + "\n</script>\n</body>", 1
    ),  # a second rank-3 history section
    "S018": lambda t: _before_history(
        t,
        '<ul data-aim-container="dupc"><li data-aim="lia">x</li></ul>'
        '<ul data-aim-container="dupc"><li data-aim="lib">y</li></ul>',
    ),
    "S019": lambda t: _before_history(
        t, '<ul data-aim-container="intro"><li data-aim="lic">y</li></ul>'
    ),  # 'intro' is already a chunk id
    "S022": lambda t: _before_history(
        t, '<ul data-aim-container="badlist"><p data-aim="pbad">x</p></ul>'
    ),
    "S026": lambda t: _before_history(
        t,
        '<aim-slide data-aim-container="sl1" style="width:420px; height:595px">'
        '<aim-slide data-aim-container="sl2" style="width:420px; height:595px">'
        "</aim-slide></aim-slide>",
    ),
    "S027": lambda t: t.replace(
        "</title>",
        '</title>\n<script type="application/aim-meta+json">{}</script>'
        '\n<script type="application/aim-meta+json">{}</script>',
        1,
    ),
    "V001": lambda t: _before_history(t, "<aim-assets><div>x</div></aim-assets>"),
    "V006": lambda t: t.replace('<p data-aim="intro">', '<p data-aim="intro" style="left">', 1),
}


class TestUncoveredLintRulesFire:
    """AF-24: 18 of 83 lint rule codes fired in no test — the only pin
    regex-grepped the code strings out of lint.py source, so a refactor
    that drops or inverts a check while the string survives passes CI
    (S022 had already rotted exactly that way). One firing probe per
    previously-uncovered code."""

    @pytest.mark.parametrize("code", sorted(_RULE_PROBES))
    def test_rule_fires(self, code, basic_doc):
        text = _RULE_PROBES[code](basic_doc.dumps())
        assert code in {f.code for f in aim.lint_text(text)}

    def test_the_probe_base_is_clean(self, basic_doc):
        assert aim.lint_text(basic_doc.dumps()) == []


class TestEveryWorkflowActionIsShaPinned:
    """AF-25: publish.yml — the OIDC→PyPI workflow, the repo's highest
    privilege — was the only workflow with floating action refs; a moved
    tag (tj-actions 2025 precedent) could publish a trojaned package with
    zero repo changes. Every `uses:` in every workflow must pin a full
    commit SHA."""

    def test_all_uses_refs_are_full_shas(self):
        import pathlib

        workflows = sorted(
            (pathlib.Path(__file__).parent.parent / ".github" / "workflows").glob("*.yml")
        )
        assert workflows, "no workflows found"
        offenders = []
        for wf in workflows:
            for line in wf.read_text().splitlines():
                stripped = line.strip()
                if not stripped.startswith("- uses:") and not stripped.startswith("uses:"):
                    continue
                ref = stripped.split("uses:", 1)[1].split("#")[0].strip()
                _, _, pin = ref.partition("@")
                if not re.fullmatch(r"[0-9a-f]{40}", pin):
                    offenders.append(f"{wf.name}: {ref}")
        assert offenders == []
